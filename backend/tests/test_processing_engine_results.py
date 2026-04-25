from pathlib import Path

from docx import Document
from pypdf import PdfReader
import pytest
from sqlalchemy.orm import Session

from app.models import SystemConfig
from app.models import TaskType
from app.services.dedup_strategies.executor import execute_dedup_strategy
from app.services.dedup_strategies.cnki_llm import _prompt as cnki_dedup_prompt
from app.services.dedup_strategies.vip_llm import _prompt as vip_dedup_prompt
from app.services.dedup_strategies.validators import validate_dedup_output
from app.services.processing_engine import ProcessingEngine
from app.services.rewrite_strategies.cnki_llm import _prompt as cnki_rewrite_prompt
from app.services.rewrite_strategies.vip_llm import _prompt as vip_rewrite_prompt
from app.services.rewrite_strategies.validators import validate_rewrite_output


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def test_cnki_llm_rewrite_uses_global_prompt_by_default(db_session: Session, monkeypatch) -> None:
    from app.services.rewrite_strategies import cnki_llm

    text = "因此，本研究在方法层面进行系统分析并提出改进路径，" * 24
    calls: list[str] = []

    def _fake_generate(_db, *, task_type, text: str):
        calls.append(text)
        if "输出JSON" in text:
            if "对比原文与改写文" in text:
                return (
                    '{"semantic_ok": true, "grammar_ok": true, "style_ok": true, '
                    '"compound_ok": true, "density": "12%", "density_ok": true, '
                    '"operation_counts": {"A功能词": 2, "B整词同义": 3, "C连接词": 1, "D句法框架": 1, "E副词": 1, "F元话语": 0}, '
                    '"issues": [], "verdict": "pass"}'
                )
        marker = "输入文本：\n"
        return text.split(marker, 1)[1].strip() if marker in text else text

    monkeypatch.setattr("app.services.rewrite_strategies.cnki_llm.generate_with_llm", _fake_generate)

    result = cnki_llm.rewrite(db_session, task=None, text=text, report_summary={})
    trace = result.get("rule_trace") or {}

    assert trace.get("chunk_count", 0) == 1
    assert trace.get("mode") == "llm_prompt_ab_strict_v16_global"
    assert trace.get("technical_chunking") is False
    assert trace.get("strategy_version") == "cnki_v16_llm_ab_strict"
    assert any("双遍执行架构" in prompt for prompt in calls)
    assert len(calls) == 2


def test_rewrite_process_uses_report_summary(tmp_path: Path, db_session: Session, monkeypatch) -> None:
    source_path = tmp_path / "paper.docx"
    report_path = tmp_path / "report.docx"
    output_path = tmp_path / "result.docx"

    _write_docx(
        source_path,
        [
            "研究表明，这一方法非常重要，而且很多场景都可以看出类似趋势。",
            "这个结论在教学和管理实践中具有较强参考价值。",
        ],
    )
    _write_docx(
        report_path,
        [
            "全文AIGC检测报告",
            "总体风险 52%",
            "高风险段落占比 24%",
        ],
    )

    monkeypatch.setattr(ProcessingEngine, "_run_llm", lambda self, *_args, **_kwargs: None)
    monkeypatch.setattr(
        "app.services.rewrite_strategies.cnki_llm.generate_with_llm",
        lambda _db, *, task_type, text: (
            '{"semantic_ok": true, "grammar_ok": true, "style_ok": true, '
            '"compound_ok": true, "density": "9%", "density_ok": true, '
            '"operation_counts": {"A功能词": 1, "B整词同义": 2, "C连接词": 1, "D句法框架": 1, "E副词": 0, "F元话语": 0}, '
            '"issues": [], "verdict": "pass"}'
            if "输出JSON" in str(text or "")
            else "研究表明，这一方法较为关键，同时多个场景能够呈现出相近趋势。这个结论在教学和管理实践中具有参考价值。"
        ),
    )

    engine = ProcessingEngine(db_session)
    result = engine.process(
        TaskType.REWRITE,
        "cnki",
        source_path,
        output_path,
        task_id=1,
        report_path=report_path,
    )

    assert output_path.exists()
    assert result.result_json["type"] == "rewrite"
    assert result.result_json["report_summary"]["available"] is True
    assert result.result_json["report_summary"]["pressure"] == "high"
    assert len(result.result_json["report_summary"]["metrics"]) >= 1
    assert len(result.result_json["review_points"]) >= 1
    assert isinstance(result.result_json["output_preview"], str)
    assert result.result_json["source_stats"]["char_count"] > 0
    assert result.result_json["output_stats"]["char_count"] > 0


def test_aigc_detect_returns_structured_result(tmp_path: Path, db_session: Session, monkeypatch) -> None:
    source_path = tmp_path / "paper.txt"
    output_path = tmp_path / "result.pdf"
    source_path.write_text(
        "本研究围绕教学改革展开讨论。"
        "研究表明该方案在多个学院中具有稳定的执行路径，因此能够快速复制。"
        "\n"
        "此外，文章采用一致的段落结构和重复性的总结语句，这会提高自动化写作特征。",
        encoding="utf-8",
    )

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.AIGC_DETECT, "cnki", source_path, output_path, task_id=2)

    assert output_path.exists()
    assert result.result_json["type"] == "aigc_detect"
    assert "summary" in result.result_json
    assert "risk_band" in result.result_json
    assert result.result_json["simulation_profile"] == "cnki_internal"
    assert result.result_json["provider_label"] == "知网AIGC检测仿真"
    assert result.result_json["source_stats"]["char_count"] > 0
    assert len(result.result_json["paragraph_details"]) >= 1
    assert "distribution" in result.result_json
    assert "fragment_distribution" in result.result_json
    assert "document_metrics" in result.result_json
    assert "decision_basis" in result.result_json
    assert result.result_json["llm_used"] is False
    assert "algo_package_used" not in result.result_json
    assert result.result_json["score_breakdown"]["strategy"] == "algorithm"
    assert "llm_score" not in result.result_json["score_breakdown"]
    assert "algo_package_score" not in result.result_json["score_breakdown"]
    content = output_path.read_bytes()
    assert content.startswith(b"%PDF-")
    assert b"/STSong-Light" in content


def test_aigc_detect_platform_profiles_are_close_but_not_identical(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "paper_platform.txt"
    source_path.write_text(
        "本研究基于教学管理数据，围绕课程评价机制展开分析。\n"
        "文章包含连续论证句、固定连接词和重复总结表达，用于模拟AIGC风险检测场景。\n"
        "在不同平台规则下，分值应保持相近，但因评分偏好不同会存在轻微差异。",
        encoding="utf-8",
    )
    engine = ProcessingEngine(db_session)
    results = {}
    for platform in ("cnki", "vip"):
        output_path = tmp_path / f"{platform}.pdf"
        result = engine.process(TaskType.AIGC_DETECT, platform, source_path, output_path, task_id=100)
        results[platform] = float(result.result_json["score_pct"])
        assert output_path.exists()

    assert len(set(round(value, 2) for value in results.values())) >= 2
    assert max(results.values()) - min(results.values()) <= 15


def test_rewrite_process_for_common_platform_uses_internal_heuristic_fallback(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "rewrite.txt"
    output_path = tmp_path / "rewrite_out.txt"
    source_path.write_text("研究表明，因此需要持续优化。", encoding="utf-8")

    monkeypatch.setattr(ProcessingEngine, "_run_llm", lambda self, *_args, **_kwargs: None)

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.REWRITE, "paperpass", source_path, output_path, task_id=3)

    assert output_path.exists()
    assert output_path.read_text(encoding="utf-8")
    assert result.result_json["type"] == "rewrite"
    assert result.result_json["output_stats"]["char_count"] > 0


def test_aigc_detect_uses_internal_strategy_chain_without_external_scores(tmp_path: Path, db_session: Session) -> None:
    source_path = tmp_path / "detect.txt"
    output_path = tmp_path / "detect.pdf"
    source_path.write_text(
        "研究表明，该路径在多个学院中具有较强复制性，并通过统一模板组织论证结构，因此可以快速迁移到不同课程场景。",
        encoding="utf-8",
    )

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.AIGC_DETECT, "cnki", source_path, output_path, task_id=4)

    assert output_path.exists()
    assert result.result_json["label"] in {"low", "medium", "high", "clean"}
    assert result.result_json["score_pct"] > 0
    assert result.result_json["llm_used"] is False
    assert "algo_package_used" not in result.result_json
    assert result.result_json["aigc_detect_strategy"]["strategy"] == "algorithm"
    assert result.result_json["aigc_detect_strategy"]["platform"] == "cnki"
    assert "algo_package_score" not in result.result_json["score_breakdown"]
    assert "llm_score" not in result.result_json["score_breakdown"]
    assert "distribution" in result.result_json
    assert "fragment_distribution" in result.result_json
    assert "document_metrics" in result.result_json
    assert "decision_basis" in result.result_json


def test_aigc_detect_full_text_pdf_report_is_parseable(tmp_path: Path, db_session: Session, monkeypatch) -> None:
    source_path = tmp_path / "full_text.txt"
    output_path = tmp_path / "full_text_report.pdf"
    paragraphs = [
        f"第{i}段：本研究围绕教学治理展开分析，研究表明该路径具备较强复制性，因此可以看出其表达具有模板化倾向，同时在多个场景中重复使用统一结论。"
        for i in range(1, 19)
    ]
    source_path.write_text("\n".join(paragraphs), encoding="utf-8")

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.AIGC_DETECT, "vip", source_path, output_path, task_id=8)

    assert output_path.exists()
    assert len(result.result_json["paragraph_details"]) == 18
    assert result.result_json["fragment_distribution"]["fragment_count"] > 0
    assert result.result_json["fragment_distribution"]["severe_fragment_count"] >= 0
    assert result.result_json["document_metrics"]["paragraph_count"] == 18
    reader = PdfReader(str(output_path))
    assert len(reader.pages) >= 2


def test_aigc_detect_pdf_uses_single_full_report_layout(tmp_path: Path, db_session: Session, monkeypatch) -> None:
    source_path = tmp_path / "layout_text.txt"
    output_path = tmp_path / "layout_report.pdf"
    paragraphs = [
        "Academic writing quality review sample title",
        "Abstract",
    ] + [
        "This report layout validation paragraph keeps a stable academic tone and repeats enough structured "
        "phrasing to trigger the full text report rendering path."
        for _ in range(10)
    ]
    source_path.write_text("\n".join(paragraphs), encoding="utf-8")

    engine = ProcessingEngine(db_session)
    engine.process(TaskType.AIGC_DETECT, "cnki", source_path, output_path, task_id=9)

    reader = PdfReader(str(output_path))
    extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "AIGC检测" in extracted
    assert "全文报告单" in extracted
    assert "全文检测结果" in extracted
    assert "AIGC片段分布图" in extracted
    assert "分段检测结果" in extracted
    assert "AI特征字符数 / 章节(部分)字符数" in extracted
    assert "原文内容" in extracted
    assert "说明" in extracted
    assert len(reader.pages) >= 2


def test_processing_engine_wrap_pdf_line_accepts_legacy_width_kwarg(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)

    lines = engine._wrap_pdf_line("这是一个用于兼容 width 参数的 PDF 换行测试。", width=10)

    assert lines
    assert all(isinstance(item, str) for item in lines)


def test_transform_text_combined_mode_runs_llm_then_heuristic_fallback(db_session: Session, monkeypatch) -> None:
    engine = ProcessingEngine(db_session)
    engine._effective_mode = "LLM_PLUS_ALGO"

    call_order: list[tuple[str, str]] = []

    def _fake_llm(self, _task_type: TaskType, text: str):
        call_order.append(("llm", text))
        self._pipeline_usage["llm_used"] = True
        return f"llm::{text}"

    monkeypatch.setattr(ProcessingEngine, "_run_llm", _fake_llm)

    output = engine._transform_text("source text", TaskType.DEDUP, "wanfang", {})

    assert output == "llm::source text"
    assert call_order == [("llm", "source text")]


def test_rewrite_cnki_uses_configured_llm_strategy(tmp_path: Path, db_session: Session, monkeypatch) -> None:
    source_path = tmp_path / "rewrite_cnki.txt"
    output_path = tmp_path / "rewrite_cnki_out.txt"
    source_path.write_text("核心素养导向下的教学设计需要兼顾活动组织与过程反馈。", encoding="utf-8")

    db_session.add(
        SystemConfig(
            category="system",
            config_key="rewrite_strategy",
            config_value={
                "cnki": {"rewrite": {"enabled": True, "active_strategy": "llm"}},
                "vip": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.commit()

    def _fake_generate(*_args, **kwargs):
        prompt = str(kwargs.get("text") or "")
        if '"ambient_formality"' in prompt:
            return (
                '{"ambient_formality": "高度书面", "p1_candidates": ["核心素养"], "p2_candidates": ["过程反馈"], '
                '"p3_candidates": ["并"], "p4_candidates": ["教学设计"], "frozen_terms": ["核心素养"], '
                '"high_freq_targets": ["设计"], "avg_n_per_sentence": 3, "needs_sentence_ops": true}'
            )
        if '"semantic_ok"' in prompt:
            return (
                '{"semantic_ok": true, "grammar_ok": true, "style_ok": true, '
                '"compound_ok": true, "density": "14%", "density_ok": true, '
                '"operation_counts": {"A功能词": 1, "B整词同义": 3, "C连接词": 1, "D句法框架": 1, "E副词": 1, "F元话语": 0}, '
                '"issues": [], "verdict": "pass"}'
            )
        return "核心素养导向下的教学设计需要兼顾活动组织、过程反馈与实施节奏，以保持原有论证脉络。"

    monkeypatch.setattr(
        "app.services.rewrite_strategies.cnki_llm.generate_with_llm",
        _fake_generate,
    )

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.REWRITE, "cnki", source_path, output_path, task_id=11)

    assert output_path.exists()
    assert result.result_json["rewrite_strategy"]["strategy"] == "llm"
    assert result.result_json["rewrite_strategy"]["platform"] == "cnki"
    assert result.result_json["llm_used"] is True
    assert "algo_package_used" not in result.result_json


def test_rewrite_vip_forces_llm_strategy_and_uses_wp2_runtime(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "rewrite_vip.txt"
    output_path = tmp_path / "rewrite_vip_out.txt"
    source_path.write_text(
        "课程评价机制的优化需要兼顾反馈时效与教学组织的可执行性，并保持实施反馈的稳定衔接，以支撑课堂治理质量提升。",
        encoding="utf-8",
    )

    def _fake_generate(*_args, **_kwargs):
        text = str(_kwargs.get("text") or "")
        if "输出JSON" in text and "对比原文与改写文" in text:
            return (
                '{"semantic_ok": true, "expansion_ratio": "+26%", "expansion_ok": true, '
                '"additive_style": true, "readability_ok": true, '
                '"mechanism_distribution": {"M1动词叠加": 2, "M2名词叠加": 2, "M3功能词并置": 1, '
                '"M4进行化": 1, "M5字符融合": 0}, "issues": [], "verdict": "pass"}'
            )
        return (
            "课程评价机制方式的优化需要兼顾反馈时效与教学组织的可执行性，并保持实施反馈的稳定衔接过程环节，"
            "以支撑课堂治理质量提升并形成更清晰的执行步骤。"
        )

    monkeypatch.setattr("app.services.rewrite_strategies.vip_llm.generate_with_llm", _fake_generate)
    monkeypatch.setattr("app.services.vip_wp2_runtime.generate_with_llm", _fake_generate)

    db_session.add(
        SystemConfig(
            category="system",
            config_key="rewrite_strategy",
            config_value={
                "cnki": {"rewrite": {"enabled": True, "active_strategy": "llm"}},
                "vip": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.commit()

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.REWRITE, "vip", source_path, output_path, task_id=12)

    assert output_path.exists()
    strategy = result.result_json["rewrite_strategy"]
    assert strategy["strategy"] == "llm"
    assert strategy["platform"] == "vip"
    assert result.result_json["llm_used"] is True
    assert "algo_package_used" not in result.result_json
    assert "vip_wp2_rewrite_llm_ab_strict" == strategy["rule_trace"]["strategy_version"]
    assert strategy["rule_trace"]["prompt_b_validation"]["verdict"] == "pass"
    assert strategy["quality_flags"]["strict_wp2_prompt_b_passed"] is True
    assert output_path.read_text(encoding="utf-8")


def test_rewrite_vip_wp2_prompt_and_rule_trace(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "rewrite_vip_assets.txt"
    output_path = tmp_path / "rewrite_vip_assets_out.txt"
    source_path.write_text(
        "将图书管理与阅读活动结合，能够提升校园阅读生态的整体质量，并形成更加稳定的实施机制，以支撑书香校园建设。",
        encoding="utf-8",
    )

    def _fake_generate(*_args, **_kwargs):
        text = str(_kwargs.get("text") or "")
        if "输出JSON" in text and "对比原文与改写文" in text:
            return (
                '{"semantic_ok": true, "expansion_ratio": "+24%", "expansion_ok": true, '
                '"additive_style": true, "readability_ok": true, '
                '"mechanism_distribution": {"M1动词叠加": 2, "M2名词叠加": 2, "M3功能词并置": 2, '
                '"M4进行化": 1, "M5字符融合": 1}, "issues": [], "verdict": "pass"}'
            )
        return (
            "将把图书管理与阅读活动结合，能够可以提升校园阅读生态的整体质量，并形成更加稳定的实施机制方式，"
            "以支撑书香校园建设过程环节。"
        )

    monkeypatch.setattr("app.services.rewrite_strategies.vip_llm.generate_with_llm", _fake_generate)
    monkeypatch.setattr("app.services.vip_wp2_runtime.generate_with_llm", _fake_generate)

    db_session.add(
        SystemConfig(
            category="system",
            config_key="rewrite_strategy",
            config_value={
                "cnki": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
                "vip": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.commit()

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.REWRITE, "vip", source_path, output_path, task_id=15)

    content = output_path.read_text(encoding="utf-8")
    trace = result.result_json["rewrite_strategy"]["rule_trace"]
    assert content
    assert content != "将图书管理与阅读活动结合，能够提升校园阅读生态的整体质量，并形成更加稳定的实施机制。"
    assert "将把" in content
    assert "能够可以" in content
    assert trace["strategy_version"] == "vip_wp2_rewrite_llm_ab_strict"
    assert trace["mode"] == "llm_prompt_ab_strict_wp2_global"
    assert trace["prompt_b_validation"]["mechanism_distribution"]["M3功能词并置"] >= 1


def test_rewrite_cnki_freezes_algorithm_config_to_strict_v16_llm(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "rewrite_cnki_structural.txt"
    output_path = tmp_path / "rewrite_cnki_structural_out.txt"
    source_path.write_text(
        "为了提升课程治理质量，需要优化实施节奏，并保持评价反馈的稳定衔接。",
        encoding="utf-8",
    )

    db_session.add(
        SystemConfig(
            category="system",
            config_key="rewrite_strategy",
            config_value={
                "cnki": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
                "vip": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.add(
        SystemConfig(
            category="system",
            config_key="llm",
            config_value={"enabled": True, "provider": "local_mock", "model": "local-mock-v1", "api_key": "", "base_url": ""},
            updated_by=1,
        )
    )
    db_session.commit()

    def _fake_generate(*_args, **kwargs):
        prompt = str(kwargs.get("text") or "")
        if '"semantic_ok"' in prompt:
            return (
                '{"semantic_ok": true, "grammar_ok": true, "style_ok": true, '
                '"compound_ok": true, "density": "9%", "density_ok": true, '
                '"operation_counts": {"A功能词": 1, "B整词同义": 2, "C连接词": 1, "D句法框架": 1, "E副词": 0, "F元话语": 0}, '
                '"issues": [], "verdict": "pass"}'
            )
        return "为了提升课程治理质量，需要同步优化实施节奏，并保持评价反馈的稳定衔接。"

    monkeypatch.setattr("app.services.rewrite_strategies.cnki_llm.generate_with_llm", _fake_generate)

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.REWRITE, "cnki", source_path, output_path, task_id=109)

    content = output_path.read_text(encoding="utf-8")
    trace = result.result_json["rewrite_strategy"]["rule_trace"]
    assert "同步优化实施节奏" in content
    assert result.result_json["rewrite_strategy"]["strategy"] == "llm"
    assert trace["mode"] == "llm_prompt_ab_strict_v16_global"
    assert not any(item.startswith("structural:") for item in trace["applied_rules"])


def test_dedup_vip_wp2_runtime_uses_additive_strategy(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "dedup_vip.txt"
    output_path = tmp_path / "dedup_vip_out.txt"
    source_path.write_text(
        "将图书管理与阅读活动结合，能够提升校园阅读生态的整体质量，并形成更加稳定的实施机制，以增强书香校园建设成效。",
        encoding="utf-8",
    )

    def _fake_generate(*_args, **_kwargs):
        text = str(_kwargs.get("text") or "")
        if "输出JSON" in text and "对比原文与改写文" in text:
            return (
                '{"semantic_ok": true, "expansion_ratio": "+23%", "expansion_ok": true, '
                '"additive_style": true, "readability_ok": true, '
                '"mechanism_distribution": {"M1动词叠加": 2, "M2名词叠加": 2, "M3功能词并置": 1, '
                '"M4进行化": 1, "M5字符融合": 1}, "issues": [], "verdict": "pass"}'
            )
        return (
            "将把图书管理与阅读活动结合，能够可以提升校园阅读生态的整体质量，并形成更加稳定的实施机制方式，"
            "以增强书香校园建设成效过程环节。"
        )

    monkeypatch.setattr("app.services.dedup_strategies.vip_llm.generate_with_llm", _fake_generate)
    monkeypatch.setattr("app.services.vip_wp2_runtime.generate_with_llm", _fake_generate)

    db_session.add(
        SystemConfig(
            category="system",
            config_key="dedup_strategy",
            config_value={
                "cnki": {"dedup": {"enabled": True, "active_strategy": "algorithm"}},
                "vip": {"dedup": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.commit()

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.DEDUP, "vip", source_path, output_path, task_id=32)

    assert output_path.exists()
    content = output_path.read_text(encoding="utf-8")
    strategy = result.result_json["dedup_strategy"]
    assert strategy["strategy"] == "llm"
    assert strategy["platform"] == "vip"
    assert content
    assert content != "将图书管理与阅读活动结合，能够提升校园阅读生态的整体质量，并形成更加稳定的实施机制。"
    assert "将把" in content
    assert "能够可以" in content
    assert strategy["rule_trace"]["strategy_version"] == "vip_wp2_dedup_llm_ab_strict"
    assert strategy["rule_trace"]["prompt_b_validation"]["verdict"] == "pass"
    assert strategy["quality_flags"]["strict_wp2_prompt_b_passed"] is True


def test_dedup_vip_prompt_uses_wp2_template() -> None:
    prompt = vip_dedup_prompt("示例文本")

    assert "维普风格中文文本改写执行器" in prompt
    assert "不删原词，在原词旁边叠加语义近邻" in prompt
    assert "M5｜字符融合" in prompt


def test_dedup_cnki_llm_strategy_uses_platform_prompt_and_keeps_dedup_meta(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "dedup_cnki_llm.txt"
    output_path = tmp_path / "dedup_cnki_llm_out.txt"
    source_path.write_text(
        "核心素养导向下的教学设计需要兼顾活动组织与过程反馈，因此需要持续优化课堂实施路径。",
        encoding="utf-8",
    )

    def _fake_generate(*_args, **_kwargs):
        text = str(_kwargs.get("text") or "")
        if "输出JSON" in text:
            if "对比原文与改写文" in text:
                return (
                    '{"semantic_ok": true, "grammar_ok": true, "style_ok": true, '
                    '"compound_ok": true, "density": "8%", "density_ok": true, '
                    '"operation_counts": {"A功能词": 1, "B整词同义": 2, "C连接词": 1, "D句法框架": 1, "E副词": 1, "F元话语": 0}, '
                    '"issues": [], "verdict": "pass"}'
                )
        return "核心素养导向下的教学设计需要兼顾活动组织、过程反馈与课堂推进节奏，由此可进一步优化实施路径。"

    monkeypatch.setattr("app.services.dedup_strategies.cnki_llm.generate_with_llm", _fake_generate)
    db_session.add(
        SystemConfig(
            category="system",
            config_key="dedup_strategy",
            config_value={
                "cnki": {"dedup": {"enabled": True, "active_strategy": "llm"}},
                "vip": {"dedup": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.add(
        SystemConfig(
            category="system",
            config_key="llm",
            config_value={"enabled": True, "provider": "local_mock", "model": "local-mock-v1", "api_key": "", "base_url": ""},
            updated_by=1,
        )
    )
    db_session.commit()

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.DEDUP, "cnki", source_path, output_path, task_id=33)

    assert output_path.exists()
    strategy = result.result_json["dedup_strategy"]
    assert strategy["strategy"] == "llm"
    assert strategy["platform"] == "cnki"
    assert result.result_json["llm_used"] is True
    assert "algo_package_used" not in result.result_json
    assert strategy["quality_flags"]["strict_v16_prompt_b_passed"] is True
    assert strategy["rule_trace"]["mode"] == "dedup_llm_prompt_ab_strict_v16_global"
    assert strategy["rule_trace"]["strategy_version"] == "cnki_v16_dedup_llm_ab_strict"


def test_dedup_validation_rejects_bad_sample_artifacts() -> None:
    source_text = "小学数学可视化教学需要保持术语准确与表达自然，并兼顾论证连贯。" * 3
    rewritten_text = "小学数学作为属于核心课程，相关研究蕴含包括着丰富经验，并形成路径方式。"

    with pytest.raises(Exception) as exc_info:
        validate_dedup_output(
            platform="vip",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "维普WP2扩写量超出允许范围" in str(exc_info.value)


def test_dedup_cnki_validation_rejects_known_bad_sample_patterns() -> None:
    source_text = "派驻监督研究需要兼顾制度逻辑、组织运行和履职约束，并保持论证表达自然。" * 3
    rewritten_text = "由此可见，这说明其属于制度推进路径，同时，进一步看，在此基础上提出结论。"

    with pytest.raises(Exception) as exc_info:
        validate_dedup_output(
            platform="cnki",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "明显异常表达" in str(exc_info.value)


def test_dedup_cnki_validation_rejects_mechanical_connective_cascade() -> None:
    source_text = "派驻监督研究需要兼顾制度逻辑、组织运行和实践约束，并保持论证表达自然。" * 3
    rewritten_text = (
        "派驻监督研究需要兼顾制度逻辑与组织运行。"
        "同时，梳理履职现状。"
        "此外，归纳监督主体结构。"
        "进一步看，分析制度运行原因。"
        "在此基础上，提出优化建议。"
    )

    with pytest.raises(Exception) as exc_info:
        validate_dedup_output(
            platform="cnki",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "机械连接词堆叠" in str(exc_info.value)


def test_rewrite_validation_accepts_borderline_ratio_with_warning() -> None:
    source_text = "研究表明该方案在多个教学场景中具有较强的可执行性与复制价值。" * 3
    rewritten_text = source_text + "同时，研究过程保持既有论证结构。"

    result = validate_rewrite_output(
        platform="cnki",
        source_text=source_text,
        rewritten_text=rewritten_text,
        rule_trace={"mode": "test"},
        strict_length=False,
    )

    assert result.text == rewritten_text
    assert result.quality_flags["length_ok"] is False
    assert any("3%~10%" in item for item in result.warnings)


def test_rewrite_validation_warns_on_soft_mechanical_prefix_usage() -> None:
    source_text = "派驻监督研究需要兼顾制度逻辑、组织运行和实践约束，并保持论证表达自然。" * 3
    rewritten_text = (
        "派驻监督研究需要兼顾制度逻辑、组织运行和实践约束。"
        "同时，围绕履职现状展开分析。"
        "此外，对制度运行影响因素进行归纳。"
    )

    result = validate_rewrite_output(
        platform="cnki",
        source_text=source_text,
        rewritten_text=rewritten_text,
        rule_trace={"mode": "test"},
        strict_length=False,
    )

    assert result.quality_flags["structure_natural_ok"] is False
    assert any("衔接表达偏机械" in item for item in result.warnings)


def test_rewrite_validation_marks_shallow_rewrite_risk() -> None:
    source_text = (
        "数字化治理研究需要同时兼顾制度逻辑、组织协同、流程重构与技术支撑，并在此基础上保持论证路径完整。"
        * 3
    )
    rewritten_text = (
        "数字化治理研究需要同时兼顾制度逻辑、组织协同、流程优化与技术支撑，并在此基础上保持论证路径完整。"
        * 3
    )

    result = validate_rewrite_output(
        platform="cnki",
        source_text=source_text,
        rewritten_text=rewritten_text,
        rule_trace={"mode": "test"},
        strict_length=False,
    )

    assert result.quality_flags["shallow_rewrite_ok"] is False
    assert any("改写幅度偏浅" in item for item in result.warnings)
    assert result.quality_score < 0.9


def test_rewrite_validation_marks_style_alignment_gap_for_fragmented_vip_text() -> None:
    source_text = (
        "书香校园建设需要整合家庭、社区和学校资源，形成稳定的阅读支持体系，同时通过活动设计和平台联动提升学生参与度。"
        * 2
    )
    rewritten_text = (
        "书香校园建设需要整合资源。学校要组织活动。家长要参与阅读。社区要提供支持。平台也要跟进。"
    )

    result = validate_rewrite_output(
        platform="vip",
        source_text=source_text,
        rewritten_text=rewritten_text,
        rule_trace={"mode": "test"},
        strict_length=False,
    )

    assert result.quality_flags["style_profile_available"] is True
    assert result.quality_flags["style_alignment_ok"] is False
    assert any("高质量改写样本" in item for item in result.warnings)


def test_dedup_validation_marks_repetitive_sentence_pattern_risk() -> None:
    source_text = (
        "企业治理研究需要把采购协同、库存效率、资金节奏与监督反馈统筹起来，以保持论证表达稳定且具备专业可信度。"
        * 2
    )
    rewritten_text = (
        "企业要统筹采购协同安排。企业要统筹库存协同安排。企业要统筹资金协同安排。企业要统筹监督协同安排。"
        "企业要统筹采购执行节奏。企业要统筹库存执行节奏。企业要统筹资金执行节奏。企业要统筹监督执行节奏。"
        "企业要统筹采购推进计划。企业要统筹库存推进计划。"
    )

    result = validate_dedup_output(
        platform="vip",
        source_text=source_text,
        rewritten_text=rewritten_text,
        rule_trace={"mode": "test"},
    )

    assert result.quality_flags["discourse_diversity_ok"] is False
    assert any("句式重复度偏高" in item for item in result.warnings)
    assert result.quality_score < 0.9


def test_dedup_validation_marks_style_alignment_gap_for_fragmented_vip_text() -> None:
    source_text = (
        "企业存货管理优化需要统筹采购节奏、库存结构、周转效率和监督机制，同时保持论证表达紧凑并与经营情境对应。"
        * 2
    )
    rewritten_text = (
        "企业要管理存货。采购要安排。库存要控制。周转要提升。监督也要跟上。"
    )

    with pytest.raises(Exception) as exc_info:
        validate_dedup_output(
            platform="vip",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "维普WP2扩写量超出允许范围" in str(exc_info.value)


def test_dedup_vip_wp2_validation_rejects_non_additive_result(db_session: Session, monkeypatch) -> None:
    source_text = (
        "企业存货管理优化需要统筹采购节奏、库存结构、周转效率和监督机制，同时保持论证表达紧凑并与经营情境对应。"
        * 2
    )

    def _fake_generate(*_args, **_kwargs):
        text = str(_kwargs.get("text") or "")
        if "输出JSON" in text and "对比原文与改写文" in text:
            return (
                '{"semantic_ok": true, "expansion_ratio": "+6%", "expansion_ok": false, '
                '"additive_style": false, "readability_ok": true, '
                '"mechanism_distribution": {"M1动词叠加": 0, "M2名词叠加": 0, "M3功能词并置": 0, '
                '"M4进行化": 0, "M5字符融合": 0}, "issues": [{"location":"第1句","type":"叠加方式","detail":"纯替换"}], '
                '"verdict": "fail"}'
            )
        return "企业通过优化采购、库存和周转管理来提升监督效果。"

    monkeypatch.setattr("app.services.dedup_strategies.vip_llm.generate_with_llm", _fake_generate)
    monkeypatch.setattr("app.services.vip_wp2_runtime.generate_with_llm", _fake_generate)

    with pytest.raises(Exception) as exc_info:
        execute_dedup_strategy(
            db_session,
            task=None,
            platform="vip",
            text=source_text,
            report_summary={},
            strategy="llm",
        )

    assert "A/B 校验未通过" in str(exc_info.value)


def test_rewrite_validation_still_rejects_extreme_length_ratio() -> None:
    source_text = "研究表明该方案在多个教学场景中具有较强的可执行性与复制价值。" * 3
    rewritten_text = source_text[:20]

    with pytest.raises(Exception) as exc_info:
        validate_rewrite_output(
            platform="cnki",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "超出可处理范围" in str(exc_info.value)


def test_rewrite_validation_rejects_bad_sample_artifacts() -> None:
    source_text = "小学数学可视化教学需要保持术语准确与表达自然。" * 3
    rewritten_text = "小学数学作为属于核心课程，相关研究蕴含包括着丰富经验，并形成路径方式。"

    with pytest.raises(Exception) as exc_info:
        validate_rewrite_output(
            platform="vip",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "明显异常表达" in str(exc_info.value)


def test_rewrite_cnki_validation_rejects_known_bad_sample_patterns() -> None:
    source_text = "可视化教学模型需要兼顾图像呈现、动态演示与智能理解三个层次。" * 3
    rewritten_text = "由此可见，这说明其属于图像表现出层，并进一步保持原有论证脉络。"

    with pytest.raises(Exception) as exc_info:
        validate_rewrite_output(
            platform="cnki",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "明显异常表达" in str(exc_info.value)


def test_rewrite_validation_rejects_cross_domain_term_weakening() -> None:
    source_text = "边缘计算架构需要兼顾实时推理、带宽约束和设备异构条件下的任务调度。" * 3
    rewritten_text = "边缘进行计算的架构需要兼顾实时性的推理、带宽方面的约束以及设备异构条件下面的任务调度过程。"

    with pytest.raises(Exception) as exc_info:
        validate_rewrite_output(
            platform="cnki",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "明显异常表达" in str(exc_info.value)


def test_rewrite_cnki_validation_rejects_mechanical_connective_cascade() -> None:
    source_text = "派驻监督研究需要兼顾制度逻辑、组织运行和实践约束，并保持论证表达自然。" * 3
    rewritten_text = (
        "摘要：派驻监督研究需要兼顾制度逻辑、组织运行和实践约束。"
        "同时，围绕履职现状展开分析。"
        "此外，归纳监督主体的结构特点。"
        "进一步看，梳理制度运行中的深层原因。"
        "在此基础上，提出优化建议。"
    )

    with pytest.raises(Exception) as exc_info:
        validate_rewrite_output(
            platform="cnki",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "机械连接词堆叠" in str(exc_info.value)


def test_rewrite_vip_wp2_allows_additive_expression_without_bad_pattern_rejection(
    tmp_path: Path, db_session: Session, monkeypatch
) -> None:
    source_path = tmp_path / "rewrite_vip_no_soft_expand.txt"
    output_path = tmp_path / "rewrite_vip_no_soft_expand_out.txt"
    source_path.write_text(
        (
            "摘要：小学图书馆作为校园文化的重要载体，其功能定位经历着深刻变革。传统图书管理以藏书保管为核心，"
            "重视馆藏数量与完整性，却疏于发挥图书的教育价值。随着基础教育改革深入推进，图书馆管理理念亟需从静态保管"
            "转向动态服务，从被动等待转向主动引导。"
        ),
        encoding="utf-8",
    )

    def _fake_generate(*_args, **_kwargs):
        text = str(_kwargs.get("text") or "")
        if "输出JSON" in text and "对比原文与改写文" in text:
            return (
                '{"semantic_ok": true, "expansion_ratio": "+22%", "expansion_ok": true, '
                '"additive_style": true, "readability_ok": true, '
                '"mechanism_distribution": {"M1动词叠加": 3, "M2名词叠加": 2, "M3功能词并置": 2, '
                '"M4进行化": 1, "M5字符融合": 0}, "issues": [], "verdict": "pass"}'
            )
        return (
            "摘要：小学图书馆作为校园文化的重要载体，其功能定位经历着深刻变革过程环节。传统图书管理以藏书保管为核心，"
            "重视馆藏数量与完整性，却疏于发挥图书的教育价值意义。随着基础教育改革深入推进，图书馆管理理念亟需要从静态保管"
            "转向动态服务，从被动等待转向主动引导，并形成更加稳定的实施机制方式。"
        )

    monkeypatch.setattr("app.services.rewrite_strategies.vip_llm.generate_with_llm", _fake_generate)
    monkeypatch.setattr("app.services.vip_wp2_runtime.generate_with_llm", _fake_generate)

    db_session.add(
        SystemConfig(
            category="system",
            config_key="rewrite_strategy",
            config_value={
                "cnki": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
                "vip": {"rewrite": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.commit()

    engine = ProcessingEngine(db_session)
    result = engine.process(TaskType.REWRITE, "vip", source_path, output_path, task_id=107)

    content = output_path.read_text(encoding="utf-8")
    strategy = result.result_json["rewrite_strategy"]

    assert strategy["strategy"] == "llm"
    assert "亟需要" in content or "机制方式" in content
    assert strategy["quality_flags"]["strict_wp2_prompt_b_passed"] is True
    assert strategy["rule_trace"]["prompt_b_validation"]["verdict"] == "pass"


def test_dedup_validation_rejects_cross_domain_term_weakening() -> None:
    source_text = "现金流风险预警是企业财务稳健性分析的重要环节。" * 3
    rewritten_text = "现金流方面的风险情况预警是企业财务稳健性分析中的关键环节和重要形式。"

    with pytest.raises(Exception) as exc_info:
        validate_dedup_output(
            platform="vip",
            source_text=source_text,
            rewritten_text=rewritten_text,
            rule_trace={"mode": "test"},
        )

    assert "维普WP2扩写量超出允许范围" in str(exc_info.value)


def test_dedup_validation_marks_cross_domain_terms_missing_without_hard_fail() -> None:
    source_text = "社区记忆的建构既受地方叙事影响，也与代际传播机制和公共空间实践相关。" * 2
    rewritten_text = (
        "相关记忆的建构主要受地方叙事持续影响，也与代际传播表达、空间实践安排和社区互动场景相关联，"
        "并在后续传播过程中不断被重新强化与重复提及，同时在不同公共讨论空间里形成持续扩散的叙述回路。"
    )

    result = validate_dedup_output(
        platform="vip",
        source_text=source_text,
        rewritten_text=rewritten_text,
        rule_trace={"mode": "test"},
    )

    assert result.quality_flags["cross_domain_terms_ok"] is False
    assert result.quality_flags["protected_content_ok"] is False
    assert any("保护术语缺失" in item for item in result.warnings)


def test_llm_prompts_include_cross_discipline_constraints() -> None:
    source = "测试文本"
    cnki_prompt = cnki_rewrite_prompt(source)
    cnki_dedup = cnki_dedup_prompt(source)
    vip_rewrite = vip_rewrite_prompt(source)
    vip_dedup = vip_dedup_prompt(source)

    assert "你是中文文本改写执行器。强制完整执行以下全部阶段，直接输出结果，不加任何说明。" in cnki_prompt
    assert "改写密度硬性要求" in cnki_prompt
    assert "阶段一：冻结保护" in cnki_prompt
    assert "P5｜正式度下调" in cnki_prompt
    assert "C4｜复合词完整性终检" in cnki_prompt
    assert "输入文本：测试文本" in cnki_prompt
    assert "输入文本：测试文本" in vip_rewrite
    assert "维普风格中文文本改写执行器" in vip_rewrite
    assert "不删原词，在原词旁边叠加语义近邻" in vip_rewrite
    assert "M5｜字符融合" in vip_rewrite
    assert "输入文本：测试文本" in cnki_dedup
    assert "阶段一：冻结保护" in cnki_dedup
    assert "P2｜整词同义替换" in cnki_dedup
    assert "C5｜全文改写密度核算" in cnki_dedup
    assert "输入文本：测试文本" in vip_dedup
    assert "维普风格中文文本改写执行器" in vip_dedup
    assert "每段字数比原文增加 20%–30%" in vip_dedup
    assert "M4｜进行化插入" in vip_dedup


def test_transform_docx_skips_blank_runs_for_dedup(tmp_path: Path, db_session: Session) -> None:
    source_path = tmp_path / "dedup_blank_runs.docx"
    output_path = tmp_path / "dedup_blank_runs_out.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("   ")
    paragraph.add_run("因此需要进一步优化课堂组织路径。")
    paragraph.add_run("\n")
    doc.save(source_path)

    db_session.add(
        SystemConfig(
            category="system",
            config_key="dedup_strategy",
            config_value={
                "cnki": {"dedup": {"enabled": True, "active_strategy": "algorithm"}},
                "vip": {"dedup": {"enabled": True, "active_strategy": "algorithm"}},
            },
            updated_by=1,
        )
    )
    db_session.commit()

    def _fake_generate(_db, *, task_type, text: str):
        if "输出JSON" in str(text or ""):
            return (
                '{"semantic_ok": true, "grammar_ok": true, "style_ok": true, '
                '"compound_ok": true, "density": "10%", "density_ok": true, '
                '"issues": [], "verdict": "pass"}'
            )
        return "因此需要进一步优化课堂组织路径，并同步改进实施节奏。"

    with pytest.MonkeyPatch.context() as monkeypatch:
        monkeypatch.setattr("app.services.dedup_strategies.cnki_llm.generate_with_llm", _fake_generate)

        engine = ProcessingEngine(db_session)
        result = engine.process(TaskType.DEDUP, "cnki", source_path, output_path, task_id=36)

        output_doc = Document(str(output_path))
        output_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
        assert output_path.exists()
        assert result.result_json["type"] == "dedup"
        assert "改进" in output_text or "优化" in output_text or "路径" in output_text


def test_transform_text_combined_mode_falls_back_to_heuristic_when_llm_empty(
    db_session: Session, monkeypatch
) -> None:
    engine = ProcessingEngine(db_session)
    engine._effective_mode = "LLM_PLUS_ALGO"

    call_order: list[str] = []

    def _fake_llm(self, _task_type: TaskType, _text: str):
        call_order.append("llm")
        return None

    monkeypatch.setattr(ProcessingEngine, "_run_llm", _fake_llm)

    output = engine._transform_text("因此需要优化", TaskType.DEDUP, "wanfang", {})

    assert output
    assert output != "algo::因此需要优化"
    assert call_order == ["llm"]


def test_split_detect_paragraphs_merges_wrapped_lines(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)
    text = "\n".join(
        [
            "第一章 绪论",
            "本研究围绕数智教学改革展开分析，系统梳理平台治理、课程设计与组织协同三方面问题，",
            "研究表明该路径具备较强复制性，并且在多个应用场景中形成稳定的实施框架，",
            "同时能够通过统一表达快速生成结构完整的段落内容。",
            "此外，本研究进一步从资源统筹角度提出建议，",
            "强调通过过程监督与结果复核提升治理质量。",
        ]
    )

    paragraphs = engine._split_detect_paragraphs(text)

    assert paragraphs[0] == "第一章 绪论"
    assert len(paragraphs) == 3
    assert paragraphs[1].startswith("本研究围绕数智教学改革展开分析")
    assert paragraphs[2].startswith("此外，本研究进一步从资源统筹角度提出建议")


def test_detect_label_supports_clean_band(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)
    profile = engine._platform_detect_profile("cnki")

    assert engine._score_to_detect_label(0.12, profile) == "clean"
    assert engine._score_to_detect_label(0.24, profile) == "low"
    assert engine._score_to_detect_label(0.50, profile) == "medium"


def test_detect_label_prefers_higher_risk_between_algo_and_score(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)

    assert engine._prefer_higher_risk_detect_label("clean", "low") == "low"
    assert engine._prefer_higher_risk_detect_label("low", "medium") == "medium"
    assert engine._prefer_higher_risk_detect_label("high", "low") == "high"


def test_cnki_borderline_paragraph_is_not_kept_clean(db_session: Session, monkeypatch) -> None:
    engine = ProcessingEngine(db_session)
    profile = engine._platform_detect_profile("cnki")

    monkeypatch.setattr(ProcessingEngine, "_local_suspicious_segments_v2", lambda self, _paragraph, _platform, _profile: [])
    monkeypatch.setattr(
        ProcessingEngine,
        "_simulate_platform_detect_score",
        lambda self, _platform, _text, _base: (
            0.155,
            profile,
            {
                "template_signal": 0.0,
                "repeat_signal": 0.0,
                "context_signal": 0.0,
                "opening_signal": 0.0,
                "artifact_signal": 0.0,
                "english_abstract_signal": 0.0,
                "citation_relief": 0.0,
                "evidence_relief": 0.0,
                "style_signal": 0.0,
                "template_hits": [],
            },
        ),
    )

    rows = engine._build_detect_paragraph_details(
        "这是一段足够长的正文内容，用于验证知网正文段落在边界分数附近不会继续被压成 clean 标签，而是至少进入 low 标签。",
        "cnki",
        profile,
    )

    assert rows[0]["label"] == "low"


def test_cnki_short_doc_heading_inherits_following_risk(db_session: Session, monkeypatch) -> None:
    engine = ProcessingEngine(db_session)
    profile = engine._platform_detect_profile("cnki")

    monkeypatch.setattr(ProcessingEngine, "_local_suspicious_segments_v2", lambda self, _paragraph, _platform, _profile: [])

    score_map = {
        "（三）课程融合：在一日生活中埋下种子": 0.11,
        "第一段正文用于模拟连续风险内容，这里保持较长文本以便进入边界判断。": 0.24,
        "第二段正文继续承接上一段，维持明显的疑似风险分数，用于让标题继承低风险标签。": 0.19,
    }

    def _fake_simulate(self, _platform: str, text: str, _base: float):
        return (
            score_map[text],
            profile,
            {
                "template_signal": 0.0,
                "repeat_signal": 0.0,
                "context_signal": 0.0,
                "opening_signal": 0.0,
                "artifact_signal": 0.0,
                "english_abstract_signal": 0.0,
                "citation_relief": 0.0,
                "evidence_relief": 0.0,
                "style_signal": 0.0,
                "template_hits": [],
            },
        )

    monkeypatch.setattr(ProcessingEngine, "_simulate_platform_detect_score", _fake_simulate)

    rows = engine._build_detect_paragraph_details(
        "\n".join(score_map.keys()),
        "cnki",
        profile,
    )

    assert rows[0]["label"] == "low"
    assert rows[0]["suspicious_segments"][0]["label"] == "low"
    assert rows[0]["suspicious_segments"][0]["reason"] == "风险区标题延续"


def test_cnki_risky_segment_uses_internal_low_label_without_external_package(
    db_session: Session, monkeypatch
) -> None:
    engine = ProcessingEngine(db_session)
    profile = engine._platform_detect_profile("cnki")

    monkeypatch.setattr(
        ProcessingEngine,
        "_local_suspicious_segments_v2",
        lambda self, _paragraph, _platform, _profile: [{"text": "可疑句段", "score": 35.0, "reason": "测试"}],
    )
    monkeypatch.setattr(
        ProcessingEngine,
        "_simulate_platform_detect_score",
        lambda self, _platform, _text, _base: (
            0.26,
            profile,
            {
                "template_signal": 0.0,
                "repeat_signal": 0.0,
                "context_signal": 0.0,
                "opening_signal": 0.0,
                "artifact_signal": 0.0,
                "english_abstract_signal": 0.0,
                "citation_relief": 0.0,
                "evidence_relief": 0.0,
                "style_signal": 0.0,
                "template_hits": [],
            },
        ),
    )

    rows = engine._build_detect_paragraph_details(
        "这是一段用于测试片段标签升级的正文内容，长度足够，且应该被识别为可疑。",
        "cnki",
        profile,
    )

    assert rows[0]["suspicious_segments"][0]["label"] == "low"


def test_local_suspicious_segments_v2_skips_weak_sentences(db_session: Session, monkeypatch) -> None:
    engine = ProcessingEngine(db_session)
    profile = engine._platform_detect_profile("cnki")

    def _fake_simulate(self, _platform: str, _text: str, _base_score: float):
        return (
            0.39,
            profile,
            {
                "template_signal": 0.05,
                "repeat_signal": 0.08,
                "context_signal": 0.10,
                "opening_signal": 0.02,
                "artifact_signal": 0.0,
                "english_abstract_signal": 0.0,
                "citation_relief": 0.0,
                "evidence_relief": 0.0,
                "template_hits": [],
            },
        )

    monkeypatch.setattr(ProcessingEngine, "_simulate_platform_detect_score", _fake_simulate)

    segments = engine._local_suspicious_segments_v2("这是一个正常句子。这里继续补充说明。", "cnki", profile)

    assert segments == []


def test_local_suspicious_segments_v2_skips_cnki_summary_wrapup_sentences(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)
    profile = engine._platform_detect_profile("cnki")

    segments = engine._local_suspicious_segments_v2(
        "这一模式的可迁移性体现在两个层面，最核心的教育逻辑在于形成系列化的家园共育体系。",
        "cnki",
        profile,
    )

    assert segments == []


def test_is_no_ai_fragment_filters_thesis_front_matter(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)

    assert engine._is_no_ai_fragment("学 位 论 文 独 创 性 说 明")
    assert engine._is_no_ai_fragment("本人郑重声明：所呈交的学位论文是我个人在导师指导下进行的研究工作。")


def test_human_case_relief_signal_prefers_practice_paragraphs(db_session: Session) -> None:
    engine = ProcessingEngine(db_session)
    text = (
        "在日常观察中发现，幼儿对“爱家人”这件事不是没有感情，而是缺少一种看得见、摸得着的表达渠道。"
        "正是立足这一现实，本园设计并推行“爱的存折”家园共育项目，尝试把孝亲教育转变为可记录、可反馈的生活实践。"
    )

    assert engine._human_case_relief_signal(text) >= 0.12
