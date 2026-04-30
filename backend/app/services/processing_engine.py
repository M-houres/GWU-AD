import json
import math
import re
import statistics
from difflib import SequenceMatcher
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models import LLMErrorLog, SwitchLog, SystemSwitch, TaskType
from app.models import Task
from app.services.aigc_detect_strategies.common import split_paragraphs as split_detect_strategy_paragraphs
from app.services.aigc_detect_strategies.executor import execute_aigc_detect_strategy
from app.services.dedup_strategies.executor import (
    STRATEGY_LLM as DEDUP_STRATEGY_LLM,
    execute_dedup_strategy,
)
from app.services.dedup_strategies.config import get_active_dedup_strategy
from app.services.llm_service import generate_with_llm, load_llm_config
from app.services.processing_detect_result_builder import (
    build_detect_result_payload,
    build_detect_summary,
    enrich_detect_breakdown,
)
from app.services.processing_result_builder import build_transform_result
from app.services.processing_report_summary import extract_percent, extract_report_summary
from app.services.processing_text_tools import split_long_sentences
from app.services.rewrite_strategies.executor import execute_rewrite_strategy
from app.utils import count_billable_chars, extract_text_from_file

MODE_LLM_PLUS_ALGO = "LLM_PLUS_ALGO"
MODE_ALGO_ONLY = "ALGO_ONLY"
settings = get_settings()


@dataclass
class ProcessResult:
    output_path: str
    result_json: dict


class ProcessingEngine:
    def __init__(self, db: Session) -> None:
        self.db = db
        self._current_switch: SystemSwitch | None = None
        self._current_task_id: int | None = None
        self._current_detect_source_text = ""
        self._pipeline_usage = {"llm_used": False}
        self._effective_mode = MODE_ALGO_ONLY
        self._rewrite_strategy_meta: dict | None = None
        self._dedup_strategy_meta: dict | None = None

    def _get_or_init_switch(self) -> SystemSwitch:
        switch = self.db.query(SystemSwitch).first()
        if switch:
            return switch
        switch = SystemSwitch(
            current_mode=MODE_LLM_PLUS_ALGO if settings.llm_enabled_default else MODE_ALGO_ONLY,
            llm_enabled=settings.llm_enabled_default,
            llm_fail_count=0,
            llm_fail_threshold=3,
        )
        self.db.add(switch)
        self.db.flush()
        return switch

    def _switch_mode(self, target_mode: str, reason: str) -> None:
        switch = self._get_or_init_switch()
        if switch.current_mode == target_mode:
            return
        self.db.add(SwitchLog(from_mode=switch.current_mode, to_mode=target_mode, reason=reason))
        switch.current_mode = target_mode
        self.db.flush()

    def _normalize_requested_mode(self, processing_mode: str | None) -> str | None:
        raw = str(processing_mode or "").strip()
        mode = raw.upper()
        if mode in {MODE_ALGO_ONLY, MODE_LLM_PLUS_ALGO}:
            return mode
        lowered = raw.lower()
        if lowered in {"algo_only", "algo"}:
            return MODE_ALGO_ONLY
        if lowered in {"algo_llm", "llm_plus_algo", "algo+llm"}:
            return MODE_LLM_PLUS_ALGO
        return None

    def _resolve_effective_mode(self, switch: SystemSwitch, requested_mode: str | None) -> str:
        if requested_mode == MODE_ALGO_ONLY:
            return MODE_ALGO_ONLY
        if requested_mode == MODE_LLM_PLUS_ALGO:
            if switch.llm_enabled and switch.current_mode == MODE_LLM_PLUS_ALGO:
                return MODE_LLM_PLUS_ALGO
            return MODE_ALGO_ONLY
        return switch.current_mode

    def process(
        self,
        task_type: TaskType,
        platform: str,
        input_path: Path,
        output_path: Path,
        task_id: int | None = None,
        report_path: Path | None = None,
        processing_mode: str | None = None,
    ) -> ProcessResult:
        switch = self._get_or_init_switch()
        self._current_switch = switch
        self._current_task_id = task_id
        self._current_detect_source_text = ""
        self._pipeline_usage = {"llm_used": False}
        self._rewrite_strategy_meta = None
        self._dedup_strategy_meta = None

        llm_cfg = load_llm_config(self.db)
        switch.llm_enabled = bool(llm_cfg.get("enabled", False))
        normalized_platform = (platform or "").strip().lower()
        if not switch.llm_enabled:
            self._switch_mode(MODE_ALGO_ONLY, "llm disabled")
        elif switch.current_mode != MODE_ALGO_ONLY or switch.llm_fail_count < switch.llm_fail_threshold:
            self._switch_mode(MODE_LLM_PLUS_ALGO, "llm healthy")
        requested_mode = self._normalize_requested_mode(processing_mode)
        self._effective_mode = self._resolve_effective_mode(switch, requested_mode)

        source_text = extract_text_from_file(input_path)
        report_text = self._load_optional_report(report_path)
        report_summary = self._extract_report_summary(task_type, report_text)

        if task_type == TaskType.AIGC_DETECT:
            self._current_detect_source_text = source_text
            detect_task = self.db.get(Task, task_id) if task_id else None
            detect_result = execute_aigc_detect_strategy(
                self.db,
                task=detect_task,
                text=source_text,
                platform=normalized_platform,
                report_summary=report_summary,
                mode=self._effective_mode,
            )
            detect_result["report_view"] = self._build_detect_report_view(detect_result)
            self._write_detect_report_pdf(output_path, detect_result)
            return ProcessResult(output_path=str(output_path), result_json=detect_result)

        if input_path.suffix.lower() == ".docx":
            self._transform_docx(input_path, output_path, task_type, normalized_platform, report_summary)
            output_text = extract_text_from_file(output_path)
        else:
            output_text = self._transform_text(source_text, task_type, normalized_platform, report_summary)
            self._write_text_output(output_path, output_text)

        result_json = self._build_transform_result(
            task_type=task_type,
            platform=normalized_platform,
            mode=self._effective_mode,
            source_text=source_text,
            output_text=output_text,
            report_summary=report_summary,
        )
        return ProcessResult(output_path=str(output_path), result_json=result_json)

    def _load_optional_report(self, report_path: Path | None) -> str:
        if report_path is None or not report_path.exists():
            return ""
        try:
            return extract_text_from_file(report_path)
        except Exception:
            return ""

    def _heuristic_ai_score(self, text: str) -> float:
        clean = text.strip()
        if not clean:
            return 0.0
        normalized = clean.replace("。", ".").replace("！", ".").replace("？", ".")
        sentences = [seg.strip() for seg in normalized.split(".") if seg.strip()]
        if not sentences:
            return 0.0
        avg_len = sum(len(s) for s in sentences) / len(sentences)
        uniq_ratio = len(set(clean)) / max(len(clean), 1)
        score = min(1.0, max(0.0, (avg_len / 80.0) * 0.6 + (1 - uniq_ratio) * 0.4))
        return round(score, 4)

    def _iter_body_runs(self, doc: Document):
        for para in doc.paragraphs:
            for run in para.runs:
                yield run
        if settings.docx_process_table_text:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                yield run

    def _iter_body_paragraphs(self, doc: Document):
        for para in doc.paragraphs:
            yield para
        if settings.docx_process_table_text:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para

    def _write_text_output(self, output_path: Path, text: str) -> None:
        if output_path.suffix.lower() != ".docx":
            output_path.write_text(text, encoding="utf-8")
            return

        paragraphs = text.splitlines()
        doc = Document()
        if paragraphs:
            doc.paragraphs[0].text = paragraphs[0]
            for paragraph_text in paragraphs[1:]:
                doc.add_paragraph(paragraph_text)
        else:
            doc.paragraphs[0].text = ""
        doc.save(str(output_path))

    def _is_docx_reference_heading(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", str(text or ""))
        return normalized in {"参考文献", "参考文献:", "参考文献：", "references", "REFERENCES"}

    def _is_docx_abstract_heading(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", str(text or ""))
        lowered = normalized.lower()
        return lowered in {"摘要", "摘要:", "摘要：", "中文摘要", "英文摘要", "abstract", "abstract:", "abstract："}

    def _is_docx_keyword_heading(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", str(text or ""))
        lowered = normalized.lower()
        return lowered in {
            "关键词",
            "关键词:",
            "关键词：",
            "关键字",
            "关键字:",
            "关键字：",
            "【关键词】",
            "【关键字】",
            "keywords",
            "keywords:",
            "keywords：",
        }

    def _is_docx_intro_heading(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", str(text or ""))
        lowered = normalized.lower()
        return lowered in {
            "引言",
            "引言:",
            "引言：",
            "绪论",
            "绪论:",
            "绪论：",
            "前言",
            "前言:",
            "前言：",
            "导论",
            "导论:",
            "导论：",
        }

    def _is_docx_title_like(self, text: str, *, nonempty_index: int) -> bool:
        content = str(text or "").strip()
        if not content:
            return False
        return nonempty_index == 1 and len(content) <= 64 and not re.search(r"[。！？；;]", content)

    def _is_docx_subtitle_like(self, text: str, *, nonempty_index: int, previous_kind: str) -> bool:
        content = str(text or "").strip()
        if not content:
            return False
        if nonempty_index > 3:
            return False
        if previous_kind not in {"title", "subtitle"}:
            return False
        if content.startswith(("——", "—", "--", "-", "副标题")):
            return True
        return len(content) <= 48 and not re.search(r"[。！？；;]", content)

    def _is_docx_keyword_body(self, text: str, *, previous_kind: str) -> bool:
        content = str(text or "").strip()
        if previous_kind != "keyword_heading" or not content:
            return False
        if len(content) > 120:
            return False
        return any(token in content for token in ("；", ";", "、", "，"))

    def _match_docx_inline_front_label(self, text: str) -> tuple[str, str, str] | None:
        content = str(text or "").strip()
        if not content:
            return None
        patterns = (
            ("abstract_inline", r"^(【摘要】|摘要[:：])\s*(.+)$"),
            ("keyword_inline", r"^(【关键词】|【关键字】|关键词[:：]|关键字[:：])\s*(.+)$"),
        )
        for kind, pattern in patterns:
            match = re.match(pattern, content, flags=re.IGNORECASE | re.DOTALL)
            if match:
                return kind, match.group(1), match.group(2).strip()
        return None

    def _classify_docx_paragraph_kind(
        self,
        paragraph,
        text: str,
        *,
        nonempty_index: int,
        previous_kind: str,
        in_reference_section: bool,
    ) -> str:
        content = str(text or "").strip()
        if not content:
            return "blank"
        if in_reference_section:
            return "reference_body"
        if self._is_docx_reference_heading(content):
            return "reference_heading"
        if self._docx_paragraph_has_superscript(paragraph):
            return "preserved"
        if self._is_docx_title_like(content, nonempty_index=nonempty_index):
            return "title"
        if self._is_docx_subtitle_like(content, nonempty_index=nonempty_index, previous_kind=previous_kind):
            return "subtitle"
        if self._is_docx_abstract_heading(content):
            return "abstract_heading"
        if self._is_docx_keyword_heading(content):
            return "keyword_heading"
        inline_label = self._match_docx_inline_front_label(content)
        if inline_label:
            return inline_label[0]
        if self._is_docx_keyword_body(content, previous_kind=previous_kind):
            return "keyword_body"
        if self._is_docx_intro_heading(content):
            return "intro_heading"
        if re.match(r"^【[^】]{1,12}】$", content):
            return "label_heading"
        if self._is_docx_heading_or_caption(paragraph, content):
            return "heading"
        return "body"

    def _is_docx_heading_or_caption(self, paragraph, text: str) -> bool:
        content = str(text or "").strip()
        if not content:
            return True
        if content.startswith(("关键词：", "关键词:", "关键字：", "关键字:")):
            return True
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
        if "heading" in style_name or "标题" in style_name:
            return True
        if re.match(r"^(图|表)\s*[0-9一二三四五六七八九十]+", content):
            return True
        if len(content) <= 32 and not re.search(r"[。！？；;]", content):
            if re.match(r"^([一二三四五六七八九十]+[、.．]|第[一二三四五六七八九十百零0-9]+[章节部分篇]|[（(][一二三四五六七八九十百零0-9]+[)）])", content):
                return True
        return False

    def _docx_paragraph_has_superscript(self, paragraph) -> bool:
        for run in getattr(paragraph, "runs", []):
            if not (run.text and run.text.strip()):
                continue
            if getattr(getattr(run, "font", None), "superscript", None) is True:
                return True
            run_xml = str(getattr(getattr(run, "_r", None), "xml", "") or "")
            if "vertAlign" in run_xml and "superscript" in run_xml:
                return True
        return False

    def _split_docx_first_sentence(self, text: str) -> tuple[str, str]:
        content = str(text or "")
        if not content:
            return "", ""
        sentence_end = -1
        for idx, char in enumerate(content):
            if char in "。！？!?":
                sentence_end = idx
                break
            if char in ".．":
                prev_char = content[idx - 1] if idx > 0 else ""
                next_char = content[idx + 1] if idx + 1 < len(content) else ""
                if prev_char.isdigit() and next_char.isdigit():
                    continue
                if idx + 1 < len(content) and next_char not in " \t\r\n\"'”’）)]】》":
                    continue
                sentence_end = idx
                break
        if sentence_end < 0:
            return content, ""
        split_at = sentence_end + 1
        while split_at < len(content) and content[split_at] in " \t\r\n\"'”’）)]】》":
            split_at += 1
        return content[:split_at], content[split_at:]

    def _redistribute_paragraph_text(self, paragraph, rewritten_text: str) -> None:
        text_runs = [run for run in paragraph.runs if run.text and run.text.strip()]
        if not text_runs:
            return
        if len(text_runs) == 1:
            text_runs[0].text = rewritten_text
            return

        source_lengths = [max(len(run.text), 1) for run in text_runs]
        total_length = sum(source_lengths)
        cursor = 0
        assigned_until = 0
        for index, run in enumerate(text_runs):
            if index == len(text_runs) - 1:
                segment = rewritten_text[cursor:]
            else:
                assigned_until += source_lengths[index]
                target_end = round(len(rewritten_text) * assigned_until / total_length)
                target_end = max(cursor, min(len(rewritten_text), target_end))
                segment = rewritten_text[cursor:target_end]
                cursor = target_end
            run.text = segment

    def _replace_docx_inline_label_paragraph(self, paragraph, label_text: str, rewritten_body: str) -> None:
        text_runs = [run for run in paragraph.runs if run.text]
        if not text_runs:
            return
        normalized_label = str(label_text or "")
        remaining = str(rewritten_body or "")
        if len(text_runs) == 1:
            text_runs[0].text = f"{normalized_label}{remaining}"
            return
        assigned_label = False
        label_run = None
        for run in text_runs:
            current = run.text or ""
            if not assigned_label and normalized_label and current.startswith(normalized_label):
                run.text = normalized_label
                assigned_label = True
                label_run = run
                continue
            run.text = ""
        if assigned_label:
            for run in text_runs:
                if run.text == "":
                    run.text = remaining
                    remaining = ""
                    break
            if remaining and label_run is not None:
                label_run.text = f"{label_run.text}{remaining}"
                remaining = ""
        else:
            text_runs[0].text = f"{normalized_label}{remaining}"
            for run in text_runs[1:]:
                run.text = ""

    def _should_transform_docx_paragraph(self, paragraph, text: str, *, in_reference_section: bool) -> bool:
        if not text.strip():
            return False
        if in_reference_section:
            return False
        if self._docx_paragraph_has_superscript(paragraph):
            return False
        if self._is_docx_heading_or_caption(paragraph, text):
            return False
        return True

    def _transform_docx(
        self,
        input_path: Path,
        output_path: Path,
        task_type: TaskType,
        platform: str,
        report_summary: dict | None = None,
    ) -> None:
        doc = Document(str(input_path))
        if platform == "cnki" and task_type in {TaskType.REWRITE, TaskType.DEDUP}:
            self._transform_docx_cnki_pipeline(doc, task_type)
            doc.save(str(output_path))
            return
        if platform == "vip" and task_type in {TaskType.REWRITE, TaskType.DEDUP}:
            self._transform_docx_vip_pipeline(doc, task_type)
            doc.save(str(output_path))
            return
        summary = report_summary or {}
        in_reference_section = False
        nonempty_index = 0
        previous_kind = ""
        for paragraph in self._iter_body_paragraphs(doc):
            source_text = paragraph.text or ""
            stripped = source_text.strip()
            if not stripped:
                continue
            nonempty_index += 1
            kind = self._classify_docx_paragraph_kind(
                paragraph,
                stripped,
                nonempty_index=nonempty_index,
                previous_kind=previous_kind,
                in_reference_section=in_reference_section,
            )
            if kind == "reference_heading":
                in_reference_section = True
                previous_kind = kind
                continue
            if kind != "body":
                if kind == "abstract_inline":
                    inline = self._match_docx_inline_front_label(source_text)
                    if inline and inline[2]:
                        rewritten_body = self._transform_text(inline[2], task_type, platform, summary)
                        rewritten_text = f"{inline[1]}{rewritten_body}" if rewritten_body else source_text
                        if rewritten_text and rewritten_text != source_text:
                            self._redistribute_paragraph_text(paragraph, rewritten_text)
                previous_kind = kind
                continue
            rewritten_text = self._transform_text(source_text, task_type, platform, summary)
            if rewritten_text and rewritten_text != source_text:
                self._redistribute_paragraph_text(paragraph, rewritten_text)
            previous_kind = kind
        doc.save(str(output_path))

    def _collect_full_docx_rewrite_targets(self, doc: Document) -> tuple[list, list[str], list[str], list[str]]:
        target_paragraphs: list = []
        source_paragraphs: list[str] = []
        target_kinds: list[str] = []
        preserved_prefixes: list[str] = []
        in_reference_section = False
        nonempty_index = 0
        previous_kind = ""
        for paragraph in self._iter_body_paragraphs(doc):
            source_text = paragraph.text or ""
            stripped = source_text.strip()
            if not stripped:
                continue
            nonempty_index += 1
            kind = self._classify_docx_paragraph_kind(
                paragraph,
                stripped,
                nonempty_index=nonempty_index,
                previous_kind=previous_kind,
                in_reference_section=in_reference_section,
            )
            if kind == "reference_heading":
                in_reference_section = True
                previous_kind = kind
                continue
            if kind == "abstract_inline":
                inline = self._match_docx_inline_front_label(source_text)
                if inline and inline[2]:
                    target_paragraphs.append(paragraph)
                    source_paragraphs.append(inline[2])
                    target_kinds.append(kind)
                    preserved_prefixes.append(inline[1])
                previous_kind = kind
                continue
            if kind != "body":
                previous_kind = kind
                continue
            target_paragraphs.append(paragraph)
            source_paragraphs.append(source_text)
            target_kinds.append(kind)
            preserved_prefixes.append("")
            previous_kind = kind
        return target_paragraphs, source_paragraphs, target_kinds, preserved_prefixes

    def _transform_docx_cnki_pipeline(self, doc: Document, task_type: TaskType) -> None:
        if task_type == TaskType.REWRITE:
            from app.services.cnki_rewrite_runtime import execute_cnki_rewrite as execute_cnki_pipeline
            quality_flags = {"cnki_rewrite_pipeline_applied": True}
            strategy_version = "cnki_rewrite_v1"
            mode = "cnki_rewrite_style_transfer"
        else:
            from app.services.cnki_dedup_runtime import execute_cnki_dedup as execute_cnki_pipeline
            quality_flags = {"cnki_dedup_pipeline_applied": True}
            strategy_version = "cnki_dedup_v1"
            mode = "cnki_dedup_style_transfer"

        target_paragraphs, source_paragraphs, _target_kinds, preserved_prefixes = self._collect_full_docx_rewrite_targets(doc)
        if not source_paragraphs:
            return
        run = execute_cnki_pipeline(self.db, task_type=task_type, paragraphs=source_paragraphs)
        rewritten_paragraphs = [item.strip() for item in re.split(r"\n\s*\n", str(run.text or "").strip()) if item.strip()]
        if len(rewritten_paragraphs) != len(target_paragraphs):
            rewritten_paragraphs = [line.strip() for line in str(run.text or "").splitlines() if line.strip()]
        if len(rewritten_paragraphs) != len(target_paragraphs):
            raise BizError(code=4626, message="知网 DOCX 输出段落数不匹配")
        for paragraph, rewritten_text, source_text, target_kind, preserved_prefix in zip(
            target_paragraphs, rewritten_paragraphs, source_paragraphs, _target_kinds, preserved_prefixes
        ):
            if preserved_prefix:
                if target_kind == "abstract_inline":
                    self._replace_docx_inline_label_paragraph(paragraph, preserved_prefix, rewritten_text)
                    continue
                rewritten_text = f"{preserved_prefix}{rewritten_text}"
            if rewritten_text and rewritten_text != source_text:
                self._redistribute_paragraph_text(paragraph, rewritten_text)
        meta = {
            "strategy": "llm",
            "platform": "cnki",
            "task_type": task_type.value,
            "length_before": run.total_chars,
            "length_after": count_billable_chars(run.text),
            "quality_score": 1.0,
            "quality_flags": quality_flags,
            "warnings": [],
            "rule_trace": {
                "mode": mode,
                "strategy_version": strategy_version,
                "paragraph_count": run.paragraph_count,
                "total_chars": run.total_chars,
                "llm_provider": run.llm_provider,
                "llm_model": run.llm_model,
            },
        }
        if task_type == TaskType.REWRITE:
            meta["change_ratio"] = round(((meta["length_after"] - meta["length_before"]) / max(meta["length_before"], 1)), 4)
            self._rewrite_strategy_meta = meta
        else:
            similarity = SequenceMatcher(None, "\n\n".join(source_paragraphs)[:4000], run.text[:4000]).ratio()
            meta["length_delta_ratio"] = round(((meta["length_after"] - meta["length_before"]) / max(meta["length_before"], 1)), 4)
            meta["similarity_ratio"] = round(similarity, 4)
            meta["change_ratio"] = round((1 - similarity) * 100, 2)
            self._dedup_strategy_meta = meta
        self._pipeline_usage["llm_used"] = True

    def _transform_docx_vip_pipeline(self, doc: Document, task_type: TaskType) -> None:
        if task_type == TaskType.REWRITE:
            from app.services.vip_rewrite_runtime import execute_vip_rewrite as execute_vip_pipeline
            quality_flags = {"vip_rewrite_pipeline_applied": True}
            strategy_version = "vip_rewrite_v1"
            mode = "vip_rewrite_style_transfer"
        else:
            from app.services.vip_dedup_runtime import execute_vip_dedup as execute_vip_pipeline
            quality_flags = {"vip_dedup_pipeline_applied": True}
            strategy_version = "vip_dedup_v1"
            mode = "vip_dedup_style_transfer"

        target_paragraphs, source_paragraphs, _target_kinds, preserved_prefixes = self._collect_full_docx_rewrite_targets(doc)
        if not source_paragraphs:
            return
        run = execute_vip_pipeline(self.db, task_type=task_type, paragraphs=source_paragraphs)
        rewritten_paragraphs = [item.strip() for item in re.split(r"\n\s*\n", str(run.text or "").strip()) if item.strip()]
        if len(rewritten_paragraphs) != len(target_paragraphs):
            rewritten_paragraphs = [line.strip() for line in str(run.text or "").splitlines() if line.strip()]
        if len(rewritten_paragraphs) != len(target_paragraphs):
            raise BizError(code=4636, message="维普 DOCX 输出段落数不匹配")
        for paragraph, rewritten_text, source_text, target_kind, preserved_prefix in zip(
            target_paragraphs, rewritten_paragraphs, source_paragraphs, _target_kinds, preserved_prefixes
        ):
            if preserved_prefix:
                if target_kind == "abstract_inline":
                    self._replace_docx_inline_label_paragraph(paragraph, preserved_prefix, rewritten_text)
                    continue
                rewritten_text = f"{preserved_prefix}{rewritten_text}"
            if rewritten_text and rewritten_text != source_text:
                self._redistribute_paragraph_text(paragraph, rewritten_text)
        meta = {
            "strategy": "llm",
            "platform": "vip",
            "task_type": task_type.value,
            "length_before": run.total_chars,
            "length_after": count_billable_chars(run.text),
            "quality_score": 1.0,
            "quality_flags": quality_flags,
            "warnings": [],
            "rule_trace": {
                "mode": mode,
                "strategy_version": strategy_version,
                "paragraph_count": run.paragraph_count,
                "total_chars": run.total_chars,
                "llm_provider": run.llm_provider,
                "llm_model": run.llm_model,
            },
        }
        if task_type == TaskType.REWRITE:
            meta["change_ratio"] = round(((meta["length_after"] - meta["length_before"]) / max(meta["length_before"], 1)), 4)
            self._rewrite_strategy_meta = meta
        else:
            similarity = SequenceMatcher(None, "\n\n".join(source_paragraphs)[:4000], run.text[:4000]).ratio()
            meta["length_delta_ratio"] = round(((meta["length_after"] - meta["length_before"]) / max(meta["length_before"], 1)), 4)
            meta["similarity_ratio"] = round(similarity, 4)
            meta["change_ratio"] = round((1 - similarity) * 100, 2)
            self._dedup_strategy_meta = meta
        self._pipeline_usage["llm_used"] = True

    def _run_llm(self, task_type: TaskType, text: str) -> str | None:
        switch = self._current_switch or self._get_or_init_switch()
        if self._effective_mode != MODE_LLM_PLUS_ALGO or not switch.llm_enabled:
            return None
        try:
            output = generate_with_llm(self.db, task_type=task_type, text=text)
            switch.llm_fail_count = 0
            self._pipeline_usage["llm_used"] = True
            self.db.flush()
            return output
        except Exception as exc:
            switch.llm_fail_count += 1
            should_downgrade = switch.llm_fail_count >= switch.llm_fail_threshold
            if should_downgrade:
                self._switch_mode(MODE_ALGO_ONLY, f"llm_fail:{str(exc)[:160]}")
            self.db.add(
                LLMErrorLog(
                    task_id=self._current_task_id,
                    error_type=exc.__class__.__name__,
                    error_detail=str(exc)[:500],
                    trigger_downgrade=should_downgrade,
                )
            )
            self.db.flush()
            return None

    def _parse_llm_detect_result(self, raw: str | None) -> dict | None:
        if not isinstance(raw, str) or not raw.strip():
            return None

        payload: dict | None = None
        text = raw.strip()
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            try:
                candidate = json.loads(match.group(0))
                if isinstance(candidate, dict):
                    payload = candidate
            except Exception:
                payload = None

        if payload is None:
            return self._parse_llm_detect_fallback(text)

        score = self._coerce_ratio(payload.get("ai_score"))
        if score is None:
            score = self._coerce_ratio(payload.get("score"))
        if score is None:
            return self._parse_llm_detect_fallback(text)

        label = self._normalize_detect_label(payload.get("label"))
        if not label:
            label = self._normalize_detect_label(payload.get("risk_level"))
        reason = payload.get("reason")
        return {
            "ai_score": score,
            "label": label,
            "reason": str(reason).strip()[:180] if isinstance(reason, str) else "",
        }

    def _parse_llm_detect_fallback(self, text: str) -> dict | None:
        percent_match = re.search(r"(\d+(?:\.\d+)?)\s*%", text)
        if percent_match:
            score = self._coerce_ratio(percent_match.group(1))
        else:
            score = self._coerce_ratio(self._first_numeric(text))
        if score is None:
            return None
        return {
            "ai_score": score,
            "label": self._normalize_detect_label(text),
            "reason": "",
        }

    def _first_numeric(self, text: str) -> str | None:
        match = re.search(r"(\d+(?:\.\d+)?)", text)
        if not match:
            return None
        return match.group(1)

    def _coerce_ratio(self, raw) -> float | None:
        if raw is None:
            return None
        if isinstance(raw, str):
            raw = raw.strip().replace("%", "")
        try:
            value = float(raw)
        except Exception:
            return None
        if value > 1:
            value = value / 100
        return round(self._clamp_score(value), 4)

    def _normalize_detect_label(self, raw) -> str:
        value = str(raw or "").strip().lower()
        if not value:
            return ""
        if any(token in value for token in ("clean", "no_ai", "safe", "normal", "none")):
            return "clean"
        if "high" in value or "高" in value:
            return "high"
        if "medium" in value or "mid" in value or "中" in value:
            return "medium"
        if "low" in value or "低" in value:
            return "low"
        return ""

    def _transform_text(self, text: str, task_type: TaskType, platform: str, report_summary: dict) -> str:
        normalized_input = self._normalize_text(text)

        if task_type == TaskType.REWRITE and platform in {"cnki", "vip"}:
            result = execute_rewrite_strategy(
                self.db,
                task=None,
                platform=platform,
                text=normalized_input,
                report_summary=report_summary,
            )
            self._rewrite_strategy_meta = {
                "strategy": result.get("strategy"),
                "platform": result.get("platform"),
                "task_type": result.get("task_type"),
                "length_before": result.get("length_before"),
                "length_after": result.get("length_after"),
                "change_ratio": result.get("change_ratio"),
                "quality_score": result.get("quality_score"),
                "quality_flags": result.get("quality_flags") or {},
                "warnings": result.get("warnings") or [],
                "rule_trace": result.get("rule_trace") or {},
            }
            if result.get("strategy") == "llm":
                self._pipeline_usage["llm_used"] = True
            return str(result.get("rewritten_text") or "")

        if task_type == TaskType.DEDUP and platform in {"cnki", "vip"}:
            dedup_strategy = get_active_dedup_strategy(self.db, platform=platform)
            result = execute_dedup_strategy(
                self.db,
                task=None,
                platform=platform,
                text=normalized_input,
                report_summary=report_summary,
                strategy=dedup_strategy,
            )
            self._dedup_strategy_meta = {
                "strategy": result.get("strategy"),
                "platform": result.get("platform"),
                "task_type": result.get("task_type"),
                "length_before": result.get("length_before"),
                "length_after": result.get("length_after"),
                "length_delta_ratio": result.get("length_delta_ratio"),
                "similarity_ratio": result.get("similarity_ratio"),
                "change_ratio": result.get("change_ratio"),
                "quality_score": result.get("quality_score"),
                "quality_flags": result.get("quality_flags") or {},
                "warnings": result.get("warnings") or [],
                "rule_trace": result.get("rule_trace") or {},
            }
            if result.get("strategy") == DEDUP_STRATEGY_LLM:
                self._pipeline_usage["llm_used"] = True
            return str(result.get("rewritten_text") or "")

        if self._effective_mode == MODE_LLM_PLUS_ALGO:
            llm_output = self._run_llm(task_type, normalized_input)
            if isinstance(llm_output, str) and llm_output.strip():
                return llm_output
            return self._heuristic_transform_text(
                text=normalized_input,
                task_type=task_type,
                report_summary=report_summary,
            )

        return self._heuristic_transform_text(
            text=normalized_input,
            task_type=task_type,
            report_summary=report_summary,
        )

    def _heuristic_transform_text(self, *, text: str, task_type: TaskType, report_summary: dict) -> str:
        pressure = report_summary.get("pressure", "low")
        if task_type == TaskType.DEDUP:
            replacements = {
                "因此": "由此可见",
                "但是": "然而",
                "首先": "第一",
                "其次": "第二",
                "总之": "综上所述",
                "可以看出": "据此可见",
                "本文认为": "本文进一步指出",
            }
            output = self._apply_replacements(text, replacements)
            output = self._split_long_sentences(output, 48 if pressure == "high" else 64)
            return output
        if task_type == TaskType.REWRITE:
            replacements = {
                "研究表明": "已有研究指出",
                "可以看出": "据此可见",
                "非常": "较为",
                "重要": "关键",
                "很多": "大量",
                "我们发现": "研究发现",
                "这个": "该",
            }
            output = self._apply_replacements(text, replacements)
            output = self._split_long_sentences(output, 54 if pressure == "high" else 72)
            return output
        return text
    def _normalize_text(self, text: str) -> str:
        output = re.sub(r"[ \t]+", " ", text)
        output = re.sub(r"\n{3,}", "\n\n", output)
        return output.strip()

    def _apply_replacements(self, text: str, replacements: dict[str, str]) -> str:
        output = text
        for src, target in replacements.items():
            output = output.replace(src, target)
        return output

    def _split_long_sentences(self, text: str, threshold: int) -> str:
        return split_long_sentences(text, threshold)

    def _text_stats(self, text: str) -> dict:
        clean = text.strip()
        sentences = [part.strip() for part in re.split(r"[。！？!?；;\n]+", clean) if part.strip()]
        paragraphs = [part.strip() for part in clean.splitlines() if part.strip()]
        sentence_count = len(sentences)
        avg_sentence_length = round(sum(len(item) for item in sentences) / sentence_count, 2) if sentence_count else 0
        return {
            "char_count": count_billable_chars(clean),
            "paragraph_count": len(paragraphs),
            "sentence_count": sentence_count,
            "avg_sentence_length": avg_sentence_length,
        }

    def _extract_report_summary(self, task_type: TaskType, report_text: str) -> dict:
        return extract_report_summary(task_type, report_text)

    def _extract_percent(self, text: str, keywords: list[str]) -> float | None:
        return extract_percent(text, keywords)


    def _clamp_score(self, value: float) -> float:
        return max(0.0, min(1.0, float(value)))

    def _legacy_platform_detect_profile_v000(self, platform: str) -> dict:
        key = (platform or "").strip().lower()
        profiles = {
            "cnki": {
                "name": "cnki_like",
                "baseline_weight": 0.7,
                "style_weight": 0.2,
                "repeat_weight": 0.1,
                "offset": 0.0,
                "high": 0.65,
                "medium": 0.35,
            },
            "vip": {
                "name": "vip_like",
                "baseline_weight": 0.64,
                "style_weight": 0.24,
                "repeat_weight": 0.12,
                "offset": -0.02,
                "high": 0.62,
                "medium": 0.33,
            },
        }
        return profiles.get(key, profiles["cnki"])

    def _legacy_simulate_platform_detect_score_v00(self, platform: str, text: str, base_score: float) -> tuple[float, dict, dict]:
        profile = self._platform_detect_profile(platform)
        stats = self._text_stats(text)
        clean = (text or "").strip()
        compact = " ".join(clean.split())
        unique_ratio = (len(set(compact)) / len(compact)) if compact else 1.0
        repeat_signal = self._clamp_score(1.0 - unique_ratio)
        avg_len = float(stats.get("avg_sentence_length") or 0.0)
        style_signal = self._clamp_score((avg_len - 18.0) / 55.0)

        weighted = (
            float(base_score) * profile["baseline_weight"]
            + style_signal * profile["style_weight"]
            + repeat_signal * profile["repeat_weight"]
            + profile["offset"]
        )
        score = round(self._clamp_score(weighted), 4)
        breakdown = {
            "base_score": round(float(base_score), 4),
            "style_signal": round(style_signal, 4),
            "repeat_signal": round(repeat_signal, 4),
            "weights": {
                "baseline": profile["baseline_weight"],
                "style": profile["style_weight"],
                "repeat": profile["repeat_weight"],
                "offset": profile["offset"],
            },
            "thresholds": {
                "high": profile["high"],
                "medium": profile["medium"],
            },
        }
        return score, profile, breakdown

    def _risk_band(self, score: float, *, high: float = 0.65, medium: float = 0.35) -> str:
        if score >= high:
            return "高风险"
        if score >= medium:
            return "中风险"
        return "低风险"

    def _legacy_top_risk_paragraphs_v00(self, text: str, platform: str = "cnki") -> list[dict]:
        paragraphs = [part.strip() for part in text.splitlines() if part.strip()]
        scored = []
        for index, paragraph in enumerate(paragraphs, start=1):
            if len(paragraph) < 20:
                continue
            base_score = self._heuristic_ai_score(paragraph)
            simulated_score, _profile, _breakdown = self._simulate_platform_detect_score(platform, paragraph, base_score)
            scored.append(
                {
                    "index": index,
                    "score": round(simulated_score * 100, 2),
                    "excerpt": self._clip_text(paragraph, 80),
                }
            )
        scored.sort(key=lambda item: item["score"], reverse=True)
        return scored[:3]


    def _wrap_pdf_line(self, text: str, width: int = 90) -> list[str]:
        compact = " ".join(str(text or "").split())
        if not compact:
            return [""]
        wrapped: list[str] = []
        remaining = compact
        while len(remaining) > width:
            cut = remaining.rfind(" ", 0, width + 1)
            if cut <= 0:
                cut = width
            wrapped.append(remaining[:cut].strip())
            remaining = remaining[cut:].strip()
        if remaining:
            wrapped.append(remaining)
        return wrapped

    def _pdf_safe_text(self, text: str) -> str:
        return str(text or "").encode("latin-1", "replace").decode("latin-1")

    def _pdf_escape(self, text: str) -> str:
        return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    def _render_pdf(self, lines: list[str]) -> bytes:
        page_width = 595
        page_height = 842
        margin_x = 46
        margin_top = 60
        margin_bottom = 52
        font_size = 11
        line_height = 15

        expanded_lines: list[str] = []
        for line in lines:
            wrapped = self._wrap_pdf_line(line, width=92)
            expanded_lines.extend(wrapped if wrapped else [""])
        if not expanded_lines:
            expanded_lines = [""]

        usable_height = page_height - margin_top - margin_bottom
        lines_per_page = max(1, int(usable_height // line_height))
        page_line_chunks = [
            expanded_lines[i : i + lines_per_page] for i in range(0, len(expanded_lines), lines_per_page)
        ]

        objects: list[tuple[int, bytes]] = [(1, b"<< /Type /Catalog /Pages 2 0 R >>")]
        page_refs: list[str] = []
        next_obj_id = 3
        font_obj_id = 2 + len(page_line_chunks) * 2 + 1

        for chunk in page_line_chunks:
            page_obj_id = next_obj_id
            content_obj_id = next_obj_id + 1
            next_obj_id += 2
            page_refs.append(f"{page_obj_id} 0 R")

            text_ops: list[str] = []
            for index, line in enumerate(chunk):
                y = page_height - margin_top - index * line_height
                safe_line = self._pdf_escape(self._pdf_safe_text(line))
                text_ops.append(f"BT /F1 {font_size} Tf 1 0 0 1 {margin_x} {y:.2f} Tm ({safe_line}) Tj ET")
            stream_text = "\n".join(text_ops)
            stream_bytes = stream_text.encode("latin-1", "replace")

            page_obj = (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] "
                f"/Resources << /Font << /F1 {font_obj_id} 0 R >> >> /Contents {content_obj_id} 0 R >>"
            ).encode("ascii")
            content_obj = (
                f"<< /Length {len(stream_bytes)} >>\nstream\n".encode("ascii")
                + stream_bytes
                + b"\nendstream"
            )
            objects.append((page_obj_id, page_obj))
            objects.append((content_obj_id, content_obj))

        pages_obj = f"<< /Type /Pages /Count {len(page_line_chunks)} /Kids [{' '.join(page_refs)}] >>".encode("ascii")
        objects.insert(1, (2, pages_obj))
        objects.append((font_obj_id, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))

        objects.sort(key=lambda item: item[0])
        output = bytearray()
        output.extend(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets: dict[int, int] = {}

        for obj_id, obj_body in objects:
            offsets[obj_id] = len(output)
            output.extend(f"{obj_id} 0 obj\n".encode("ascii"))
            output.extend(obj_body)
            output.extend(b"\nendobj\n")

        xref_offset = len(output)
        max_obj_id = max(offsets)
        output.extend(f"xref\n0 {max_obj_id + 1}\n".encode("ascii"))
        output.extend(b"0000000000 65535 f \n")
        for obj_id in range(1, max_obj_id + 1):
            offset = offsets.get(obj_id, 0)
            output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        output.extend(f"trailer\n<< /Size {max_obj_id + 1} /Root 1 0 R >>\n".encode("ascii"))
        output.extend(f"startxref\n{xref_offset}\n%%EOF".encode("ascii"))
        return bytes(output)
    def _legacy_platform_detect_profile_v0(self, platform: str) -> dict:
        key = (platform or "").strip().lower()
        profiles = {
            "cnki": {
                "name": "cnki_like",
                "provider_label": "知网AIGC检测仿真",
                "score_label": "AIGC值",
                "baseline_weight": 0.56,
                "style_weight": 0.14,
                "repeat_weight": 0.12,
                "template_weight": 0.10,
                "context_weight": 0.08,
                "offset": 0.0,
                "high": 0.67,
                "medium": 0.42,
                "overall_paragraph_weight": 0.56,
                "overall_peak_weight": 0.26,
                "overall_segment_weight": 0.18,
            },
            "vip": {
                "name": "vip_like",
                "provider_label": "维普AIGC检测仿真",
                "score_label": "AIGC疑似度",
                "baseline_weight": 0.52,
                "style_weight": 0.16,
                "repeat_weight": 0.12,
                "template_weight": 0.11,
                "context_weight": 0.09,
                "offset": -0.02,
                "high": 0.64,
                "medium": 0.40,
                "overall_paragraph_weight": 0.54,
                "overall_peak_weight": 0.24,
                "overall_segment_weight": 0.22,
            },
        }
        return profiles.get(key, profiles["cnki"])

    def _platform_detect_profile(self, platform: str) -> dict:
        key = (platform or "").strip().lower()
        profiles = {
            "cnki": {
                "name": "cnki_like",
                "provider_label": "知网AIGC检测仿真",
                "score_label": "AI特征值",
                "baseline_weight": 0.56,
                "style_weight": 0.14,
                "repeat_weight": 0.12,
                "template_weight": 0.10,
                "context_weight": 0.08,
                "opening_weight": 0.06,
                "offset": 0.0,
                "high": 0.67,
                "medium": 0.42,
                "clean": 0.18,
                "coverage_weight": 0.06,
                "section_weight": 0.08,
                "streak_weight": 0.03,
                "opening_similarity_weight": 0.02,
                "evidence_relief_weight": 0.06,
                "colloquial_relief_weight": 0.42,
                "specificity_relief_weight": 0.30,
                "artifact_weight": 0.12,
                "english_abstract_weight": 0.08,
                "abstract_section_weight": 0.08,
                "intro_section_weight": 0.05,
                "fragment_display_thresholds": {"mild": 0.70, "moderate": 0.80, "severe": 0.90},
            },
            "vip": {
                "name": "vip_like",
                "provider_label": "维普AIGC检测仿真",
                "score_label": "全文疑似AIGC生成",
                "baseline_weight": 0.52,
                "style_weight": 0.16,
                "repeat_weight": 0.12,
                "template_weight": 0.11,
                "context_weight": 0.09,
                "opening_weight": 0.07,
                "offset": -0.02,
                "high": 0.64,
                "medium": 0.40,
                "clean": 0.18,
                "coverage_weight": 0.08,
                "section_weight": 0.08,
                "streak_weight": 0.03,
                "opening_similarity_weight": 0.02,
                "evidence_relief_weight": 0.05,
                "colloquial_relief_weight": 0.48,
                "specificity_relief_weight": 0.34,
                "artifact_weight": 0.06,
                "english_abstract_weight": 0.03,
                "abstract_section_weight": 0.03,
                "intro_section_weight": 0.02,
                "fragment_display_thresholds": {"mild": 0.70, "moderate": 0.80, "severe": 0.90},
            },
        }
        return profiles.get(key, profiles["cnki"])

    def _legacy_template_signal_v1(self, text: str) -> tuple[float, list[str]]:
        hits: list[str] = []
        phrases = [
            "研究表明",
            "本研究旨在",
            "本文基于",
            "在此背景下",
            "可以看出",
            "由此可见",
            "综上所述",
            "总而言之",
            "值得注意的是",
            "首先",
            "其次",
            "再次",
            "最后",
            "此外",
            "与此同时",
            "另一方面",
            "进一步而言",
            "从整体上看",
            "不难发现",
            "基于此",
            "具有重要的理论意义和实践意义",
            "为此提供参考",
            "从而实现",
        ]
        content = str(text or "")
        for phrase in phrases:
            if phrase in content:
                hits.append(phrase)
        density = len(hits) / max(len(content) / 85.0, 1.0)
        return round(self._clamp_score(density), 4), hits[:4]

    def _legacy_opening_signal_v1(self, text: str) -> float:
        clean = " ".join(str(text or "").split())
        normalized = re.sub(
            r"^(?:第[一二三四五六七八九十百零0-9]+[章节部分篇]|[一二三四五六七八九十百零]+[、.．]|"
            r"\d+(?:\.\d+){0,3}[、.．]?|[（(][一二三四五六七八九十百零0-9]+[)）])",
            "",
            clean,
        ).lstrip("：:.。;；、，, ")
        if not normalized:
            return 0.0

        starters = (
            "本研究",
            "本文",
            "研究表明",
            "综上所述",
            "总而言之",
            "值得注意的是",
            "首先",
            "其次",
            "再次",
            "此外",
            "与此同时",
            "另一方面",
            "进一步",
            "从整体上看",
            "由此可见",
            "在此背景下",
            "基于此",
        )
        signal = 0.0
        if any(normalized.startswith(token) for token in starters):
            signal += 0.48
        order_hits = sum(1 for token in ("首先", "其次", "再次", "最后", "此外", "另一方面") if token in normalized[:24])
        signal += min(0.36, order_hits * 0.12)
        if re.search(r"^在.{0,8}背景下", normalized):
            signal += 0.18
        return round(self._clamp_score(signal), 4)

    def _legacy_citation_relief_signal_v1(self, text: str) -> float:
        content = str(text or "")
        patterns = [r"\[\d+\]", r"（\d{4}）", r"\(\d{4}\)", r"表\d+", r"图\d+", r"\d+%"]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content))
        return round(min(0.18, hit_count * 0.025), 4)

    def _legacy_evidence_relief_signal_v1(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"\bN\s*=\s*\d+",
            r"样本量",
            r"问卷",
            r"访谈",
            r"实验",
            r"受访者",
            r"标准差",
            r"均值",
            r"统计",
            r"案例",
            r"观察记录",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        return round(min(0.16, hit_count * 0.018), 4)

    def _legacy_uniformity_signal_v1(self, sentences: list[str]) -> float:
        lengths = [len(item.strip()) for item in sentences if item.strip()]
        if not lengths:
            return 0.0
        if len(lengths) == 1:
            return 0.5
        avg_len = sum(lengths) / len(lengths)
        variance = sum((value - avg_len) ** 2 for value in lengths) / len(lengths)
        signal = 1.0 - min(1.0, variance / max(avg_len * 14.0, 1.0))
        return round(self._clamp_score(signal), 4)

    def _legacy_normalized_opening_key_v1(self, paragraph: str) -> str:
        clean = " ".join(str(paragraph or "").split())
        if not clean:
            return ""
        heading_level, _heading = self._detect_outline_heading(clean)
        if heading_level is not None:
            return ""
        normalized = re.sub(
            r"^(?:第[一二三四五六七八九十百零0-9]+[章节部分篇]|[一二三四五六七八九十百零]+[、.．]|"
            r"\d+(?:\.\d+){0,3}[、.．]?|[（(][一二三四五六七八九十百零0-9]+[)）])",
            "",
            clean,
        ).lstrip("：:.。;；、，, ")
        normalized = re.sub(r"\s+", "", normalized)
        return normalized[:12]

    def _legacy_paragraph_opening_similarity_v1(self, paragraphs: list[str]) -> float:
        keys: list[str] = []
        starter_counts: dict[str, int] = {}
        starters = (
            "本研究",
            "本文",
            "研究表明",
            "首先",
            "其次",
            "再次",
            "最后",
            "此外",
            "另一方面",
            "值得注意的是",
            "综上所述",
            "由此可见",
            "基于此",
        )
        for paragraph in paragraphs:
            key = self._normalized_opening_key(paragraph)
            if len(key) < 6:
                continue
            keys.append(key)
            for starter in starters:
                if key.startswith(starter):
                    starter_counts[starter] = starter_counts.get(starter, 0) + 1
                    break
        if not keys:
            return 0.0
        duplicate_ratio = max(0, len(keys) - len(set(keys))) / max(len(keys), 1)
        starter_repeat_ratio = sum(max(0, count - 1) for count in starter_counts.values()) / max(len(keys), 1)
        return round(self._clamp_score(duplicate_ratio * 0.68 + starter_repeat_ratio * 0.62), 4)

    def _legacy_fragment_display_band_v1(self, score: float, profile: dict) -> str:
        thresholds = profile.get("fragment_display_thresholds") or {}
        if score >= float(thresholds.get("severe", 0.9)):
            return "severe"
        if score >= float(thresholds.get("moderate", 0.8)):
            return "moderate"
        if score >= float(thresholds.get("mild", 0.7)):
            return "mild"
        return ""

    def _legacy_simulate_platform_detect_score_v0(self, platform: str, text: str, base_score: float) -> tuple[float, dict, dict]:
        profile = self._platform_detect_profile(platform)
        stats = self._text_stats(text)
        clean = (text or "").strip()
        compact = " ".join(clean.split())
        unique_ratio = (len(set(compact)) / len(compact)) if compact else 1.0
        repeat_signal = self._clamp_score(1.0 - unique_ratio)
        avg_len = float(stats.get("avg_sentence_length") or 0.0)
        style_signal = self._clamp_score((avg_len - 18.0) / 52.0)
        template_signal, template_hits = self._template_signal(clean)
        context_signal = self._uniformity_signal(self._split_detect_sentences(clean))
        opening_signal = self._opening_signal(clean)
        citation_relief = self._citation_relief_signal(clean)
        evidence_relief = self._evidence_relief_signal(clean)

        weighted = (
            float(base_score) * profile["baseline_weight"]
            + style_signal * profile["style_weight"]
            + repeat_signal * profile["repeat_weight"]
            + template_signal * profile["template_weight"]
            + context_signal * profile["context_weight"]
            + opening_signal * profile["opening_weight"]
            - citation_relief
            - evidence_relief
            + profile["offset"]
        )
        score = round(self._clamp_score(weighted), 4)
        breakdown = {
            "base_score": round(float(base_score), 4),
            "style_signal": round(style_signal, 4),
            "repeat_signal": round(repeat_signal, 4),
            "template_signal": round(template_signal, 4),
            "context_signal": round(context_signal, 4),
            "opening_signal": round(opening_signal, 4),
            "citation_relief": round(citation_relief, 4),
            "evidence_relief": round(evidence_relief, 4),
            "template_hits": template_hits,
            "weights": {
                "baseline": profile["baseline_weight"],
                "style": profile["style_weight"],
                "repeat": profile["repeat_weight"],
                "template": profile["template_weight"],
                "context": profile["context_weight"],
                "opening": profile["opening_weight"],
                "offset": profile["offset"],
            },
            "thresholds": {"high": profile["high"], "medium": profile["medium"]},
        }
        return score, profile, breakdown

    def _legacy_local_suspicious_segments_v1(self, paragraph: str, platform: str, profile: dict) -> list[dict]:
        segments: list[dict] = []
        for sentence in self._split_detect_sentences(paragraph):
            if len(sentence) < 8:
                continue
            base_score = self._heuristic_ai_score(sentence)
            score, _profile, breakdown = self._simulate_platform_detect_score(platform, sentence, base_score)
            if score < float(profile.get("medium", 0.35)) and len(breakdown.get("template_hits") or []) < 2:
                continue
            reasons: list[str] = []
            if float(breakdown.get("template_signal") or 0.0) >= 0.22:
                reasons.append("模板连接词偏多")
            if float(breakdown.get("repeat_signal") or 0.0) >= 0.32:
                reasons.append("重复表达偏多")
            if float(breakdown.get("context_signal") or 0.0) >= 0.58:
                reasons.append("句式波动偏小")
            if float(breakdown.get("opening_signal") or 0.0) >= 0.44:
                reasons.append("段首表述模板化")
            segments.append(
                {
                    "text": self._clip_text(sentence, 76),
                    "score": round(score * 100, 2),
                    "reason": "、".join(reasons[:2]) or "综合风险偏高",
                }
            )
        segments.sort(key=lambda item: item["score"], reverse=True)
        return segments[:3]

    def _legacy_paragraph_reason_tags_v1(self, breakdown: dict) -> list[str]:
        tags: list[str] = []
        if float(breakdown.get("template_signal") or 0.0) >= 0.22:
            tags.append("模板连接词偏多")
        if float(breakdown.get("repeat_signal") or 0.0) >= 0.30:
            tags.append("重复表达偏高")
        if float(breakdown.get("context_signal") or 0.0) >= 0.56:
            tags.append("句式波动偏小")
        if float(breakdown.get("style_signal") or 0.0) >= 0.55:
            tags.append("长句占比较高")
        if float(breakdown.get("opening_signal") or 0.0) >= 0.44:
            tags.append("段首表述模板化")
        return tags[:3]

    def _legacy_build_fragment_distribution_v1(
        self,
        text: str,
        platform: str,
        profile: dict,
        paragraph_details: list[dict],
        document_outline: list[dict] | None = None,
    ) -> dict:
        paragraphs = self._split_detect_paragraphs(text)
        paragraph_score_map = {
            int(item.get("index") or 0): round(float(item.get("score") or 0.0) / 100.0, 4) for item in paragraph_details
        }
        count_map = {"high": 0, "medium": 0, "low": 0, "no_ai": 0}
        char_map = {"high": 0, "medium": 0, "low": 0, "no_ai": 0}
        display_count_map = {"mild": 0, "moderate": 0, "severe": 0}
        display_char_map = {"mild": 0, "moderate": 0, "severe": 0}
        weighted_scores: list[float] = []

        for paragraph_index, paragraph in enumerate(paragraphs, start=1):
            paragraph_ratio = paragraph_score_map.get(paragraph_index, 0.0)
            for sentence in self._split_detect_sentences(paragraph):
                compact = " ".join(sentence.split())
                if not compact:
                    continue
                char_count = count_billable_chars(compact)
                if char_count <= 0:
                    continue
                if self._is_no_ai_fragment(compact):
                    label = "no_ai"
                    score_ratio = 0.0
                else:
                    base_score = self._heuristic_ai_score(compact)
                    score_ratio, _profile, _breakdown = self._simulate_platform_detect_score(platform, compact, base_score)
                    if paragraph_ratio > 0:
                        score_ratio = round(self._clamp_score(score_ratio * 0.72 + paragraph_ratio * 0.28), 4)
                    label = self._score_to_detect_label(score_ratio, profile)
                    if label == "clean":
                        label = "no_ai"
                    else:
                        weighted_scores.append(score_ratio)
                        display_band = self._fragment_display_band(score_ratio, profile)
                        if display_band:
                            display_count_map[display_band] += 1
                            display_char_map[display_band] += char_count
                count_map[label] += 1
                char_map[label] += char_count

        total_fragments = sum(count_map.values())
        total_chars = sum(char_map.values())
        if total_fragments <= 0 or total_chars <= 0:
            return {
                "fragment_count": 0,
                "high_fragment_count": 0,
                "middle_fragment_count": 0,
                "low_fragment_count": 0,
                "no_ai_fragment_count": 0,
                "high_suspected_fragment_ratio": 0.0,
                "middle_suspected_fragment_ratio": 0.0,
                "low_suspected_fragment_ratio": 0.0,
                "no_ai_fragment_ratio": 0.0,
                "high_and_middle_suspected_fragment_ratio": 0.0,
                "total_suspected_fragment_ratio": 0.0,
                "high_suspected_text_ratio": 0.0,
                "middle_suspected_text_ratio": 0.0,
                "low_suspected_text_ratio": 0.0,
                "no_ai_suspected_text_ratio": 0.0,
                "high_and_middle_suspected_text_ratio": 0.0,
                "total_suspected_text_ratio": 0.0,
                "weighted_score_pct": 0.0,
                "mild_fragment_count": 0,
                "moderate_fragment_count": 0,
                "severe_fragment_count": 0,
                "mild_fragment_ratio": 0.0,
                "moderate_fragment_ratio": 0.0,
                "severe_fragment_ratio": 0.0,
                "mild_text_ratio": 0.0,
                "moderate_text_ratio": 0.0,
                "severe_text_ratio": 0.0,
                "display_thresholds": {"mild": 70.0, "moderate": 80.0, "severe": 90.0},
            }

        def ratio(part: int, whole: int) -> float:
            return round(part / max(whole, 1) * 100, 2)

        high_middle_fragments = count_map["high"] + count_map["medium"]
        suspected_fragments = high_middle_fragments + count_map["low"]
        high_middle_chars = char_map["high"] + char_map["medium"]
        suspected_chars = high_middle_chars + char_map["low"]
        thresholds = profile.get("fragment_display_thresholds") or {}
        return {
            "fragment_count": total_fragments,
            "high_fragment_count": count_map["high"],
            "middle_fragment_count": count_map["medium"],
            "low_fragment_count": count_map["low"],
            "no_ai_fragment_count": count_map["no_ai"],
            "high_suspected_fragment_ratio": ratio(count_map["high"], total_fragments),
            "middle_suspected_fragment_ratio": ratio(count_map["medium"], total_fragments),
            "low_suspected_fragment_ratio": ratio(count_map["low"], total_fragments),
            "no_ai_fragment_ratio": ratio(count_map["no_ai"], total_fragments),
            "high_and_middle_suspected_fragment_ratio": ratio(high_middle_fragments, total_fragments),
            "total_suspected_fragment_ratio": ratio(suspected_fragments, total_fragments),
            "high_suspected_text_ratio": ratio(char_map["high"], total_chars),
            "middle_suspected_text_ratio": ratio(char_map["medium"], total_chars),
            "low_suspected_text_ratio": ratio(char_map["low"], total_chars),
            "no_ai_suspected_text_ratio": ratio(char_map["no_ai"], total_chars),
            "high_and_middle_suspected_text_ratio": ratio(high_middle_chars, total_chars),
            "total_suspected_text_ratio": ratio(suspected_chars, total_chars),
            "weighted_score_pct": round(sum(weighted_scores) / max(len(weighted_scores), 1) * 100, 2)
            if weighted_scores
            else 0.0,
            "mild_fragment_count": display_count_map["mild"],
            "moderate_fragment_count": display_count_map["moderate"],
            "severe_fragment_count": display_count_map["severe"],
            "mild_fragment_ratio": ratio(display_count_map["mild"], total_fragments),
            "moderate_fragment_ratio": ratio(display_count_map["moderate"], total_fragments),
            "severe_fragment_ratio": ratio(display_count_map["severe"], total_fragments),
            "mild_text_ratio": ratio(display_char_map["mild"], total_chars),
            "moderate_text_ratio": ratio(display_char_map["moderate"], total_chars),
            "severe_text_ratio": ratio(display_char_map["severe"], total_chars),
            "display_thresholds": {
                "mild": round(float(thresholds.get("mild", 0.7)) * 100, 2),
                "moderate": round(float(thresholds.get("moderate", 0.8)) * 100, 2),
                "severe": round(float(thresholds.get("severe", 0.9)) * 100, 2),
            },
        }

    def _legacy_build_document_metrics_v1(
        self,
        *,
        text: str,
        paragraph_details: list[dict],
        section_distribution: list[dict],
        document_outline: list[dict],
        profile: dict,
    ) -> dict:
        total = len(paragraph_details)
        total_chars = sum(int(item.get("char_count") or 0) for item in paragraph_details)
        high_medium_count = sum(1 for item in paragraph_details if item.get("label") in {"high", "medium"})
        high_medium_chars = sum(
            int(item.get("char_count") or 0) for item in paragraph_details if item.get("label") in {"high", "medium"}
        )
        longest_streak = 0
        current_streak = 0
        for item in paragraph_details:
            if item.get("label") in {"high", "medium"}:
                current_streak += 1
                longest_streak = max(longest_streak, current_streak)
            else:
                current_streak = 0
        medium_threshold = float(profile.get("medium", 0.35)) * 100
        suspicious_sections = sum(
            1
            for item in section_distribution
            if float(item.get("avg_score") or 0.0) >= medium_threshold or int(item.get("high_count") or 0) > 0
        )
        abstract_scores = [
            float(item.get("avg_score") or 0.0)
            for item in section_distribution
            if self._section_category(item.get("section") or "") == "abstract"
        ]
        intro_scores = [
            float(item.get("avg_score") or 0.0)
            for item in section_distribution
            if self._section_category(item.get("section") or "") == "intro"
        ]
        review_scores = [
            float(item.get("avg_score") or 0.0)
            for item in section_distribution
            if self._section_category(item.get("section") or "") == "review"
        ]
        paragraphs = self._split_detect_paragraphs(text)
        opening_similarity = self._paragraph_opening_similarity(paragraphs)
        evidence_relief = max(self._citation_relief_signal_v2(text), self._evidence_relief_signal_v2(text))
        return {
            "paragraph_count": total,
            "high_medium_paragraph_count": high_medium_count,
            "high_medium_paragraph_ratio": round(high_medium_count / max(total, 1) * 100, 2) if total else 0.0,
            "high_medium_text_ratio": round(high_medium_chars / max(total_chars, 1) * 100, 2) if total_chars else 0.0,
            "longest_risk_streak": longest_streak,
            "longest_risk_streak_ratio": round(longest_streak / max(total, 1) * 100, 2) if total else 0.0,
            "opening_similarity_ratio": round(opening_similarity * 100, 2),
            "section_count": len(section_distribution),
            "suspicious_section_count": suspicious_sections,
            "section_coverage_ratio": round(suspicious_sections / max(len(section_distribution), 1) * 100, 2)
            if section_distribution
            else 0.0,
            "outline_sections": len(document_outline),
            "distribution_mode": "outline" if document_outline else "band",
            "evidence_relief_pct": round(evidence_relief * 100, 2),
            "abstract_avg_score": round(sum(abstract_scores) / max(len(abstract_scores), 1), 2) if abstract_scores else 0.0,
            "intro_avg_score": round(sum(intro_scores) / max(len(intro_scores), 1), 2) if intro_scores else 0.0,
            "review_avg_score": round(sum(review_scores) / max(len(review_scores), 1), 2) if review_scores else 0.0,
        }

    def _legacy_build_decision_basis_v1(
        self,
        *,
        breakdown: dict,
        document_metrics: dict,
        fragment_distribution: dict,
        suspicious_segments: list[dict],
    ) -> list[dict]:
        items: list[dict] = []
        template_signal = float(breakdown.get("template_signal") or 0.0)
        context_signal = float(breakdown.get("context_signal") or 0.0)
        repeat_signal = float(breakdown.get("repeat_signal") or 0.0)
        opening_signal = float(breakdown.get("opening_signal") or 0.0)
        template_hits = [str(item) for item in (breakdown.get("template_hits") or []) if str(item).strip()]

        if template_signal >= 0.2:
            hit_text = "、".join(template_hits[:3]) if template_hits else "高频模板连接词"
            items.append(
                {
                    "title": "模板连接词密度偏高",
                    "detail": f"命中 {hit_text}，模板信号 {round(template_signal * 100, 2)}%。",
                    "direction": "risk",
                }
            )
        if context_signal >= 0.55:
            items.append(
                {
                    "title": "句长波动偏小",
                    "detail": f"上下文均匀度 {round(context_signal * 100, 2)}%，更接近机器批量生成的平直节奏。",
                    "direction": "risk",
                }
            )
        if repeat_signal >= 0.3:
            items.append(
                {
                    "title": "重复表达偏多",
                    "detail": f"重复信号 {round(repeat_signal * 100, 2)}%，存在近义重复或固定总结句反复出现。",
                    "direction": "risk",
                }
            )
        if opening_signal >= 0.44 or float(document_metrics.get("opening_similarity_ratio") or 0.0) >= 25:
            items.append(
                {
                    "title": "段首表达重复度较高",
                    "detail": f"段首相似度 {document_metrics.get('opening_similarity_ratio', 0.0)}%，段落展开方式较集中。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("section_coverage_ratio") or 0.0) >= 35:
            items.append(
                {
                    "title": "中高风险内容跨章节分布",
                    "detail": f"可疑章节覆盖率 {document_metrics.get('section_coverage_ratio', 0.0)}%，不是单点异常。",
                    "direction": "risk",
                }
            )
        if int(document_metrics.get("longest_risk_streak") or 0) >= 3:
            items.append(
                {
                    "title": "存在连续风险片段带",
                    "detail": f"最长连续中高风险段落 {document_metrics.get('longest_risk_streak', 0)} 段。",
                    "direction": "risk",
                }
            )
        if float(fragment_distribution.get("high_and_middle_suspected_text_ratio") or 0.0) >= 20:
            items.append(
                {
                    "title": "高中风险文字占比较高",
                    "detail": f"高中风险文字占比 {fragment_distribution.get('high_and_middle_suspected_text_ratio', 0.0)}%。",
                    "direction": "risk",
                }
            )
        if float(fragment_distribution.get("severe_fragment_ratio") or 0.0) >= 5:
            items.append(
                {
                    "title": "存在重度疑似片段",
                    "detail": f"90%以上片段占比 {fragment_distribution.get('severe_fragment_ratio', 0.0)}%。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("evidence_relief_pct") or 0.0) >= 6:
            items.append(
                {
                    "title": "引文与样本证据较多",
                    "detail": f"实证减权 {document_metrics.get('evidence_relief_pct', 0.0)}%，用于抑制纯模板误判。",
                    "direction": "relief",
                }
            )
        if not items:
            items.append(
                {
                    "title": "以全文统计与片段聚合综合判定",
                    "detail": f"当前识别到 {len(suspicious_segments)} 个可疑片段，建议人工复核摘要、引言与结论段。",
                    "direction": "risk",
                }
            )
        return items[:6]

    def _split_detect_sentences(self, text: str) -> list[str]:
        return [seg.strip() for seg in re.split(r"[。！？!?；;\n]+", str(text or "")) if seg.strip()]

    def _normalize_detect_text_line(self, text: str) -> str:
        return re.sub(r"\s+", " ", str(text or "")).strip()

    def _looks_like_wrapped_detect_block(self, lines: list[str]) -> bool:
        compact_lengths = [len(re.sub(r"\s+", "", line)) for line in lines if str(line or "").strip()]
        if len(compact_lengths) < 4:
            return False
        median_len = float(statistics.median(compact_lengths))
        short_ratio = sum(1 for item in compact_lengths if item <= 46) / max(len(compact_lengths), 1)
        long_ratio = sum(1 for item in compact_lengths if item >= 95) / max(len(compact_lengths), 1)
        if len(compact_lengths) >= 80 and median_len <= 32 and short_ratio >= 0.5:
            return True
        return median_len <= 46 and short_ratio >= 0.55 and long_ratio <= 0.45

    def _looks_like_detect_paragraph_start(self, text: str) -> bool:
        compact = re.sub(r"\s+", "", str(text or ""))
        if not compact:
            return False
        starters = (
            "首先",
            "其次",
            "再次",
            "最后",
            "此外",
            "另外",
            "与此同时",
            "综上",
            "总之",
            "由此可见",
            "由此可知",
            "一方面",
            "另一方面",
            "针对上述",
            "基于上述",
        )
        if compact.startswith(starters):
            return True
        return bool(re.match(r"^(注|说明|备注|来源)[:：]", compact))

    def _join_detect_wrapped_lines(self, left: str, right: str) -> str:
        left_text = self._normalize_detect_text_line(left)
        right_text = self._normalize_detect_text_line(right)
        if not left_text:
            return right_text
        if not right_text:
            return left_text
        if re.search(r"[\u4e00-\u9fff（([“‘]$", left_text) or re.match(r"^[\u4e00-\u9fff，。；：！？、）】》」]", right_text):
            return f"{left_text}{right_text}"
        return f"{left_text} {right_text}".strip()

    def _should_break_wrapped_detect_paragraph(self, buffer_text: str, next_line: str) -> bool:
        current = self._normalize_detect_text_line(next_line)
        if not current:
            return False
        current_level, _current_heading = self._detect_outline_heading(current)
        if current_level is not None:
            return True

        buffer_compact = re.sub(r"\s+", "", str(buffer_text or ""))
        if not buffer_compact:
            return False
        buffer_level, _buffer_heading = self._detect_outline_heading(buffer_compact)
        if buffer_level is not None and len(buffer_compact) <= 40:
            return True

        previous_complete = bool(re.search(r"[。！？!?；;：:]$", self._normalize_detect_text_line(buffer_text)))
        if not previous_complete:
            return False

        if len(buffer_compact) >= 170:
            return True
        if len(buffer_compact) >= 110 and len(re.sub(r"\s+", "", current)) >= 18:
            return True
        if len(buffer_compact) >= 85 and self._looks_like_detect_paragraph_start(current):
            return True
        return False

    def _split_detect_paragraphs(self, text: str) -> list[str]:
        paragraphs = split_detect_strategy_paragraphs(str(text or ""))
        if paragraphs:
            return paragraphs
        normalized = self._normalize_detect_text_line(text)
        return [normalized] if normalized else []

    def _build_detect_llm_excerpt(self, text: str, platform: str) -> str:
        clean = str(text or "").strip()
        paragraphs = self._split_detect_paragraphs(clean)
        if len(clean) <= 4800 and len(paragraphs) <= 12:
            return clean

        scored: list[tuple[int, float]] = []
        for index, paragraph in enumerate(paragraphs, start=1):
            if len(paragraph) < 20:
                continue
            base_score = self._heuristic_ai_score(paragraph)
            score, _profile, _breakdown = self._simulate_platform_detect_score(platform, paragraph, base_score)
            scored.append((index, score))

        keep: set[int] = set(range(1, min(len(paragraphs), 2) + 1))
        if paragraphs:
            keep.update(range(max(1, len(paragraphs) - 1), len(paragraphs) + 1))
        for index, _score in sorted(scored, key=lambda item: item[1], reverse=True)[:4]:
            keep.add(index)

        excerpt = "\n".join(paragraphs[index - 1] for index in sorted(keep) if 1 <= index <= len(paragraphs)).strip()
        if len(excerpt) > 4200:
            excerpt = excerpt[:4200]
        return excerpt or clean[:4200]

    def _template_signal(self, text: str) -> tuple[float, list[str]]:
        hits: list[str] = []
        phrases = [
            "研究表明",
            "本研究旨在",
            "本文基于",
            "在此背景下",
            "可以看出",
            "由此可见",
            "综上所述",
            "总而言之",
            "值得注意的是",
            "首先",
            "其次",
            "再次",
            "最后",
            "此外",
            "与此同时",
            "另一方面",
            "进一步而言",
            "从整体上看",
            "不难发现",
            "基于此",
        ]
        content = str(text or "")
        for phrase in phrases:
            if phrase in content:
                hits.append(phrase)
        density = len(hits) / max(len(content) / 85.0, 1.0)
        return round(self._clamp_score(density), 4), hits[:4]

    def _citation_relief_signal(self, text: str) -> float:
        content = str(text or "")
        patterns = [r"\[\d+\]", r"（\d{4}）", r"\(\d{4}\)", r"表\d+", r"图\d+", r"\d+%"]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content))
        return round(min(0.18, hit_count * 0.025), 4)

    def _uniformity_signal(self, sentences: list[str]) -> float:
        lengths = [len(item.strip()) for item in sentences if item.strip()]
        if not lengths:
            return 0.0
        if len(lengths) == 1:
            return 0.5
        avg_len = sum(lengths) / len(lengths)
        variance = sum((value - avg_len) ** 2 for value in lengths) / len(lengths)
        signal = 1.0 - min(1.0, variance / max(avg_len * 14.0, 1.0))
        return round(self._clamp_score(signal), 4)

    def _opening_signal(self, text: str) -> float:
        clean = re.sub(
            r"^(?:第[一二三四五六七八九十百零0-9]+[章节部分篇]|[一二三四五六七八九十百零]+[、.．]|"
            r"\d+(?:\.\d+){0,3}[、.．]?|[（(][一二三四五六七八九十百零0-9]+[)）])",
            "",
            " ".join(str(text or "").split()),
        ).lstrip("：:.。;；、，, ")
        starters = (
            "本研究",
            "本文",
            "研究表明",
            "综上所述",
            "总而言之",
            "值得注意的是",
            "首先",
            "其次",
            "再次",
            "此外",
            "另一方面",
            "由此可见",
            "基于此",
        )
        signal = 0.48 if any(clean.startswith(token) for token in starters) else 0.0
        if re.search(r"^在.{0,8}背景下", clean):
            signal += 0.18
        return round(self._clamp_score(signal), 4)

    def _evidence_relief_signal(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"\bN\s*=\s*\d+",
            r"样本量",
            r"问卷",
            r"访谈",
            r"实验",
            r"受访者",
            r"标准差",
            r"均值",
            r"统计",
            r"案例",
            r"观察记录",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        return round(min(0.16, hit_count * 0.018), 4)

    def _colloquial_relief_signal(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"觉得",
            r"就.{0,4}(?:而言|来看|来说|情况)",
            r"还要|还需|还应|还会",
            r"让(?:其|学生|教师|家长|学校)?",
            r"能够",
            r"很多|这样|不太|其实|并不是|不只是|很难",
            r"给(?:予|出|足|全体|学生|孩子)?",
            r"做法",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content))
        density = hit_count / max(len(content) / 120.0, 1.0)
        return round(min(0.18, density * 0.12), 4)

    def _specificity_relief_signal(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"(?:19|20)\d{2}年",
            r"《[^》]{2,30}》",
            r"国家|教育部|课程标准|指导意见|实施方案|条例|通知|文件|制度",
            r"案例|问卷|调查|样本|统计|实验|访谈|观察记录",
            r"表\d+|图\d+|\[\d+\]|（\d{4}）|\(\d{4}\)",
            r"某市|某校|A公司|B公司",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        density = hit_count / max(len(content) / 150.0, 1.0)
        return round(min(0.16, density * 0.08), 4)

    def _template_signal_v2(self, text: str) -> tuple[float, list[str]]:
        hits: list[str] = []
        phrases = [
            "研究表明",
            "本研究旨在",
            "本研究以",
            "本文基于",
            "在此背景下",
            "可以看出",
            "由此可见",
            "综上所述",
            "总而言之",
            "值得注意的是",
            "首先",
            "其次",
            "再次",
            "最后",
            "此外",
            "与此同时",
            "另一方面",
            "进一步而言",
            "从整体上看",
            "不难发现",
            "基于上述",
            "研究结论显示",
            "系统分析",
            "构建",
            "搭建",
            "实施保障",
        ]
        content = str(text or "")
        for phrase in phrases:
            if phrase in content:
                hits.append(phrase)
        density = len(hits) / max(len(content) / 85.0, 1.0)
        return round(self._clamp_score(density), 4), hits[:5]

    def _citation_relief_signal_v2(self, text: str) -> float:
        content = str(text or "")
        patterns = [r"\[\d+\]", r"（\d{4}）", r"\(\d{4}\)", r"表\d+", r"图\d+", r"\d+%"]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content))
        return round(min(0.18, hit_count * 0.025), 4)

    def _opening_signal_v2(self, text: str) -> float:
        clean = re.sub(
            r"^(?:第[一二三四五六七八九十百零0-9]+[章节部分篇]|[一二三四五六七八九十百零]+[、.．]|"
            r"\d+(?:\.\d+){0,3}[、.．]?|[（(][一二三四五六七八九十百零0-9]+[)）])",
            "",
            " ".join(str(text or "").split()),
        ).lstrip("：:.。;；、，, ")
        starters = (
            "本研究",
            "本文",
            "研究表明",
            "研究发现",
            "综上所述",
            "总而言之",
            "值得注意的是",
            "首先",
            "其次",
            "再次",
            "此外",
            "另一方面",
            "由此可见",
            "基于上述",
        )
        signal = 0.48 if any(clean.startswith(token) for token in starters) else 0.0
        if re.search(r"^在.{0,8}背景下", clean):
            signal += 0.18
        return round(self._clamp_score(signal), 4)

    def _evidence_relief_signal_v2(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"\bN\s*=\s*\d+",
            r"样本量",
            r"问卷",
            r"访谈",
            r"实验",
            r"受访者",
            r"标准差",
            r"均值",
            r"统计",
            r"案例",
            r"观察记录",
            r"调查",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        return round(min(0.16, hit_count * 0.018), 4)

    def _specificity_relief_signal_v2(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"(?:19|20)\d{2}年",
            r"《[^》]{2,30}》",
            r"国家|教育部|课程标准|指导意见|实施方案|条例|通知|文件|制度",
            r"案例|问卷|调查|样本|统计|实验|访谈|观察记录",
            r"表\d+|图\d+|\[\d+\]|（\d{4}）|\(\d{4}\)",
            r"某市|某校|A公司|B公司",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        density = hit_count / max(len(content) / 150.0, 1.0)
        return round(min(0.16, density * 0.08), 4)

    def _human_case_relief_signal(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"“[^”]{2,20}”",
            r"——",
            r"幼儿|孩子|家长|教师|老师|顾客|员工|门店|会员|妈妈|爸爸|爷爷|奶奶",
            r"本园|我园|本校|本班|班级|课堂|晨圈|园所",
            r"案例|项目|活动|分享|记录|观察|反馈|实践",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        density = hit_count / max(len(content) / 110.0, 1.0)
        return round(min(0.24, density * 0.09), 4)

    def _cnki_practice_chain_signal(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"课程融合|晨圈分享|家长参与|成效与反思",
            r"绘本阅读|手工绘画|演唱活动|家人小调查|家长体验课|家长助教活动|家长感悟接龙",
            r"角色互换|双向流动|情感共鸣|进入社区|主动参与|表达更主动|参与质量",
            r"制作贺卡|爱心甜汤|调查活动|分享环节|晨圈时间|节日活动",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        density = hit_count / max(len(content) / 120.0, 1.0)
        return round(min(0.24, density * 0.11), 4)

    def _cnki_summary_wrapup_relief(self, text: str) -> float:
        content = str(text or "")
        patterns = [
            r"最核心的教育逻辑|最本质的教育逻辑",
            r"可迁移性体现在两个层面|具有极强的可迁移性",
            r"形成系列化的家园共育体系",
            r"推广价值|生命底色|不长在课堂里",
        ]
        hit_count = 0
        for pattern in patterns:
            hit_count += len(re.findall(pattern, content, flags=re.IGNORECASE))
        density = hit_count / max(len(content) / 110.0, 1.0)
        return round(min(0.20, density * 0.12), 4)

    def _artifact_signal(self, text: str) -> tuple[float, list[str]]:
        content = str(text or "")
        markers = [
            "需要求",
            "知识得到",
            "主题着眼于",
            "读全面本书",
            "都衡",
            "全面性",
            "让用",
            "成效",
            "机械替换",
            "目的模糊等困难",
        ]
        hits = [marker for marker in markers if marker in content]
        density = len(hits) / max(len(content) / 120.0, 1.0)
        score = min(0.2, density * 0.14 + (0.04 if hits else 0.0))
        return round(self._clamp_score(score), 4), hits[:4]

    def _english_abstract_signal(self, text: str) -> float:
        content = str(text or "")
        lower = content.lower()
        ascii_ratio = len(re.findall(r"[A-Za-z]", content)) / max(len(content), 1)
        hints = ("abstract", "keywords", "this study", "this paper", "based on", "findings show", "research object")
        hit_count = sum(1 for hint in hints if hint in lower)
        score = 0.0
        if "abstract" in lower or "keywords" in lower:
            score += 0.06
        if ascii_ratio >= 0.18:
            score += min(0.06, ascii_ratio * 0.18)
        if hit_count >= 2:
            score += min(0.05, hit_count * 0.02)
        return round(min(0.18, score), 4)

    def _calibrate_detect_score(
        self,
        *,
        platform: str,
        raw_score: float,
        breakdown: dict,
        distribution: dict,
        fragment_distribution: dict,
        document_metrics: dict,
        source_stats: dict,
    ) -> float:
        raw_pct = round(float(raw_score or 0.0) * 100.0, 2)
        template_pct = float(breakdown.get("template_signal") or 0.0) * 100.0
        context_pct = float(breakdown.get("context_signal") or 0.0) * 100.0
        opening_pct = float(breakdown.get("opening_signal") or 0.0) * 100.0
        citation_pct = float(breakdown.get("citation_relief") or 0.0) * 100.0
        evidence_pct = float(breakdown.get("evidence_relief") or 0.0) * 100.0
        colloquial_pct = float(breakdown.get("colloquial_relief") or 0.0) * 100.0
        specificity_pct = float(breakdown.get("specificity_relief") or 0.0) * 100.0
        human_case_pct = float(breakdown.get("human_case_relief") or 0.0) * 100.0
        weighted_fragment_pct = float(fragment_distribution.get("weighted_score_pct") or 0.0)
        high_middle_text_pct = float(fragment_distribution.get("high_and_middle_suspected_text_ratio") or 0.0)
        high_ratio_pct = float(distribution.get("high_ratio") or 0.0)
        char_count = int(source_stats.get("char_count") or 0)
        long_doc_relief = max(0.0, min(8.0, (char_count - 4500) / 650.0)) if char_count >= 4500 else 0.0

        key = (platform or "").strip().lower()
        if key == "cnki":
            calibrated_pct = (
                raw_pct * 0.30
                + template_pct * 0.50
                + context_pct * 0.08
                + opening_pct * 0.20
                + weighted_fragment_pct * 0.08
                + high_ratio_pct * 0.12
                - citation_pct * 0.12
                - evidence_pct * 0.60
                - colloquial_pct * 0.40
                - specificity_pct * 0.45
                - human_case_pct * 0.55
                - long_doc_relief * 0.30
                - 0.60
            )
            if template_pct < 8.0 and weighted_fragment_pct < 10.0:
                calibrated_pct -= 1.5
            if raw_pct >= 20.0:
                calibrated_pct = max(calibrated_pct, raw_pct * 0.38)
            if calibrated_pct < 2.0:
                calibrated_pct = 0.0
        elif key == "vip":
            calibrated_pct = (
                template_pct * 0.42
                + context_pct * 0.26
                + opening_pct * 0.22
                + weighted_fragment_pct * 0.08
                + high_ratio_pct * 0.22
                - citation_pct * 0.18
                - evidence_pct * 1.45
                - colloquial_pct * 0.88
                - specificity_pct * 0.95
                - long_doc_relief * 0.90
                - 3.80
            )
            if template_pct < 10.0 and high_ratio_pct < 3.0:
                calibrated_pct -= 2.0
            if calibrated_pct < 3.0:
                calibrated_pct = 0.0
        else:
            calibrated_pct = raw_pct

        return round(self._clamp_score(max(0.0, calibrated_pct) / 100.0), 4)

    def _should_expand_detect_details(self, platform: str, score_pct: float) -> bool:
        return score_pct >= 8.0

    def _normalized_opening_key(self, paragraph: str) -> str:
        clean = " ".join(str(paragraph or "").split())
        if not clean:
            return ""
        heading_level, _heading = self._detect_outline_heading(clean)
        if heading_level is not None:
            return ""
        normalized = re.sub(
            r"^(?:第[一二三四五六七八九十百零0-9]+[章节部分篇]|[一二三四五六七八九十百零]+[、.．]|"
            r"\d+(?:\.\d+){0,3}[、.．]?|[（(][一二三四五六七八九十百零0-9]+[)）])",
            "",
            clean,
        ).lstrip("：:.。;；、，, ")
        normalized = re.sub(r"\s+", "", normalized)
        return normalized[:12]

    def _paragraph_opening_similarity(self, paragraphs: list[str]) -> float:
        keys: list[str] = []
        starter_counts: dict[str, int] = {}
        starters = (
            "本研究",
            "本文",
            "研究表明",
            "首先",
            "其次",
            "再次",
            "最后",
            "此外",
            "另一方面",
            "值得注意的是",
            "综上所述",
            "由此可见",
            "基于此",
        )
        for paragraph in paragraphs:
            key = self._normalized_opening_key(paragraph)
            if len(key) < 6:
                continue
            keys.append(key)
            for starter in starters:
                if key.startswith(starter):
                    starter_counts[starter] = starter_counts.get(starter, 0) + 1
                    break
        if not keys:
            return 0.0
        duplicate_ratio = max(0, len(keys) - len(set(keys))) / max(len(keys), 1)
        starter_repeat_ratio = sum(max(0, count - 1) for count in starter_counts.values()) / max(len(keys), 1)
        return round(self._clamp_score(duplicate_ratio * 0.68 + starter_repeat_ratio * 0.62), 4)

    def _fragment_display_band(self, score: float, profile: dict) -> str:
        thresholds = profile.get("fragment_display_thresholds") or {}
        if score >= float(thresholds.get("severe", 0.9)):
            return "severe"
        if score >= float(thresholds.get("moderate", 0.8)):
            return "moderate"
        if score >= float(thresholds.get("mild", 0.7)):
            return "mild"
        return ""

    def _simulate_platform_detect_score(self, platform: str, text: str, base_score: float) -> tuple[float, dict, dict]:
        profile = self._platform_detect_profile(platform)
        stats = self._text_stats(text)
        clean = (text or "").strip()
        compact = " ".join(clean.split())
        unique_ratio = (len(set(compact)) / len(compact)) if compact else 1.0
        repeat_signal = self._clamp_score(1.0 - unique_ratio)
        avg_len = float(stats.get("avg_sentence_length") or 0.0)
        style_signal = self._clamp_score((avg_len - 18.0) / 52.0)
        template_signal, template_hits = self._template_signal_v2(clean)
        context_signal = self._uniformity_signal(self._split_detect_sentences(clean))
        opening_signal = self._opening_signal_v2(clean)
        citation_relief = self._citation_relief_signal_v2(clean)
        evidence_relief = self._evidence_relief_signal_v2(clean)
        colloquial_relief = self._colloquial_relief_signal(clean)
        specificity_relief = self._specificity_relief_signal_v2(clean)
        human_case_relief = self._human_case_relief_signal(clean)
        artifact_signal, artifact_hits = self._artifact_signal(clean)
        english_abstract_signal = self._english_abstract_signal(clean)
        practice_chain_signal = 0.0
        summary_wrapup_relief = 0.0
        effective_human_case_relief = human_case_relief
        if (platform or "").strip().lower() == "cnki":
            practice_chain_signal = self._cnki_practice_chain_signal(clean)
            summary_wrapup_relief = self._cnki_summary_wrapup_relief(clean)
            if practice_chain_signal >= 0.08 and summary_wrapup_relief < 0.10:
                effective_human_case_relief = round(
                    max(0.0, human_case_relief * (1.0 - min(0.58, practice_chain_signal * 2.2))),
                    4,
                )

        weighted = (
            float(base_score) * profile["baseline_weight"]
            + style_signal * profile["style_weight"]
            + repeat_signal * profile["repeat_weight"]
            + template_signal * profile["template_weight"]
            + context_signal * profile["context_weight"]
            + opening_signal * profile["opening_weight"]
            + artifact_signal * profile.get("artifact_weight", 0.0)
            + english_abstract_signal * profile.get("english_abstract_weight", 0.0)
            + practice_chain_signal * 0.16
            - citation_relief
            - evidence_relief
            - colloquial_relief * profile.get("colloquial_relief_weight", 0.0)
            - specificity_relief * profile.get("specificity_relief_weight", 0.0)
            - effective_human_case_relief
            - summary_wrapup_relief * 0.24
            + profile["offset"]
        )
        score = round(self._clamp_score(weighted), 4)
        breakdown = {
            "base_score": round(float(base_score), 4),
            "style_signal": round(style_signal, 4),
            "repeat_signal": round(repeat_signal, 4),
            "template_signal": round(template_signal, 4),
            "context_signal": round(context_signal, 4),
            "opening_signal": round(opening_signal, 4),
            "citation_relief": round(citation_relief, 4),
            "evidence_relief": round(evidence_relief, 4),
            "colloquial_relief": round(colloquial_relief, 4),
            "specificity_relief": round(specificity_relief, 4),
            "human_case_relief": round(effective_human_case_relief, 4),
            "human_case_relief_raw": round(human_case_relief, 4),
            "artifact_signal": round(artifact_signal, 4),
            "artifact_hits": artifact_hits,
            "english_abstract_signal": round(english_abstract_signal, 4),
            "practice_chain_signal": round(practice_chain_signal, 4),
            "summary_wrapup_relief": round(summary_wrapup_relief, 4),
            "template_hits": template_hits,
            "weights": {
                "baseline": profile["baseline_weight"],
                "style": profile["style_weight"],
                "repeat": profile["repeat_weight"],
                "template": profile["template_weight"],
                "context": profile["context_weight"],
                "opening": profile["opening_weight"],
                "artifact": profile.get("artifact_weight", 0.0),
                "english_abstract": profile.get("english_abstract_weight", 0.0),
                "colloquial_relief": profile.get("colloquial_relief_weight", 0.0),
                "specificity_relief": profile.get("specificity_relief_weight", 0.0),
                "human_case_relief": 1.0,
                "offset": profile["offset"],
            },
            "thresholds": {"high": profile["high"], "medium": profile["medium"]},
        }
        return score, profile, breakdown

    def _score_to_detect_label(self, score: float, profile: dict) -> str:
        if score >= float(profile.get("high", 0.65)):
            return "high"
        if score >= float(profile.get("medium", 0.35)):
            return "medium"
        if score >= float(profile.get("clean", 0.18)):
            return "low"
        return "clean"

    def _prefer_higher_risk_detect_label(self, primary: str, secondary: str) -> str:
        rank = {"clean": 0, "low": 1, "medium": 2, "high": 3}
        left = self._normalize_detect_label(primary)
        right = self._normalize_detect_label(secondary)
        if not left:
            return right
        if not right:
            return left
        return left if rank.get(left, -1) >= rank.get(right, -1) else right

    def _local_suspicious_segments(self, paragraph: str, platform: str, profile: dict) -> list[dict]:
        segments: list[dict] = []
        for sentence in self._split_detect_sentences(paragraph):
            if len(sentence) < 8:
                continue
            base_score = self._heuristic_ai_score(sentence)
            score, _profile, breakdown = self._simulate_platform_detect_score(platform, sentence, base_score)
            if score < float(profile.get("medium", 0.35)) and len(breakdown.get("template_hits") or []) < 2:
                continue
            reasons: list[str] = []
            if float(breakdown.get("template_signal") or 0.0) >= 0.22:
                reasons.append("模板衔接词偏多")
            if float(breakdown.get("repeat_signal") or 0.0) >= 0.32:
                reasons.append("重复表达偏多")
            if float(breakdown.get("context_signal") or 0.0) >= 0.58:
                reasons.append("句式波动较小")
            if float(breakdown.get("opening_signal") or 0.0) >= 0.44:
                reasons.append("段首表述模板化")
            segments.append(
                {
                    "text": self._clip_text(sentence, 76),
                    "score": round(score * 100, 2),
                    "reason": "、".join(reasons[:2]) or "综合风险偏高",
                }
            )
        segments.sort(key=lambda item: item["score"], reverse=True)
        return segments[:3]

    def _merge_suspicious_segments(self, *group_lists) -> list[dict]:
        merged: list[dict] = []
        seen: set[str] = set()
        for group in group_lists:
            if not isinstance(group, list):
                continue
            for item in group:
                if not isinstance(item, dict):
                    continue
                text = self._clip_text(str(item.get("text") or item.get("excerpt") or ""), 76)
                if not text or text in seen:
                    continue
                seen.add(text)
                merged.append(
                    {
                        "text": text,
                        "score": round(float(item.get("score") or 0.0), 2),
                        "reason": str(item.get("reason") or "").strip(),
                    }
                )
        merged.sort(key=lambda item: item["score"], reverse=True)
        return merged[:4]

    def _paragraph_reason_tags(self, breakdown: dict) -> list[str]:
        tags: list[str] = []
        if float(breakdown.get("template_signal") or 0.0) >= 0.22:
            tags.append("模板衔接词偏多")
        if float(breakdown.get("repeat_signal") or 0.0) >= 0.30:
            tags.append("重复表达偏高")
        if float(breakdown.get("context_signal") or 0.0) >= 0.56:
            tags.append("句式波动偏小")
        if float(breakdown.get("style_signal") or 0.0) >= 0.55:
            tags.append("长句占比较高")
        if float(breakdown.get("opening_signal") or 0.0) >= 0.44:
            tags.append("段首表述模板化")
        return tags[:3]

    def _local_suspicious_segments_v2(self, paragraph: str, platform: str, profile: dict) -> list[dict]:
        segments: list[dict] = []
        for sentence in self._split_detect_sentences(paragraph):
            if len(sentence) < 8:
                continue
            base_score = self._heuristic_ai_score(sentence)
            score, _profile, breakdown = self._simulate_platform_detect_score(platform, sentence, base_score)
            strong_signal_count = sum(
                1
                for passed in (
                    float(breakdown.get("template_signal") or 0.0) >= 0.22,
                    float(breakdown.get("repeat_signal") or 0.0) >= 0.32,
                    float(breakdown.get("context_signal") or 0.0) >= 0.58,
                    float(breakdown.get("opening_signal") or 0.0) >= 0.44,
                    float(breakdown.get("artifact_signal") or 0.0) >= 0.08,
                    float(breakdown.get("english_abstract_signal") or 0.0) >= 0.08,
                )
                if passed
            )
            relief_signal = max(
                float(breakdown.get("citation_relief") or 0.0),
                float(breakdown.get("evidence_relief") or 0.0),
            )
            practice_chain_signal = float(breakdown.get("practice_chain_signal") or 0.0)
            summary_wrapup_relief = float(breakdown.get("summary_wrapup_relief") or 0.0)
            cnki_practice_allow = (
                platform == "cnki"
                and practice_chain_signal >= 0.10
                and summary_wrapup_relief < 0.10
                and score >= float(profile.get("clean", 0.18)) + 0.04
            )
            if not cnki_practice_allow and score < float(profile.get("medium", 0.35)) + 0.08 and strong_signal_count < 2:
                continue
            if relief_signal >= 0.08 and score < float(profile.get("high", 0.65)):
                continue
            if platform == "cnki" and summary_wrapup_relief >= 0.10 and score < float(profile.get("high", 0.65)):
                continue
            if len(sentence) < 16 and score < float(profile.get("high", 0.65)):
                continue
            reasons: list[str] = []
            if float(breakdown.get("template_signal") or 0.0) >= 0.22:
                reasons.append("模板连接词偏多")
            if float(breakdown.get("repeat_signal") or 0.0) >= 0.32:
                reasons.append("重复表达偏多")
            if float(breakdown.get("context_signal") or 0.0) >= 0.58:
                reasons.append("句式波动偏小")
            if float(breakdown.get("opening_signal") or 0.0) >= 0.44:
                reasons.append("段首展开方式过于集中")
            if float(breakdown.get("artifact_signal") or 0.0) >= 0.08:
                reasons.append("存在异常改写痕迹")
            if float(breakdown.get("english_abstract_signal") or 0.0) >= 0.08:
                reasons.append("英文摘要镜像痕迹偏强")
            segments.append(
                {
                    "text": self._clip_text(sentence, 76),
                    "score": round(score * 100, 2),
                    "reason": "、".join(reasons[:2]) or "综合风险偏高",
                }
            )
        segments.sort(key=lambda item: item["score"], reverse=True)
        return segments[:2]

    def _paragraph_reason_tags_v2(self, breakdown: dict) -> list[str]:
        tags: list[str] = []
        if float(breakdown.get("template_signal") or 0.0) >= 0.22:
            tags.append("模板连接词偏多")
        if float(breakdown.get("repeat_signal") or 0.0) >= 0.30:
            tags.append("重复表达偏高")
        if float(breakdown.get("context_signal") or 0.0) >= 0.56:
            tags.append("句式波动偏小")
        if float(breakdown.get("style_signal") or 0.0) >= 0.55:
            tags.append("长句占比偏高")
        if float(breakdown.get("opening_signal") or 0.0) >= 0.44:
            tags.append("段首展开方式集中")
        if float(breakdown.get("artifact_signal") or 0.0) >= 0.08:
            tags.append("存在异常改写痕迹")
        if float(breakdown.get("english_abstract_signal") or 0.0) >= 0.08:
            tags.append("英文摘要镜像痕迹")
        return tags[:3]

    def _build_detect_paragraph_details(self, text: str, platform: str, profile: dict) -> list[dict]:
        paragraphs = self._split_detect_paragraphs(text)

        heading_map = {
            index: (self._detect_outline_heading(paragraph)[1] or "")
            for index, paragraph in enumerate(paragraphs, start=1)
        }
        rows: list[dict] = []
        for index, paragraph in enumerate(paragraphs, start=1):
            base_score = self._heuristic_ai_score(paragraph)
            score_ratio, _profile, breakdown = self._simulate_platform_detect_score(platform, paragraph, base_score)

            local_segments = self._local_suspicious_segments_v2(paragraph, platform, profile)
            merged_segments = self._merge_suspicious_segments(local_segments, [])
            if merged_segments:
                score_ratio = round(self._clamp_score(score_ratio + min(0.04, len(merged_segments) * 0.01)), 4)

            score_label = self._score_to_detect_label(score_ratio, profile)
            if (
                platform == "cnki"
                and score_label == "clean"
                and score_ratio >= 0.15
                and self._detect_outline_heading(paragraph)[0] is None
            ):
                score_label = "low"
            label = score_label
            segment_rows = []
            for segment in merged_segments:
                segment_rows.append(
                    {
                        **segment,
                        "label": "low",
                    }
                )
            rows.append(
                {
                    "index": index,
                    "label": label,
                    "risk_band": self._risk_band(score_ratio, high=profile["high"], medium=profile["medium"]),
                    "score": round(score_ratio * 100, 2),
                    "char_count": count_billable_chars(paragraph),
                    "sentence_count": len(self._split_detect_sentences(paragraph)),
                    "excerpt": self._clip_text(paragraph, 110),
                    "reason_tags": [] if label == "clean" else self._paragraph_reason_tags_v2(breakdown),
                    "suspicious_segments": [] if label == "clean" else segment_rows,
                }
            )

        if platform == "cnki" and len(rows) <= 60:
            excluded_heading_pattern = r"摘要|关键词|结论|结语|参考文献|附录|特色|推广|价值"
            for position, row in enumerate(rows):
                if row["label"] != "clean":
                    continue
                heading_text = re.sub(r"\s+", "", heading_map.get(row["index"], ""))
                if not heading_text or re.search(excluded_heading_pattern, heading_text):
                    continue
                following = rows[position + 1 : position + 3]
                if len(following) < 2:
                    continue
                avg_follow_score = sum(float(item["score"]) for item in following) / len(following)
                following_risky = sum(1 for item in following if item["label"] != "clean")
                if following_risky >= 1 and avg_follow_score >= 18.0:
                    row["label"] = "low"
                    row["suspicious_segments"] = [
                        {
                            "text": self._clip_text(heading_text, 76),
                            "score": round(avg_follow_score, 2),
                            "reason": "风险区标题延续",
                            "label": "low",
                        }
                    ]
        return rows

    def _detect_outline_heading(self, text: str) -> tuple[int | None, str]:
        clean = " ".join(str(text or "").split())
        normalized = re.sub(r"\s+", "", clean).strip("：:.。;；")
        if not normalized or len(normalized) > 40:
            return None, ""

        special_labels = ("摘要", "关键词", "引言", "绪论", "前言", "结语", "结论", "参考文献", "附录", "致谢")
        if normalized in special_labels:
            return 1, clean
        if len(normalized) <= 12:
            for label in special_labels:
                if normalized.startswith(label):
                    return 1, clean

        numeric_match = re.match(r"^(\d+(?:\.\d+){0,3})[、.．]?", normalized)
        if numeric_match:
            level = min(4, numeric_match.group(1).count(".") + 1)
            return level, clean

        patterns = [
            (r"^第[一二三四五六七八九十百零0-9]+[章节部分篇]", 1),
            (r"^[一二三四五六七八九十百零]+[、.．]", 1),
            (r"^[（(][一二三四五六七八九十百零0-9]+[)）]", 2),
        ]
        for pattern, level in patterns:
            if re.match(pattern, normalized):
                return level, clean
        return None, ""

    def _is_no_ai_fragment(self, text: str) -> bool:
        clean = " ".join(str(text or "").split())
        compact = re.sub(r"\s+", "", clean)
        if not clean:
            return True
        heading_level, _heading = self._detect_outline_heading(clean)
        if heading_level is not None:
            return True
        if re.search(
            r"分类号|密级|学号|独创性说明|知识产权声明|学位论文作者签名|指导教师签名|保密论文待解密后适用本声明|论文题目[:：]|学科名称[:：]",
            compact,
        ):
            return True
        if re.search(r"本人郑重声明|本人完全了解学校有关保护知识产权的规定", compact):
            return True
        if self._citation_relief_signal(clean) >= 0.08 and len(clean) <= 90:
            return True
        if len(clean) <= 10 and not re.search(r"[\u4e00-\u9fffA-Za-z]{4,}", clean):
            return True
        return False

    def _build_document_outline(self, text: str) -> list[dict]:
        paragraphs = self._split_detect_paragraphs(text)
        outline: list[dict] = []
        seen: set[str] = set()
        for index, paragraph in enumerate(paragraphs, start=1):
            level, heading = self._detect_outline_heading(paragraph)
            if level is None:
                continue
            key = re.sub(r"\s+", "", heading)
            if key in seen:
                continue
            seen.add(key)
            outline.append(
                {
                    "section": self._clip_text(heading, 32),
                    "level": level,
                    "start_index": index,
                }
            )

        for pos, item in enumerate(outline):
            next_start = outline[pos + 1]["start_index"] if pos + 1 < len(outline) else len(paragraphs) + 1
            item["end_index"] = max(int(item["start_index"]), next_start - 1)
            item["paragraph_count"] = item["end_index"] - int(item["start_index"]) + 1
        return outline[:24]

    def _build_section_distribution(
        self,
        paragraph_details: list[dict],
        document_outline: list[dict] | None = None,
    ) -> list[dict]:
        if not paragraph_details:
            return []
        if document_outline:
            outline_rows: list[dict] = []
            for item in document_outline:
                start_index = int(item.get("start_index") or 0)
                end_index = int(item.get("end_index") or 0)
                rows = [
                    detail
                    for detail in paragraph_details
                    if start_index <= int(detail.get("index") or 0) <= end_index
                ]
                if not rows:
                    continue
                avg_score = round(sum(float(detail["score"]) for detail in rows) / len(rows), 2)
                high_count = sum(1 for detail in rows if detail.get("label") == "high")
                outline_rows.append(
                    {
                        "section": item.get("section") or "-",
                        "level": int(item.get("level") or 1),
                        "paragraph_count": len(rows),
                        "avg_score": avg_score,
                        "high_count": high_count,
                        "start_index": start_index,
                        "end_index": end_index,
                    }
                )
            if outline_rows:
                return outline_rows

        total = len(paragraph_details)
        buckets = [
            ("开篇", range(1, max(2, total // 3 + 1))),
            ("中段", range(max(2, total // 3 + 1), max(3, total * 2 // 3 + 1))),
            ("结尾", range(max(3, total * 2 // 3 + 1), total + 1)),
        ]
        results: list[dict] = []
        for label, index_range in buckets:
            rows = [item for item in paragraph_details if item["index"] in index_range]
            if not rows:
                continue
            avg_score = round(sum(float(item["score"]) for item in rows) / len(rows), 2)
            high_count = sum(1 for item in rows if item.get("label") == "high")
            results.append(
                {
                    "section": label,
                    "paragraph_count": len(rows),
                    "avg_score": avg_score,
                    "high_count": high_count,
                }
            )
        return results

    def _cnki_short_doc_focus_indexes(self, paragraphs: list[str]) -> set[int]:
        focus_indexes: set[int] = set()
        if len(paragraphs) > 60:
            return focus_indexes
        for index, paragraph in enumerate(paragraphs, start=1):
            _level, heading = self._detect_outline_heading(paragraph)
            compact = re.sub(r"\s+", "", heading or "")
            if not compact or not re.search(r"课程融合|家长参与|成效与反思", compact):
                continue
            focus_indexes.add(index)
            for offset in (1, 2):
                target = index + offset
                if target <= len(paragraphs):
                    focus_indexes.add(target)
        return focus_indexes

    def _build_fragment_distribution(
        self,
        text: str,
        platform: str,
        profile: dict,
        paragraph_details: list[dict],
        document_outline: list[dict] | None = None,
    ) -> dict:
        paragraphs = self._split_detect_paragraphs(text)
        paragraph_score_map = {
            int(item.get("index") or 0): round(float(item.get("score") or 0.0) / 100.0, 4) for item in paragraph_details
        }
        paragraph_label_map = {int(item.get("index") or 0): str(item.get("label") or "") for item in paragraph_details}
        cnki_focus_indexes = (
            self._cnki_short_doc_focus_indexes(paragraphs) if (platform or "").strip().lower() == "cnki" else set()
        )

        count_map = {"high": 0, "medium": 0, "low": 0, "no_ai": 0}
        char_map = {"high": 0, "medium": 0, "low": 0, "no_ai": 0}
        display_count_map = {"mild": 0, "moderate": 0, "severe": 0}
        display_char_map = {"mild": 0, "moderate": 0, "severe": 0}
        weighted_scores: list[float] = []

        for paragraph_index, paragraph in enumerate(paragraphs, start=1):
            paragraph_ratio = paragraph_score_map.get(paragraph_index, 0.0)
            for sentence in self._split_detect_sentences(paragraph):
                compact = " ".join(sentence.split())
                if not compact:
                    continue
                char_count = count_billable_chars(compact)
                if char_count <= 0:
                    continue

                focus_heading = (
                    platform == "cnki"
                    and paragraph_index in cnki_focus_indexes
                    and self._detect_outline_heading(compact)[0] is not None
                )
                if self._is_no_ai_fragment(compact) and not focus_heading:
                    label = "no_ai"
                    score_ratio = 0.0
                else:
                    base_score = self._heuristic_ai_score(compact)
                    score_ratio, _profile, _breakdown = self._simulate_platform_detect_score(platform, compact, base_score)
                    if paragraph_ratio > 0:
                        score_ratio = round(self._clamp_score(score_ratio * 0.72 + paragraph_ratio * 0.28), 4)
                    label = self._score_to_detect_label(score_ratio, profile)
                    if platform == "cnki" and paragraph_index in cnki_focus_indexes:
                        paragraph_label = paragraph_label_map.get(paragraph_index, "")
                        if paragraph_label != "clean" and paragraph_ratio >= 0.30:
                            label = "high"
                        elif paragraph_label != "clean" and paragraph_ratio >= 0.22 and label == "low":
                            label = "medium"
                    elif platform == "cnki" and cnki_focus_indexes and paragraph_index != 1:
                        wrapup_relief = float(_breakdown.get("summary_wrapup_relief") or 0.0)
                        if wrapup_relief >= 0.10:
                            label = "no_ai"
                        elif label in ("low", "medium") and paragraph_ratio < 0.30:
                            label = "no_ai"
                    if label == "clean":
                        label = "no_ai"
                    else:
                        weighted_scores.append(score_ratio)
                        display_band = self._fragment_display_band(score_ratio, profile)
                        if display_band:
                            display_count_map[display_band] += 1
                            display_char_map[display_band] += char_count

                count_map[label] += 1
                char_map[label] += char_count

        total_fragments = sum(count_map.values())
        total_chars = sum(char_map.values())
        if total_fragments <= 0 or total_chars <= 0:
            return {
                "fragment_count": 0,
                "high_fragment_count": 0,
                "middle_fragment_count": 0,
                "low_fragment_count": 0,
                "no_ai_fragment_count": 0,
                "high_suspected_fragment_ratio": 0.0,
                "middle_suspected_fragment_ratio": 0.0,
                "low_suspected_fragment_ratio": 0.0,
                "no_ai_fragment_ratio": 0.0,
                "high_and_middle_suspected_fragment_ratio": 0.0,
                "total_suspected_fragment_ratio": 0.0,
                "high_suspected_text_ratio": 0.0,
                "middle_suspected_text_ratio": 0.0,
                "low_suspected_text_ratio": 0.0,
                "no_ai_suspected_text_ratio": 0.0,
                "high_and_middle_suspected_text_ratio": 0.0,
                "total_suspected_text_ratio": 0.0,
                "weighted_score_pct": 0.0,
                "mild_fragment_count": 0,
                "moderate_fragment_count": 0,
                "severe_fragment_count": 0,
                "mild_fragment_ratio": 0.0,
                "moderate_fragment_ratio": 0.0,
                "severe_fragment_ratio": 0.0,
                "mild_text_ratio": 0.0,
                "moderate_text_ratio": 0.0,
                "severe_text_ratio": 0.0,
                "display_thresholds": {"mild": 70.0, "moderate": 80.0, "severe": 90.0},
            }

        def ratio(part: int, whole: int) -> float:
            return round(part / max(whole, 1) * 100, 2)

        high_middle_fragments = count_map["high"] + count_map["medium"]
        suspected_fragments = high_middle_fragments + count_map["low"]
        high_middle_chars = char_map["high"] + char_map["medium"]
        suspected_chars = high_middle_chars + char_map["low"]
        thresholds = profile.get("fragment_display_thresholds") or {}

        return {
            "fragment_count": total_fragments,
            "high_fragment_count": count_map["high"],
            "middle_fragment_count": count_map["medium"],
            "low_fragment_count": count_map["low"],
            "no_ai_fragment_count": count_map["no_ai"],
            "high_suspected_fragment_ratio": ratio(count_map["high"], total_fragments),
            "middle_suspected_fragment_ratio": ratio(count_map["medium"], total_fragments),
            "low_suspected_fragment_ratio": ratio(count_map["low"], total_fragments),
            "no_ai_fragment_ratio": ratio(count_map["no_ai"], total_fragments),
            "high_and_middle_suspected_fragment_ratio": ratio(high_middle_fragments, total_fragments),
            "total_suspected_fragment_ratio": ratio(suspected_fragments, total_fragments),
            "high_suspected_text_ratio": ratio(char_map["high"], total_chars),
            "middle_suspected_text_ratio": ratio(char_map["medium"], total_chars),
            "low_suspected_text_ratio": ratio(char_map["low"], total_chars),
            "no_ai_suspected_text_ratio": ratio(char_map["no_ai"], total_chars),
            "high_and_middle_suspected_text_ratio": ratio(high_middle_chars, total_chars),
            "total_suspected_text_ratio": ratio(suspected_chars, total_chars),
            "weighted_score_pct": round(sum(weighted_scores) / max(len(weighted_scores), 1) * 100, 2)
            if weighted_scores
            else 0.0,
            "mild_fragment_count": display_count_map["mild"],
            "moderate_fragment_count": display_count_map["moderate"],
            "severe_fragment_count": display_count_map["severe"],
            "mild_fragment_ratio": ratio(display_count_map["mild"], total_fragments),
            "moderate_fragment_ratio": ratio(display_count_map["moderate"], total_fragments),
            "severe_fragment_ratio": ratio(display_count_map["severe"], total_fragments),
            "mild_text_ratio": ratio(display_char_map["mild"], total_chars),
            "moderate_text_ratio": ratio(display_char_map["moderate"], total_chars),
            "severe_text_ratio": ratio(display_char_map["severe"], total_chars),
            "display_thresholds": {
                "mild": round(float(thresholds.get("mild", 0.7)) * 100, 2),
                "moderate": round(float(thresholds.get("moderate", 0.8)) * 100, 2),
                "severe": round(float(thresholds.get("severe", 0.9)) * 100, 2),
            },
        }

    def _section_category(self, section_name: str) -> str:
        normalized = re.sub(r"\s+", "", str(section_name or "")).lower()
        if not normalized:
            return ""
        if any(token in normalized for token in ("摘要", "中英文摘要", "英文摘要", "abstract")):
            return "abstract"
        if any(token in normalized for token in ("绪论", "引言", "前言")):
            return "intro"
        if any(token in normalized for token in ("相关概念", "理论基础", "文献综述", "相关研究", "理论综述")):
            return "review"
        if any(token in normalized for token in ("结论", "结语", "总结")):
            return "conclusion"
        return ""

    def _build_document_metrics(
        self,
        *,
        text: str,
        paragraph_details: list[dict],
        section_distribution: list[dict],
        document_outline: list[dict],
        profile: dict,
    ) -> dict:
        total = len(paragraph_details)
        total_chars = sum(int(item.get("char_count") or 0) for item in paragraph_details)
        high_medium_count = sum(1 for item in paragraph_details if item.get("label") in {"high", "medium"})
        high_medium_chars = sum(
            int(item.get("char_count") or 0) for item in paragraph_details if item.get("label") in {"high", "medium"}
        )
        longest_streak = 0
        current_streak = 0
        for item in paragraph_details:
            if item.get("label") in {"high", "medium"}:
                current_streak += 1
                longest_streak = max(longest_streak, current_streak)
            else:
                current_streak = 0
        medium_threshold = float(profile.get("medium", 0.35)) * 100
        suspicious_sections = sum(
            1
            for item in section_distribution
            if float(item.get("avg_score") or 0.0) >= medium_threshold or int(item.get("high_count") or 0) > 0
        )
        abstract_scores = [
            float(item.get("avg_score") or 0.0)
            for item in section_distribution
            if self._section_category(item.get("section") or "") == "abstract"
        ]
        intro_scores = [
            float(item.get("avg_score") or 0.0)
            for item in section_distribution
            if self._section_category(item.get("section") or "") == "intro"
        ]
        review_scores = [
            float(item.get("avg_score") or 0.0)
            for item in section_distribution
            if self._section_category(item.get("section") or "") == "review"
        ]
        paragraphs = self._split_detect_paragraphs(text)
        opening_similarity = self._paragraph_opening_similarity(paragraphs)
        evidence_relief = max(self._citation_relief_signal_v2(text), self._evidence_relief_signal_v2(text))
        return {
            "paragraph_count": total,
            "high_medium_paragraph_count": high_medium_count,
            "high_medium_paragraph_ratio": round(high_medium_count / max(total, 1) * 100, 2) if total else 0.0,
            "high_medium_text_ratio": round(high_medium_chars / max(total_chars, 1) * 100, 2) if total_chars else 0.0,
            "longest_risk_streak": longest_streak,
            "longest_risk_streak_ratio": round(longest_streak / max(total, 1) * 100, 2) if total else 0.0,
            "opening_similarity_ratio": round(opening_similarity * 100, 2),
            "section_count": len(section_distribution),
            "suspicious_section_count": suspicious_sections,
            "section_coverage_ratio": round(suspicious_sections / max(len(section_distribution), 1) * 100, 2)
            if section_distribution
            else 0.0,
            "outline_sections": len(document_outline),
            "distribution_mode": "outline" if document_outline else "band",
            "evidence_relief_pct": round(evidence_relief * 100, 2),
            "abstract_avg_score": round(sum(abstract_scores) / max(len(abstract_scores), 1), 2) if abstract_scores else 0.0,
            "intro_avg_score": round(sum(intro_scores) / max(len(intro_scores), 1), 2) if intro_scores else 0.0,
            "review_avg_score": round(sum(review_scores) / max(len(review_scores), 1), 2) if review_scores else 0.0,
        }

    def _build_decision_basis(
        self,
        *,
        breakdown: dict,
        document_metrics: dict,
        fragment_distribution: dict,
        suspicious_segments: list[dict],
    ) -> list[dict]:
        items: list[dict] = []
        template_signal = float(breakdown.get("template_signal") or 0.0)
        context_signal = float(breakdown.get("context_signal") or 0.0)
        repeat_signal = float(breakdown.get("repeat_signal") or 0.0)
        opening_signal = float(breakdown.get("opening_signal") or 0.0)
        template_hits = [str(item) for item in (breakdown.get("template_hits") or []) if str(item).strip()]

        if template_signal >= 0.2:
            hit_text = "、".join(template_hits[:3]) if template_hits else "高频模板连接词"
            items.append(
                {
                    "title": "模板连接词密度偏高",
                    "detail": f"命中 {hit_text}，模板信号 {round(template_signal * 100, 2)}%。",
                    "direction": "risk",
                }
            )
        if context_signal >= 0.55:
            items.append(
                {
                    "title": "句长波动偏小",
                    "detail": f"上下文均匀度 {round(context_signal * 100, 2)}%，更接近机器批量生成的平直节奏。",
                    "direction": "risk",
                }
            )
        if repeat_signal >= 0.3:
            items.append(
                {
                    "title": "重复表达偏多",
                    "detail": f"重复信号 {round(repeat_signal * 100, 2)}%，存在近义重复或固定总结句反复出现。",
                    "direction": "risk",
                }
            )
        if opening_signal >= 0.44 or float(document_metrics.get("opening_similarity_ratio") or 0.0) >= 25:
            items.append(
                {
                    "title": "段首表达重复度较高",
                    "detail": f"段首相似度 {document_metrics.get('opening_similarity_ratio', 0.0)}%，段落展开方式较集中。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("section_coverage_ratio") or 0.0) >= 35:
            items.append(
                {
                    "title": "中高风险内容跨章节分布",
                    "detail": f"可疑章节覆盖率 {document_metrics.get('section_coverage_ratio', 0.0)}%，不是单点异常。",
                    "direction": "risk",
                }
            )
        if int(document_metrics.get("longest_risk_streak") or 0) >= 3:
            items.append(
                {
                    "title": "存在连续风险片段带",
                    "detail": f"最长连续中高风险段落 {document_metrics.get('longest_risk_streak', 0)} 段。",
                    "direction": "risk",
                }
            )
        if float(fragment_distribution.get("high_and_middle_suspected_text_ratio") or 0.0) >= 20:
            items.append(
                {
                    "title": "高中风险文字占比较高",
                    "detail": f"高中风险文字占比 {fragment_distribution.get('high_and_middle_suspected_text_ratio', 0.0)}%。",
                    "direction": "risk",
                }
            )
        if float(fragment_distribution.get("severe_fragment_ratio") or 0.0) >= 5:
            items.append(
                {
                    "title": "存在重度疑似片段",
                    "detail": f"90%以上片段占比 {fragment_distribution.get('severe_fragment_ratio', 0.0)}%。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("evidence_relief_pct") or 0.0) >= 6:
            items.append(
                {
                    "title": "引文与样本证据较多",
                    "detail": f"实证减权 {document_metrics.get('evidence_relief_pct', 0.0)}%，用于抑制纯模板误判。",
                    "direction": "relief",
                }
            )
        if not items:
            items.append(
                {
                    "title": "以全文统计与片段聚合综合判定",
                    "detail": f"当前识别到 {len(suspicious_segments)} 个可疑片段，建议人工复核摘要、引言与结论段。",
                    "direction": "risk",
                }
            )
        return items[:6]

    def _collect_suspicious_segments(self, paragraph_details: list[dict]) -> list[dict]:
        items: list[dict] = []
        seen: set[str] = set()
        for paragraph in paragraph_details:
            for segment in paragraph.get("suspicious_segments") or []:
                text = str(segment.get("text") or "").strip()
                if not text or text in seen:
                    continue
                seen.add(text)
                items.append(
                    {
                        "paragraph_index": paragraph.get("index"),
                        "score": round(float(segment.get("score") or 0.0), 2),
                        "text": text,
                        "reason": str(segment.get("reason") or "").strip(),
                    }
                )
        items.sort(key=lambda item: item["score"], reverse=True)
        return items[:12]

    def _build_detect_distribution(self, paragraph_details: list[dict]) -> dict:
        total = len(paragraph_details)
        if total == 0:
            return {
                "paragraph_count": 0,
                "high_count": 0,
                "medium_count": 0,
                "low_count": 0,
                "clean_count": 0,
                "high_ratio": 0.0,
                "medium_ratio": 0.0,
                "low_ratio": 0.0,
                "clean_ratio": 0.0,
                "avg_score": 0.0,
                "max_score": 0.0,
            }
        high_count = sum(1 for item in paragraph_details if item.get("label") == "high")
        medium_count = sum(1 for item in paragraph_details if item.get("label") == "medium")
        low_count = sum(1 for item in paragraph_details if item.get("label") == "low")
        clean_count = sum(1 for item in paragraph_details if item.get("label") == "clean")
        scores = [float(item.get("score") or 0.0) for item in paragraph_details]
        return {
            "paragraph_count": total,
            "high_count": high_count,
            "medium_count": medium_count,
            "low_count": low_count,
            "clean_count": clean_count,
            "high_ratio": round(high_count / total * 100, 2),
            "medium_ratio": round(medium_count / total * 100, 2),
            "low_ratio": round(low_count / total * 100, 2),
            "clean_ratio": round(clean_count / total * 100, 2),
            "avg_score": round(sum(scores) / total, 2),
            "max_score": round(max(scores), 2),
        }

    def _build_decision_basis_v2(
        self,
        *,
        breakdown: dict,
        document_metrics: dict,
        fragment_distribution: dict,
        suspicious_segments: list[dict],
    ) -> list[dict]:
        items: list[dict] = []
        template_signal = float(breakdown.get("template_signal") or 0.0)
        context_signal = float(breakdown.get("context_signal") or 0.0)
        repeat_signal = float(breakdown.get("repeat_signal") or 0.0)
        opening_signal = float(breakdown.get("opening_signal") or 0.0)
        artifact_signal = float(breakdown.get("artifact_signal") or 0.0)
        english_signal = float(breakdown.get("english_abstract_signal") or 0.0)
        template_hits = [str(item) for item in (breakdown.get("template_hits") or []) if str(item).strip()]
        artifact_hits = [str(item) for item in (breakdown.get("artifact_hits") or []) if str(item).strip()]

        if template_signal >= 0.2:
            hit_text = "、".join(template_hits[:3]) if template_hits else "高频模板连接词"
            items.append(
                {
                    "title": "模板连接词密度偏高",
                    "detail": f"命中 {hit_text}，模板信号 {round(template_signal * 100, 2)}%。",
                    "direction": "risk",
                }
            )
        if context_signal >= 0.55:
            items.append(
                {
                    "title": "句式波动偏小",
                    "detail": f"上下文均匀度 {round(context_signal * 100, 2)}%，更接近批量生成文本的平直节奏。",
                    "direction": "risk",
                }
            )
        if repeat_signal >= 0.3:
            items.append(
                {
                    "title": "重复表达偏多",
                    "detail": f"重复信号 {round(repeat_signal * 100, 2)}%，存在固定总结句反复出现的倾向。",
                    "direction": "risk",
                }
            )
        if opening_signal >= 0.44 or float(document_metrics.get("opening_similarity_ratio") or 0.0) >= 25:
            items.append(
                {
                    "title": "段首展开方式较集中",
                    "detail": f"段首相似度 {document_metrics.get('opening_similarity_ratio', 0.0)}%，段落起手方式较为集中。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("abstract_avg_score") or 0.0) >= 30:
            items.append(
                {
                    "title": "摘要区域模板化明显",
                    "detail": f"摘要相关章节平均风险 {document_metrics.get('abstract_avg_score', 0.0)}%，与样本中的高风险摘要特征一致。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("intro_avg_score") or 0.0) >= 12:
            items.append(
                {
                    "title": "绪论区域风险偏高",
                    "detail": f"绪论相关章节平均风险 {document_metrics.get('intro_avg_score', 0.0)}%，常见于研究对象、目的、路径的模板化展开。",
                    "direction": "risk",
                }
            )
        if artifact_signal >= 0.08 or artifact_hits:
            hit_text = "、".join(artifact_hits[:3]) if artifact_hits else "异常改写痕迹"
            items.append(
                {
                    "title": "存在异常改写痕迹",
                    "detail": f"命中 {hit_text}，疑似存在机械替换或不自然表达。",
                    "direction": "risk",
                }
            )
        if english_signal >= 0.08:
            items.append(
                {
                    "title": "英文摘要镜像痕迹偏强",
                    "detail": f"英文摘要信号 {round(english_signal * 100, 2)}%，中英摘要结构与措辞较为镜像。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("section_coverage_ratio") or 0.0) >= 35:
            items.append(
                {
                    "title": "中高风险内容跨章节分布",
                    "detail": f"可疑章节覆盖率 {document_metrics.get('section_coverage_ratio', 0.0)}%，不是单点异常。",
                    "direction": "risk",
                }
            )
        if int(document_metrics.get("longest_risk_streak") or 0) >= 3:
            items.append(
                {
                    "title": "存在连续风险片段带",
                    "detail": f"最长连续中高风险段落 {document_metrics.get('longest_risk_streak', 0)} 段。",
                    "direction": "risk",
                }
            )
        if float(fragment_distribution.get("high_and_middle_suspected_text_ratio") or 0.0) >= 20:
            items.append(
                {
                    "title": "高中风险文字占比较高",
                    "detail": f"高中风险文字占比 {fragment_distribution.get('high_and_middle_suspected_text_ratio', 0.0)}%。",
                    "direction": "risk",
                }
            )
        if float(document_metrics.get("evidence_relief_pct") or 0.0) >= 6:
            items.append(
                {
                    "title": "实证证据较充足",
                    "detail": f"实证减权 {document_metrics.get('evidence_relief_pct', 0.0)}%，用于抑制纯模板误判。",
                    "direction": "relief",
                }
            )
        if not items:
            items.append(
                {
                    "title": "以全文统计与片段聚合综合判定",
                    "detail": f"当前识别到 {len(suspicious_segments)} 个可疑片段，建议人工复核摘要、引言与结论段。",
                    "direction": "risk",
                }
            )
        return items[:6]

    def _build_detect_result(
        self,
        *,
        text: str,
        platform: str,
        mode: str,
        report_summary: dict,
        llm_result: dict | None,
    ) -> dict:
        base_score = self._heuristic_ai_score(text)
        score, profile, breakdown = self._simulate_platform_detect_score(platform, text, base_score)
        paragraph_details = self._build_detect_paragraph_details(text, platform, profile)
        distribution = self._build_detect_distribution(paragraph_details)
        document_outline = self._build_document_outline(text)
        section_distribution = self._build_section_distribution(paragraph_details, document_outline=document_outline)
        suspicious_segments = self._collect_suspicious_segments(paragraph_details)
        source_stats = self._text_stats(text)
        fragment_distribution = self._build_fragment_distribution(
            text,
            platform,
            profile,
            paragraph_details,
            document_outline=document_outline,
        )
        document_metrics = self._build_document_metrics(
            text=text,
            paragraph_details=paragraph_details,
            section_distribution=section_distribution,
            document_outline=document_outline,
            profile=profile,
        )

        mean_paragraph_ratio = round(float(distribution.get("avg_score") or 0.0) / 100.0, 4)
        max_paragraph_ratio = round(float(distribution.get("max_score") or 0.0) / 100.0, 4)
        segment_ratio = round(
            min(
                1.0,
                (
                    sum(float(item.get("score") or 0.0) for item in suspicious_segments[:5])
                    / max(len(suspicious_segments[:5]), 1)
                )
                / 100.0,
            ),
            4,
        )
        fragment_weighted_ratio = round(float(fragment_distribution.get("weighted_score_pct") or 0.0) / 100.0, 4)
        high_middle_text_ratio = round(
            float(fragment_distribution.get("high_and_middle_suspected_text_ratio") or 0.0) / 100.0,
            4,
        )
        coverage_ratio = round(float(document_metrics.get("high_medium_paragraph_ratio") or 0.0) / 100.0, 4)
        section_coverage_ratio = round(float(document_metrics.get("section_coverage_ratio") or 0.0) / 100.0, 4)
        streak_ratio = round(float(document_metrics.get("longest_risk_streak_ratio") or 0.0) / 100.0, 4)
        opening_similarity_ratio = round(float(document_metrics.get("opening_similarity_ratio") or 0.0) / 100.0, 4)
        evidence_relief_ratio = round(float(document_metrics.get("evidence_relief_pct") or 0.0) / 100.0, 4)
        abstract_section_ratio = round(float(document_metrics.get("abstract_avg_score") or 0.0) / 100.0, 4)
        intro_section_ratio = round(float(document_metrics.get("intro_avg_score") or 0.0) / 100.0, 4)
        score = round(
            self._clamp_score(
                score * 0.32
                + mean_paragraph_ratio * 0.22
                + max_paragraph_ratio * 0.14
                + segment_ratio * 0.10
                + fragment_weighted_ratio * 0.12
                + high_middle_text_ratio * 0.10
                + coverage_ratio * float(profile.get("coverage_weight", 0.0))
                + section_coverage_ratio * float(profile.get("section_weight", 0.0))
                + streak_ratio * float(profile.get("streak_weight", 0.0))
                + opening_similarity_ratio * float(profile.get("opening_similarity_weight", 0.0))
                + abstract_section_ratio * float(profile.get("abstract_section_weight", 0.0))
                + intro_section_ratio * float(profile.get("intro_section_weight", 0.0))
                - evidence_relief_ratio * float(profile.get("evidence_relief_weight", 0.0))
            ),
            4,
        )
        llm_label = ""

        if isinstance(llm_result, dict):
            llm_score = self._coerce_ratio(llm_result.get("ai_score"))
            if llm_score is not None:
                score = round(self._clamp_score(score * 0.82 + llm_score * 0.18), 4)
                breakdown["llm_score"] = round(llm_score, 4)
                breakdown["llm_blended"] = True
            else:
                breakdown["llm_blended"] = False

            llm_label = self._normalize_detect_label(llm_result.get("label"))

            reason = llm_result.get("reason")
            if isinstance(reason, str) and reason.strip():
                breakdown["llm_reason"] = reason.strip()[:180]
        else:
            breakdown["llm_blended"] = False

        raw_score = score
        score = self._calibrate_detect_score(
            platform=platform,
            raw_score=raw_score,
            breakdown=breakdown,
            distribution=distribution,
            fragment_distribution=fragment_distribution,
            document_metrics=document_metrics,
            source_stats=source_stats,
        )
        score_pct = round(score * 100, 2)

        if llm_label:
            label = llm_label
        else:
            label = self._score_to_detect_label(score, profile)

        decision_basis = self._build_decision_basis_v2(
            breakdown=breakdown,
            document_metrics=document_metrics,
            fragment_distribution=fragment_distribution,
            suspicious_segments=suspicious_segments,
        )

        band = self._risk_band(score, high=profile["high"], medium=profile["medium"])
        detail_expanded = self._should_expand_detect_details(platform, score_pct)
        display_segments = suspicious_segments if detail_expanded else []
        risk_paragraphs = sorted(paragraph_details, key=lambda item: item["score"], reverse=True)[:5] if detail_expanded else []
        report_no = f"GW-AIGC-{platform.upper()}-{datetime.now().strftime('%Y%m%d%H%M%S')}-{self._current_task_id or 0}"
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        breakdown = enrich_detect_breakdown(
            breakdown,
            mean_paragraph_ratio=mean_paragraph_ratio,
            max_paragraph_ratio=max_paragraph_ratio,
            segment_ratio=segment_ratio,
            fragment_weighted_ratio=fragment_weighted_ratio,
            high_middle_text_ratio=high_middle_text_ratio,
            outline_sections=len(document_outline),
            coverage_ratio=coverage_ratio,
            section_coverage_ratio=section_coverage_ratio,
            streak_ratio=streak_ratio,
            opening_similarity_ratio=opening_similarity_ratio,
            evidence_relief_ratio=evidence_relief_ratio,
            abstract_section_ratio=abstract_section_ratio,
            intro_section_ratio=intro_section_ratio,
            raw_score=raw_score,
            score_pct=score_pct,
        )

        basis_titles = [str(item.get("title") or "").strip() for item in decision_basis if item.get("direction") == "risk"]
        summary = build_detect_summary(
            platform=platform,
            profile=profile,
            score_pct=score_pct,
            source_stats=source_stats,
            detail_expanded=detail_expanded,
            basis_titles=basis_titles,
            fragment_distribution=fragment_distribution,
            distribution=distribution,
        )

        result = build_detect_result_payload(
            platform=platform,
            profile=profile,
            report_no=report_no,
            generated_at=generated_at,
            mode=mode,
            pipeline_usage=self._pipeline_usage,
            score=score,
            score_pct=score_pct,
            label=label,
            band=band,
            summary=summary,
            source_stats=source_stats,
            report_summary=report_summary,
            breakdown=breakdown,
            distribution=distribution,
            fragment_distribution=fragment_distribution,
            document_metrics=document_metrics,
            decision_basis=decision_basis,
            document_outline=document_outline,
            section_distribution=section_distribution,
            detail_expanded=detail_expanded,
            risk_paragraphs=risk_paragraphs,
            paragraph_details=paragraph_details,
            suspicious_segments=display_segments,
        )
        result["report_view"] = self._build_detect_report_view(result)
        return result

    def _risk_band(self, score: float, *, high: float = 0.65, medium: float = 0.35) -> str:
        if score >= high:
            return "高风险"
        if score >= medium:
            return "中风险"
        return "低风险"

    def _legacy_top_risk_paragraphs_v0(self, text: str, platform: str = "cnki") -> list[dict]:
        profile = self._platform_detect_profile(platform)
        rows = self._build_detect_paragraph_details(text, platform, profile)
        return sorted(rows, key=lambda item: item["score"], reverse=True)[:5]

    def _wrap_pdf_line(self, text: str, width: int = 56) -> list[str]:
        compact = " ".join(str(text or "").split())
        if not compact:
            return [""]

        lines: list[str] = []
        current = ""
        current_width = 0
        for char in compact:
            char_width = 1 if ord(char) < 128 else 2
            if current and current_width + char_width > width:
                lines.append(current.rstrip())
                current = char
                current_width = char_width
                continue
            current += char
            current_width += char_width
        if current:
            lines.append(current.rstrip())
        return lines or [""]

    def _pdf_text_hex(self, text: str) -> str:
        value = str(text or " ")
        return value.encode("utf-16-be").hex().upper()

    def _render_pdf(self, lines: list[str]) -> bytes:
        page_width = 595
        page_height = 842
        margin_x = 44
        margin_top = 56
        margin_bottom = 48
        font_size = 11
        line_height = 16

        expanded_lines: list[str] = []
        for line in lines:
            wrapped = self._wrap_pdf_line(line, width=56)
            expanded_lines.extend(wrapped if wrapped else [""])
        if not expanded_lines:
            expanded_lines = [""]

        usable_height = page_height - margin_top - margin_bottom
        lines_per_page = max(1, int(usable_height // line_height))
        page_line_chunks = [
            expanded_lines[i : i + lines_per_page] for i in range(0, len(expanded_lines), lines_per_page)
        ]

        objects: list[tuple[int, bytes]] = [(1, b"<< /Type /Catalog /Pages 2 0 R >>")]
        page_refs: list[str] = []
        next_obj_id = 3
        font_obj_id = 2 + len(page_line_chunks) * 2 + 1
        descendant_font_obj_id = font_obj_id + 1

        for chunk in page_line_chunks:
            page_obj_id = next_obj_id
            content_obj_id = next_obj_id + 1
            next_obj_id += 2
            page_refs.append(f"{page_obj_id} 0 R")

            text_ops: list[str] = []
            for index, line in enumerate(chunk):
                y = page_height - margin_top - index * line_height
                hex_text = self._pdf_text_hex(line)
                text_ops.append(f"BT /F1 {font_size} Tf 1 0 0 1 {margin_x} {y:.2f} Tm <{hex_text}> Tj ET")
            stream_text = "\n".join(text_ops)
            stream_bytes = stream_text.encode("ascii")

            page_obj = (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] "
                f"/Resources << /Font << /F1 {font_obj_id} 0 R >> >> /Contents {content_obj_id} 0 R >>"
            ).encode("ascii")
            content_obj = (
                f"<< /Length {len(stream_bytes)} >>\nstream\n".encode("ascii")
                + stream_bytes
                + b"\nendstream"
            )
            objects.append((page_obj_id, page_obj))
            objects.append((content_obj_id, content_obj))

        pages_obj = f"<< /Type /Pages /Count {len(page_line_chunks)} /Kids [{' '.join(page_refs)}] >>".encode("ascii")
        objects.insert(1, (2, pages_obj))
        objects.append(
            (
                font_obj_id,
                (
                    f"<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light "
                    f"/Encoding /UniGB-UCS2-H /DescendantFonts [{descendant_font_obj_id} 0 R] >>"
                ).encode("ascii"),
            )
        )
        objects.append(
            (
                descendant_font_obj_id,
                b"<< /Type /Font /Subtype /CIDFontType0 /BaseFont /STSong-Light "
                b"/CIDSystemInfo << /Registry (Adobe) /Ordering (GB1) /Supplement 4 >> /DW 1000 >>",
            )
        )

        objects.sort(key=lambda item: item[0])
        output = bytearray()
        output.extend(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets: dict[int, int] = {}

        for obj_id, obj_body in objects:
            offsets[obj_id] = len(output)
            output.extend(f"{obj_id} 0 obj\n".encode("ascii"))
            output.extend(obj_body)
            output.extend(b"\nendobj\n")

        xref_offset = len(output)
        max_obj_id = max(offsets)
        output.extend(f"xref\n0 {max_obj_id + 1}\n".encode("ascii"))
        output.extend(b"0000000000 65535 f \n")
        for obj_id in range(1, max_obj_id + 1):
            offset = offsets.get(obj_id, 0)
            output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        output.extend(f"trailer\n<< /Size {max_obj_id + 1} /Root 1 0 R >>\n".encode("ascii"))
        output.extend(f"startxref\n{xref_offset}\n%%EOF".encode("ascii"))
        return bytes(output)

    def _build_transform_result(
        self,
        *,
        task_type: TaskType,
        platform: str,
        mode: str,
        source_text: str,
        output_text: str,
        report_summary: dict,
    ) -> dict:
        return build_transform_result(
            task_type=task_type,
            platform=platform,
            mode=mode,
            source_text=source_text,
            output_text=output_text,
            report_summary=report_summary,
            text_stats=self._text_stats,
            clip_text=self._clip_text,
            pipeline_usage=self._pipeline_usage,
            rewrite_strategy_meta=self._rewrite_strategy_meta,
            dedup_strategy_meta=self._dedup_strategy_meta,
        )

    def _clip_text(self, text: str, limit: int) -> str:
        compact = " ".join((text or "").split())
        if len(compact) <= limit:
            return compact
        return f"{compact[:limit].rstrip()}..."

    def _current_task_report_meta(self) -> dict:
        from app.models import Task

        if not self._current_task_id:
            return {}
        task = self.db.get(Task, self._current_task_id)
        if task is None:
            return {}
        payload = task.result_json if isinstance(task.result_json, dict) else {}
        return {
            "paper_title": str(payload.get("paper_title") or "").strip(),
            "authors": str(payload.get("authors") or "").strip(),
            "source_filename": str(task.source_filename or "").strip(),
        }

    def _build_detect_band_rows(self, paragraph_details: list[dict]) -> list[dict]:
        total = len(paragraph_details)
        if total <= 0:
            return [
                {"label": "前部20%", "paragraph_count": 0, "avg_score": 0.0, "high_count": 0, "char_count": 0},
                {"label": "中部60%", "paragraph_count": 0, "avg_score": 0.0, "high_count": 0, "char_count": 0},
                {"label": "后部20%", "paragraph_count": 0, "avg_score": 0.0, "high_count": 0, "char_count": 0},
            ]

        front_end = max(1, math.ceil(total * 0.2))
        middle_end = max(front_end, math.ceil(total * 0.8))
        chunks = [
            ("前部20%", paragraph_details[:front_end]),
            ("中部60%", paragraph_details[front_end:middle_end]),
            ("后部20%", paragraph_details[middle_end:]),
        ]

        rows: list[dict] = []
        for label, items in chunks:
            paragraph_count = len(items)
            avg_score = round(sum(float(item.get("score") or 0.0) for item in items) / paragraph_count, 2) if items else 0.0
            high_count = sum(1 for item in items if str(item.get("label") or "").lower() == "high")
            char_count = sum(int(item.get("char_count") or 0) for item in items)
            rows.append(
                {
                    "label": label,
                    "paragraph_count": paragraph_count,
                    "avg_score": avg_score,
                    "high_count": high_count,
                    "char_count": char_count,
                }
            )
        return rows

    def _build_detect_report_view(self, result: dict) -> dict:
        platform = str(result.get("platform") or "").strip().lower()
        stats = result.get("source_stats") or {}
        fragment_distribution = result.get("fragment_distribution") or {}
        document_metrics = result.get("document_metrics") or {}
        paragraph_details = result.get("paragraph_details") or []
        band_rows = self._build_detect_band_rows(paragraph_details)
        score_pct = round(float(result.get("score_pct") or 0.0), 2)
        total_chars = int(stats.get("char_count") or 0)
        ai_chars = min(total_chars, int(round(total_chars * score_pct / 100.0)))
        high_chars = min(
            total_chars,
            int(round(total_chars * float(fragment_distribution.get("high_suspected_text_ratio") or 0.0) / 100.0)),
        )
        middle_chars = min(
            total_chars,
            int(round(total_chars * float(fragment_distribution.get("middle_suspected_text_ratio") or 0.0) / 100.0)),
        )
        detail_expanded = bool(result.get("detail_expanded"))
        display_segments = list(result.get("suspicious_segments") or [])
        display_paragraphs = paragraph_details[:24] if detail_expanded else []
        risk_paragraphs = list(result.get("risk_paragraphs") or [])

        if platform == "cnki":
            headline_metrics = [
                {"label": "AI特征值", "value": f"{score_pct}%"},
                {"label": "AI特征字符数", "value": str(ai_chars)},
                {"label": "总字符数", "value": str(total_chars)},
                {"label": "风险等级", "value": str(result.get("risk_band") or "-")},
            ]
            secondary_metrics = [
                {"label": "AI特征显著", "value": str(min(ai_chars, high_chars))},
                {"label": "AI特征疑似", "value": str(max(ai_chars - min(ai_chars, high_chars), 0))},
                {"label": "未标识部分", "value": str(max(total_chars - ai_chars, 0))},
            ]
            distribution_metrics = [
                {
                    "label": item["label"],
                    "value": f"{item['avg_score']}%",
                    "detail": f"段落 {item['paragraph_count']} | 高风险 {item['high_count']}",
                }
                for item in band_rows
            ]
            detail_hint = "当前未形成稳定可标注的高置信片段，全文报告仅保留基础判定信息。"
        elif platform == "vip":
            human_pct = round(max(0.0, 100.0 - score_pct), 2)
            headline_metrics = [
                {"label": "全文疑似AIGC生成", "value": f"{score_pct}%"},
                {"label": "全文人写概率", "value": f"{human_pct}%"},
                {"label": "AI生成文字", "value": str(ai_chars)},
                {"label": "风险等级", "value": str(result.get("risk_band") or "-")},
            ]
            secondary_metrics = [
                {"label": "章节覆盖率", "value": f"{document_metrics.get('section_coverage_ratio', 0)}%"},
                {"label": "最长连续风险段落", "value": str(document_metrics.get("longest_risk_streak", 0))},
                {"label": "总字符数", "value": str(total_chars)},
            ]
            distribution_metrics = [
                {"label": "高疑似文字占比", "value": f"{fragment_distribution.get('high_suspected_text_ratio', 0)}%"},
                {"label": "中疑似文字占比", "value": f"{fragment_distribution.get('middle_suspected_text_ratio', 0)}%"},
                {"label": "人工占比", "value": f"{human_pct}%"},
            ]
            detail_hint = "当前结果以低风险结论为主，未展开正文级疑似片段。"
        else:
            headline_metrics = [
                {"label": "AIGC总体疑似度", "value": f"{score_pct}%"},
                {
                    "label": "片段加权结果",
                    "value": f"{round(float(fragment_distribution.get('weighted_score_pct') or 0.0), 2)}%",
                },
                {"label": "高度疑似文字", "value": str(high_chars)},
                {"label": "中度疑似文字", "value": str(middle_chars)},
            ]
            secondary_metrics = [
                {"label": "总字符数", "value": str(total_chars)},
                {"label": "高风险片段数", "value": str(fragment_distribution.get("high_fragment_count", 0))},
                {"label": "中风险片段数", "value": str(fragment_distribution.get("middle_fragment_count", 0))},
            ]
            distribution_metrics = [
                {"label": "高度疑似", "value": f"{fragment_distribution.get('high_suspected_text_ratio', 0)}%"},
                {"label": "中度疑似", "value": f"{fragment_distribution.get('middle_suspected_text_ratio', 0)}%"},
                {"label": "轻度疑似", "value": f"{fragment_distribution.get('low_suspected_text_ratio', 0)}%"},
                {
                    "label": "人工占比",
                    "value": f"{round(max(0.0, 100.0 - float(fragment_distribution.get('total_suspected_text_ratio') or 0.0)), 2)}%",
                },
            ]
            detail_hint = "当前报告保留重点疑似片段与段落级判定明细。"

        return {
            "brand_title": "格物学术 AIGC 检测报告",
            "platform_badge": str(result.get("provider_label") or result.get("platform") or "-"),
            "compact_title": "简洁报告",
            "full_title": "全文报告",
            "headline_metrics": headline_metrics,
            "secondary_metrics": secondary_metrics,
            "distribution_metrics": distribution_metrics,
            "band_rows": band_rows,
            "detail_expanded": detail_expanded,
            "display_segments": display_segments,
            "display_paragraphs": display_paragraphs,
            "risk_paragraphs": risk_paragraphs,
            "detail_hint": detail_hint,
            "report_note": "本报告参考公开AIGC检测报告的栏目组织方式，由格物学术生成仿真结果，用于内部研判与人工复核。",
        }

    def _escape_pdf_text(self, text: str) -> str:
        return (
            str(text or "")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    def _write_detect_report_pdf(self, output_path: Path, result: dict) -> None:
        output_path.write_bytes(self._render_detect_report_pdf(result))

    def _render_detect_report_pdf(self, result: dict) -> bytes:
        try:
            from app.services.detect_report_renderer import render_detect_report_pdf_reportlab

            return render_detect_report_pdf_reportlab(
                result,
                meta=self._current_task_report_meta(),
                source_text=self._current_detect_source_text,
                build_detect_band_rows=self._build_detect_band_rows,
                split_detect_paragraphs=self._split_detect_paragraphs,
                escape_pdf_text=self._escape_pdf_text,
                detect_outline_heading=self._detect_outline_heading,
            )
        except Exception as exc:
            if isinstance(exc, ModuleNotFoundError):
                module_name = (getattr(exc, "name", "") or "").split(".", 1)[0]
                if module_name != "reportlab":
                    raise
            return self._render_detect_report_pdf_fallback(result)

    def _render_detect_report_pdf_reportlab(self, result: dict) -> bytes:
        from io import BytesIO

        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=16 * mm,
            rightMargin=16 * mm,
            topMargin=16 * mm,
            bottomMargin=14 * mm,
            title="格物学术 AIGC 全文检测报告",
            author="格物学术",
        )

        base_styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DetectTitle",
            parent=base_styles["Title"],
            fontName="STSong-Light",
            fontSize=20,
            leading=26,
            textColor=colors.HexColor("#111111"),
            alignment=1,
            spaceAfter=6,
        )
        subtitle_style = ParagraphStyle(
            "DetectSubtitle",
            parent=base_styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#666666"),
            alignment=1,
            spaceAfter=12,
        )
        section_style = ParagraphStyle(
            "DetectSection",
            parent=base_styles["Heading2"],
            fontName="STSong-Light",
            fontSize=13,
            leading=18,
            textColor=colors.HexColor("#111111"),
            spaceBefore=8,
            spaceAfter=8,
        )
        body_style = ParagraphStyle(
            "DetectBody",
            parent=base_styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#222222"),
        )
        small_style = ParagraphStyle(
            "DetectSmall",
            parent=body_style,
            fontSize=9,
            leading=13,
            textColor=colors.HexColor("#666666"),
        )
        emphasis_style = ParagraphStyle(
            "DetectEmphasis",
            parent=body_style,
            fontSize=11,
            leading=16,
            textColor=colors.HexColor("#111111"),
        )

        meta = self._current_task_report_meta()
        stats = result.get("source_stats") or {}
        distribution = result.get("distribution") or {}
        fragment_distribution = result.get("fragment_distribution") or {}
        document_metrics = result.get("document_metrics") or {}
        decision_basis = result.get("decision_basis") or []
        document_outline = result.get("document_outline") or []
        section_distribution = result.get("section_distribution") or []
        paragraph_details = result.get("paragraph_details") or []
        suspicious_segments = result.get("suspicious_segments") or []
        band_rows = self._build_detect_band_rows(paragraph_details)
        paper_title = meta.get("paper_title") or "未填写"
        authors = meta.get("authors") or "未填写"
        source_filename = meta.get("source_filename") or "未记录"

        def para(text: str, style: ParagraphStyle) -> Paragraph:
            return Paragraph(self._escape_pdf_text(text), style)

        def build_table(rows, col_widths=None, header=False, font_size=9.5):
            table = Table(rows, colWidths=col_widths, repeatRows=1 if header else 0)
            base_style = [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("LEADING", (0, 0), (-1, -1), font_size + 4),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#222222")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f5f1e8") if header else colors.white),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#d7d0c2")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
            if header:
                base_style.extend(
                    [
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111111")),
                        ("FONTNAME", (0, 0), (-1, 0), "STSong-Light"),
                    ]
                )
            table.setStyle(TableStyle(base_style))
            return table

        story: list = [
            para("AIGC检测 · 全文报告单", title_style),
            para("格物学术仿真版，结构参考知网公开报告形态，结果仅作内部研判与人工复核辅助。", subtitle_style),
        ]

        summary_rows = [
            [para("报告编号", body_style), para(str(result.get("report_no") or "-"), body_style)],
            [para("检测时间", body_style), para(str(result.get("generated_at") or "-"), body_style)],
            [para("检测平台", body_style), para(str(result.get("provider_label") or result.get("platform") or "-"), body_style)],
            [para("检测模式", body_style), para(str(result.get("mode") or "-"), body_style)],
            [para("篇名", body_style), para(paper_title, body_style)],
            [para("作者", body_style), para(authors, body_style)],
            [para("文件名", body_style), para(source_filename, body_style)],
        ]
        story.append(build_table(summary_rows, col_widths=[28 * mm, 150 * mm]))
        story.append(Spacer(1, 6))

        story.append(para("一、全文检测结果", section_style))
        headline_rows = [
            [para("AI特征值", small_style), para("风险等级", small_style), para("总字符数", small_style), para("可疑片段数", small_style)],
            [
                para(f"{result.get('score_pct', 0)}%", emphasis_style),
                para(str(result.get("risk_band") or "-"), emphasis_style),
                para(str(stats.get("char_count", 0)), emphasis_style),
                para(str(len(suspicious_segments)), emphasis_style),
            ],
        ]
        story.append(build_table(headline_rows, col_widths=[42 * mm, 42 * mm, 42 * mm, 42 * mm], header=True, font_size=10))
        story.append(Spacer(1, 4))
        story.append(para(str(result.get("summary") or ""), body_style))
        if decision_basis:
            story.append(Spacer(1, 6))
            story.append(para("判定依据摘要", small_style))
            basis_rows = [["方向", "依据", "说明"]]
            for item in decision_basis[:6]:
                basis_rows.append(
                    [
                        "减权" if str(item.get("direction") or "") == "relief" else "风险",
                        str(item.get("title") or "-"),
                        str(item.get("detail") or "-"),
                    ]
                )
            story.append(
                build_table(
                    basis_rows,
                    col_widths=[18 * mm, 46 * mm, 114 * mm],
                    header=True,
                    font_size=8.8,
                )
            )

        story.append(Spacer(1, 10))
        story.append(para("二、AIGC文字与片段分布", section_style))
        ratio_rows = [["等级", "文字占比", "片段数", "片段占比"]]
        ratio_rows.extend(
            [
                [
                    "高风险",
                    f"{fragment_distribution.get('high_suspected_text_ratio', 0)}%",
                    str(fragment_distribution.get("high_fragment_count", 0)),
                    f"{fragment_distribution.get('high_suspected_fragment_ratio', 0)}%",
                ],
                [
                    "中风险",
                    f"{fragment_distribution.get('middle_suspected_text_ratio', 0)}%",
                    str(fragment_distribution.get("middle_fragment_count", 0)),
                    f"{fragment_distribution.get('middle_suspected_fragment_ratio', 0)}%",
                ],
                [
                    "低风险",
                    f"{fragment_distribution.get('low_suspected_text_ratio', 0)}%",
                    str(fragment_distribution.get("low_fragment_count", 0)),
                    f"{fragment_distribution.get('low_suspected_fragment_ratio', 0)}%",
                ],
                [
                    "不计入片段",
                    f"{fragment_distribution.get('no_ai_suspected_text_ratio', 0)}%",
                    str(fragment_distribution.get("no_ai_fragment_count", 0)),
                    f"{fragment_distribution.get('no_ai_fragment_ratio', 0)}%",
                ],
            ]
        )
        story.append(build_table(ratio_rows, col_widths=[36 * mm, 44 * mm, 44 * mm, 44 * mm], header=True))
        story.append(Spacer(1, 6))
        story.append(
            para(
                f"高中风险文字占比 {fragment_distribution.get('high_and_middle_suspected_text_ratio', 0)}%，"
                f"总疑似文字占比 {fragment_distribution.get('total_suspected_text_ratio', 0)}%，"
                f"片段加权值 {fragment_distribution.get('weighted_score_pct', 0)}%。",
                small_style,
            )
        )
        story.append(Spacer(1, 8))
        story.append(para("按全文位置统计", small_style))
        band_table = [["区段", "段落数", "均值", "高风险段落", "字符数"]]
        for item in band_rows:
            band_table.append(
                [
                    item["label"],
                    str(item["paragraph_count"]),
                    f"{item['avg_score']}%",
                    str(item["high_count"]),
                    str(item["char_count"]),
                ]
            )
        story.append(build_table(band_table, col_widths=[32 * mm, 28 * mm, 28 * mm, 34 * mm, 34 * mm], header=True))

        story.append(Spacer(1, 10))
        story.append(para("三、章节目录与分布", section_style))
        if document_outline:
            outline_map = {str(item.get("section") or ""): item for item in section_distribution}
            outline_rows = [["章节", "起止段落", "层级", "平均风险"]]
            for item in document_outline[:18]:
                section_row = outline_map.get(str(item.get("section") or ""))
                outline_rows.append(
                    [
                        str(item.get("section") or "-"),
                        f"P{int(item.get('start_index') or 0):02d}-P{int(item.get('end_index') or 0):02d}",
                        str(item.get("level") or 1),
                        f"{(section_row or {}).get('avg_score', 0)}%",
                    ]
                )
            story.append(build_table(outline_rows, col_widths=[88 * mm, 34 * mm, 20 * mm, 34 * mm], header=True))
        elif section_distribution:
            section_rows = [["章节区段", "段落数", "均值", "高风险段落"]]
            for item in section_distribution:
                section_rows.append(
                    [
                        str(item.get("section") or "-"),
                        str(item.get("paragraph_count") or 0),
                        f"{item.get('avg_score', 0)}%",
                        str(item.get("high_count") or 0),
                    ]
                )
            story.append(build_table(section_rows, col_widths=[54 * mm, 30 * mm, 30 * mm, 44 * mm], header=True))
        else:
            story.append(para("未识别到可用的章节目录信息。", body_style))

        story.append(Spacer(1, 10))
        story.append(para("四、片段指标列表", section_style))
        indicator_rows = [["序号", "片段名称", "所在段落", "字符数", "风险值", "判定"]]
        for idx, item in enumerate(suspicious_segments[:12], start=1):
            indicator_rows.append(
                [
                    str(idx),
                    f"片段{idx}",
                    f"P{int(item.get('paragraph_index') or 0):02d}",
                    str(len(str(item.get("text") or ""))),
                    f"{item.get('score', 0)}%",
                    str(item.get("reason") or "综合风险偏高"),
                ]
            )
        if len(indicator_rows) == 1:
            indicator_rows.append(["-", "-", "-", "0", "0%", "未识别到高风险片段"])
        story.append(
            build_table(
                indicator_rows,
                col_widths=[14 * mm, 24 * mm, 22 * mm, 20 * mm, 20 * mm, 72 * mm],
                header=True,
                font_size=8.8,
            )
        )

        story.append(Spacer(1, 10))
        story.append(para("五、原文重点片段详情", section_style))
        if suspicious_segments:
            for idx, item in enumerate(suspicious_segments[:10], start=1):
                detail_rows = [
                    [para(f"NO.{idx} 片段{idx}", emphasis_style), para(f"段落：P{int(item.get('paragraph_index') or 0):02d}    风险值：{item.get('score', 0)}%", emphasis_style)],
                    [para("风险判定", body_style), para(str(item.get("reason") or "综合风险偏高"), body_style)],
                    [para("原文内容", body_style), para(self._clip_text(str(item.get("text") or ""), 220), body_style)],
                ]
                story.append(KeepTogether([build_table(detail_rows, col_widths=[28 * mm, 150 * mm]), Spacer(1, 6)]))
        else:
            story.append(para("本次未识别到高置信的可疑片段。", body_style))

        story.append(PageBreak())
        story.append(para("六、全文段落明细", section_style))
        if paragraph_details:
            for item in paragraph_details:
                reason_tags = "、".join(item.get("reason_tags") or []) or "综合语义判定"
                lines = [
                    para(
                        f"P{int(item.get('index') or 0):02d} | {item.get('risk_band') or '-'} | {item.get('score', 0)}% | 字符 {item.get('char_count', 0)} | 句子 {item.get('sentence_count', 0)}",
                        emphasis_style,
                    ),
                    para(f"信号：{reason_tags}", small_style),
                    para(f"摘要：{self._clip_text(str(item.get('excerpt') or ''), 260)}", body_style),
                ]
                for segment in (item.get("suspicious_segments") or [])[:2]:
                    lines.append(
                        para(
                            f"片段：{segment.get('score', 0)}% | {self._clip_text(str(segment.get('text') or ''), 180)}",
                            small_style,
                        )
                    )
                lines.append(Spacer(1, 6))
                story.append(KeepTogether(lines))
        else:
            story.append(para("未生成段落级明细。", body_style))

        story.append(Spacer(1, 8))
        story.append(para("七、辅助指标与说明", section_style))
        metric_rows = [["文档级指标", "值"]]
        metric_rows.extend(
            [
                ["中高风险段落占比", f"{document_metrics.get('high_medium_paragraph_ratio', 0)}%"],
                ["中高风险文字占比", f"{document_metrics.get('high_medium_text_ratio', 0)}%"],
                ["最长连续风险段落", str(document_metrics.get("longest_risk_streak", 0))],
                ["段首相似度", f"{document_metrics.get('opening_similarity_ratio', 0)}%"],
                ["可疑章节覆盖率", f"{document_metrics.get('section_coverage_ratio', 0)}%"],
                ["实证减权", f"{document_metrics.get('evidence_relief_pct', 0)}%"],
            ]
        )
        story.append(build_table(metric_rows, col_widths=[62 * mm, 116 * mm], header=True))
        story.append(Spacer(1, 6))
        aux_rows = [
            [para("句子数", body_style), para(str(stats.get("sentence_count", 0)), body_style)],
            [para("段落数", body_style), para(str(stats.get("paragraph_count", 0)), body_style)],
            [para("平均句长", body_style), para(str(stats.get("avg_sentence_length", 0)), body_style)],
            [para("高风险段落", body_style), para(f"{distribution.get('high_count', 0)}（{distribution.get('high_ratio', 0)}%）", body_style)],
            [
                para("高中风险文字占比", body_style),
                para(f"{fragment_distribution.get('high_and_middle_suspected_text_ratio', 0)}%", body_style),
            ],
            [
                para("总疑似片段占比", body_style),
                para(f"{fragment_distribution.get('total_suspected_fragment_ratio', 0)}%", body_style),
            ],
            [
                para("重度疑似片段", body_style),
                para(
                    f"{fragment_distribution.get('severe_fragment_count', 0)}（{fragment_distribution.get('severe_fragment_ratio', 0)}%）",
                    body_style,
                ),
            ],
        ]
        story.append(build_table(aux_rows, col_widths=[34 * mm, 144 * mm]))
        report_summary = result.get("report_summary") or {}
        if report_summary.get("available"):
            metric_rows = [["辅助报告指标", "值"]]
            for metric in (report_summary.get("metrics") or [])[:6]:
                metric_rows.append(
                    [
                        str(metric.get("label") or "指标"),
                        f"{metric.get('value', '-')}{metric.get('unit', '')}",
                    ]
                )
            if len(metric_rows) > 1:
                story.append(Spacer(1, 6))
                story.append(build_table(metric_rows, col_widths=[72 * mm, 106 * mm], header=True))
        if section_distribution and not document_outline:
            story.append(Spacer(1, 6))
            section_rows = [["章节区段", "段落数", "均值", "高风险段落"]]
            for item in section_distribution:
                section_rows.append(
                    [
                        str(item.get("section") or "-"),
                        str(item.get("paragraph_count") or 0),
                        f"{item.get('avg_score', 0)}%",
                        str(item.get("high_count") or 0),
                    ]
                )
            story.append(build_table(section_rows, col_widths=[54 * mm, 30 * mm, 30 * mm, 44 * mm], header=True))

        story.append(Spacer(1, 8))
        story.append(para("说明：", section_style))
        story.extend(
            [
                para("1. 本报告参考知网公开报告的栏目结构组织全文结果，但不是知网官方出具报告。", body_style),
                para("2. AI特征值属于概率研判结果，应结合人工复核、学校制度与正式送检结果综合理解。", body_style),
                para("3. 红色/棕色等官方标色不直接复刻，当前版本以结构化文本与段落指标替代。", body_style),
            ]
        )

        def draw_page(canvas, _doc):
            width, height = A4
            canvas.setFont("STSong-Light", 8)
            canvas.setFillColor(colors.HexColor("#666666"))
            canvas.drawString(_doc.leftMargin, height - 8 * mm, "格物学术 AIGC检测仿真报告")
            canvas.drawRightString(width - _doc.rightMargin, 8 * mm, f"第 {_doc.page} 页")

        doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
        return buffer.getvalue()

    def _render_detect_report_pdf_fallback(self, result: dict) -> bytes:
        lines = self._build_detect_report_fallback_lines(result)
        pages = self._paginate_pdf_lines(lines, lines_per_page=40)
        return self._build_text_pdf_document(pages)

    def _build_detect_report_fallback_lines(self, result: dict) -> list[str]:
        meta = self._current_task_report_meta()
        stats = result.get("source_stats") or {}
        distribution = result.get("distribution") or {}
        fragment_distribution = result.get("fragment_distribution") or {}
        document_metrics = result.get("document_metrics") or {}
        decision_basis = result.get("decision_basis") or []
        view = result.get("report_view") or self._build_detect_report_view(result)
        band_rows = view.get("band_rows") or self._build_detect_band_rows(result.get("paragraph_details") or [])
        display_segments = view.get("display_segments") or []
        display_paragraphs = view.get("display_paragraphs") or []
        headline_metrics = view.get("headline_metrics") or []
        secondary_metrics = view.get("secondary_metrics") or []
        distribution_metrics = view.get("distribution_metrics") or []

        lines: list[str] = []

        def add_line(text: str = "") -> None:
            if not text:
                lines.append("")
                return
            lines.extend(self._wrap_pdf_line(text))

        def add_kv(label: str, value) -> None:
            value_text = str(value or "").strip()
            if not value_text:
                return
            add_line(f"{label}: {value_text}")

        add_line(str(view.get("brand_title") or "格物学术 AIGC 检测报告"))
        add_line(f"{view.get('platform_badge') or result.get('provider_label') or result.get('platform')} | 简洁报告 + 全文报告")
        add_line()

        add_line("一、基本信息")
        add_kv("报告编号", result.get("report_no"))
        add_kv("生成时间", result.get("generated_at"))
        add_kv("检测平台", result.get("provider_label") or result.get("platform"))
        add_kv("处理模式", result.get("mode"))
        add_kv("篇名", meta.get("paper_title"))
        add_kv("作者", meta.get("authors"))
        add_kv("文件名", meta.get("source_filename"))
        add_line()

        add_line("二、简洁报告")
        for metric in headline_metrics:
            add_kv(metric.get("label"), metric.get("value"))
        add_kv("总段落数", stats.get("paragraph_count", 0))
        add_line(str(result.get("summary") or "").strip())
        add_line()

        if secondary_metrics:
            add_line("三、辅助指标")
            for metric in secondary_metrics:
                add_kv(metric.get("label"), metric.get("value"))
            add_line()

        add_line("四、区段与平台指标")
        add_kv("高风险段落", f"{distribution.get('high_count', 0)}（{distribution.get('high_ratio', 0)}%）")
        add_kv("中高风险文字占比", f"{fragment_distribution.get('high_and_middle_suspected_text_ratio', 0)}%")
        add_kv("总疑似文字占比", f"{fragment_distribution.get('total_suspected_text_ratio', 0)}%")
        add_kv("总疑似片段占比", f"{fragment_distribution.get('total_suspected_fragment_ratio', 0)}%")
        for metric in distribution_metrics:
            label = str(metric.get("label") or "-").strip()
            value = str(metric.get("value") or "-").strip()
            detail = str(metric.get("detail") or "").strip()
            if detail:
                add_line(f"- {label}: {value} | {detail}")
            else:
                add_line(f"- {label}: {value}")
        for item in band_rows:
            add_line(
                f"- {item['label']}: 段落 {item['paragraph_count']} | 均值 {item['avg_score']}% | "
                f"高风险 {item['high_count']} | 字符 {item['char_count']}"
            )
        add_line()

        add_line("五、全文报告")
        if decision_basis:
            add_line("判定依据")
            for item in decision_basis[:8]:
                direction = str(item.get("direction") or "").strip().lower()
                prefix = "减权" if direction == "relief" else "风险"
                title = str(item.get("title") or "-").strip()
                detail = self._clip_text(str(item.get("detail") or "-"), 120)
                add_line(f"- {prefix} | {title}: {detail}")
            add_line()

        add_line("六、文档指标")
        add_kv("中高风险段落占比", f"{document_metrics.get('high_medium_paragraph_ratio', 0)}%")
        add_kv("中高风险文字占比", f"{document_metrics.get('high_medium_text_ratio', 0)}%")
        add_kv("最长连续风险段落", document_metrics.get("longest_risk_streak", 0))
        add_kv("段首相似度", f"{document_metrics.get('opening_similarity_ratio', 0)}%")
        add_kv("可疑章节覆盖率", f"{document_metrics.get('section_coverage_ratio', 0)}%")
        add_kv("实证减权", f"{document_metrics.get('evidence_relief_pct', 0)}%")
        add_line()

        add_line("七、重点疑似片段")
        if display_segments:
            for idx, item in enumerate(display_segments[:10], start=1):
                reason_text = str(item.get("reason") or "综合风险偏高")
                add_line(
                    f"{idx}. P{int(item.get('paragraph_index') or 0):02d} | "
                    f"风险 {item.get('score', 0)}% | "
                    f"{reason_text}"
                )
                add_line(self._clip_text(str(item.get("text") or ""), 180))
                add_line()
        else:
            add_line(str(view.get("detail_hint") or "当前未展开正文级疑似片段。"))
            add_line()

        add_line("八、段落明细")
        if display_paragraphs:
            for item in display_paragraphs[:24]:
                reason_tags = " / ".join(item.get("reason_tags") or []) or "综合语义判定"
                add_line(
                    f"P{int(item.get('index') or 0):02d} | "
                    f"{item.get('risk_band') or '-'} | "
                    f"{item.get('score', 0)}% | "
                    f"字符 {item.get('char_count', 0)} | "
                    f"句子 {item.get('sentence_count', 0)}"
                )
                add_line(f"信号: {reason_tags}")
                add_line(f"摘要: {self._clip_text(str(item.get('excerpt') or ''), 180)}")
                add_line()
        else:
            add_line("当前报告未展开段落级明细。")
            add_line()

        add_line("九、说明")
        add_line("1) 本报告为格物学术生成的仿真检测结果，参考公开报告的栏目结构组织，不等同于官方报告。")
        add_line("2) 结果适用于内部研判、人工复核与版本对比，建议结合正式送检结果综合判断。")
        add_line(f"3) {view.get('report_note') or '报告样式由格物学术统一设计。'}")
        return lines

    def _wrap_pdf_line(self, text: str, max_units: float = 34.0, width: float | None = None) -> list[str]:
        if width is not None:
            max_units = float(width)
        compact = " ".join(str(text or "").split())
        if not compact:
            return [""]
        lines: list[str] = []
        current: list[str] = []
        width = 0.0
        for char in compact:
            char_width = self._pdf_char_width(char)
            if current and width + char_width > max_units:
                lines.append("".join(current).rstrip())
                current = []
                width = 0.0
                if char == " ":
                    continue
            current.append(char)
            width += char_width
        if current:
            lines.append("".join(current).rstrip())
        return lines or [""]

    def _pdf_char_width(self, char: str) -> float:
        if not char:
            return 0.0
        if char in " .,;:!'|`ilI[]()":
            return 0.35
        if ord(char) < 128:
            return 0.62
        return 1.0

    def _paginate_pdf_lines(self, lines: list[str], lines_per_page: int) -> list[list[str]]:
        if lines_per_page <= 0:
            return [lines]
        if not lines:
            return [[""]]
        pages: list[list[str]] = []
        for start in range(0, len(lines), lines_per_page):
            page = lines[start : start + lines_per_page]
            pages.append(page if page else [""])
        return pages or [[""]]

    def _build_text_pdf_document(self, pages: list[list[str]]) -> bytes:
        objects: dict[int, bytes] = {
            1: b"<< /Type /Catalog /Pages 2 0 R >>",
            3: (
                b"<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light "
                b"/Encoding /UniGB-UCS2-H /DescendantFonts [4 0 R] >>"
            ),
            4: (
                b"<< /Type /Font /Subtype /CIDFontType0 /BaseFont /STSong-Light "
                b"/CIDSystemInfo << /Registry (Adobe) /Ordering (GB1) /Supplement 4 >> "
                b"/DW 1000 >>"
            ),
        }
        page_ids: list[int] = []
        next_id = 5
        total_pages = max(len(pages), 1)

        for page_number, page_lines in enumerate(pages or [[""]], start=1):
            stream = self._build_text_pdf_page_stream(page_lines, page_number, total_pages)
            objects[next_id] = b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream"
            content_id = next_id
            next_id += 1
            objects[next_id] = (
                b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                b"/Resources << /Font << /F1 3 0 R >> >> /Contents "
                + str(content_id).encode("ascii")
                + b" 0 R >>"
            )
            page_ids.append(next_id)
            next_id += 1

        kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
        objects[2] = f"<< /Type /Pages /Count {len(page_ids)} /Kids [{kids}] >>".encode("ascii")

        max_id = max(objects)
        pdf = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0] * (max_id + 1)
        for object_id in range(1, max_id + 1):
            offsets[object_id] = len(pdf)
            pdf.extend(f"{object_id} 0 obj\n".encode("ascii"))
            pdf.extend(objects[object_id])
            pdf.extend(b"\nendobj\n")

        xref_offset = len(pdf)
        pdf.extend(f"xref\n0 {max_id + 1}\n".encode("ascii"))
        pdf.extend(b"0000000000 65535 f \n")
        for object_id in range(1, max_id + 1):
            pdf.extend(f"{offsets[object_id]:010d} 00000 n \n".encode("ascii"))
        pdf.extend(
            (
                f"trailer\n<< /Size {max_id + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_offset}\n%%EOF"
            ).encode("ascii")
        )
        return bytes(pdf)

    def _build_text_pdf_page_stream(self, lines: list[str], page_number: int, total_pages: int) -> bytes:
        header = "\u683c\u7269\u5b66\u672f AIGC\u68c0\u6d4b\u62a5\u544a"
        footer = f"\u7b2c {page_number}/{total_pages} \u9875"
        body_lines = lines or [""]
        stream_lines = [
            "BT",
            "/F1 9 Tf",
            "1 0 0 1 42 818 Tm",
            f"<{self._pdf_hex_text(header)}> Tj",
            "ET",
            "BT",
            "/F1 10 Tf",
            "14 TL",
            "1 0 0 1 42 792 Tm",
        ]
        for line in body_lines:
            stream_lines.append(f"<{self._pdf_hex_text(line or ' ')}> Tj")
            stream_lines.append("T*")
        stream_lines.extend(
            [
                "ET",
                "BT",
                "/F1 8 Tf",
                "1 0 0 1 520 18 Tm",
                f"<{self._pdf_hex_text(footer)}> Tj",
                "ET",
            ]
        )
        return "\n".join(stream_lines).encode("ascii")

    def _pdf_hex_text(self, text: str) -> str:
        payload = str(text or " ").encode("utf-16-be", errors="ignore")
        return payload.hex().upper() or "0020"

    # Override AIGC detection logic with the latest multi-signal rules.
    def _platform_detect_profile_legacy_v2(self, platform: str) -> dict:
        key = (platform or "").strip().lower()
        profiles = {
            "cnki": {
                "name": "cnki_like",
                "provider_label": "仿知网",
                "score_label": "AI特征值",
                "baseline_weight": 0.56,
                "style_weight": 0.14,
                "repeat_weight": 0.12,
                "template_weight": 0.10,
                "context_weight": 0.08,
                "opening_weight": 0.06,
                "offset": 0.0,
                "high": 0.67,
                "medium": 0.42,
                "coverage_weight": 0.06,
                "section_weight": 0.08,
                "streak_weight": 0.03,
                "opening_similarity_weight": 0.02,
                "evidence_relief_weight": 0.06,
                "colloquial_relief_weight": 0.42,
                "specificity_relief_weight": 0.30,
                "fragment_display_thresholds": {"mild": 0.70, "moderate": 0.80, "severe": 0.90},
            },
            "vip": {
                "name": "vip_like",
                "provider_label": "仿维普",
                "score_label": "全文疑似AIGC生成",
                "baseline_weight": 0.52,
                "style_weight": 0.16,
                "repeat_weight": 0.12,
                "template_weight": 0.11,
                "context_weight": 0.09,
                "opening_weight": 0.07,
                "offset": -0.02,
                "high": 0.64,
                "medium": 0.40,
                "coverage_weight": 0.08,
                "section_weight": 0.08,
                "streak_weight": 0.03,
                "opening_similarity_weight": 0.02,
                "evidence_relief_weight": 0.05,
                "colloquial_relief_weight": 0.48,
                "specificity_relief_weight": 0.34,
                "fragment_display_thresholds": {"mild": 0.70, "moderate": 0.80, "severe": 0.90},
            },
        }
        return profiles.get(key, profiles["cnki"])
