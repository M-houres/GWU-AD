import random
import re
import shutil
import string
import subprocess
from os import environ
from pathlib import Path
import zipfile

from docx import Document
from pypdf import PdfReader

def make_order_no() -> str:
    prefix = "OD"
    rand = "".join(random.choices(string.digits, k=12))
    return f"{prefix}{rand}"


def gen_code() -> str:
    return "".join(random.choices(string.digits, k=6))


def is_phone_valid(phone: str) -> bool:
    return bool(re.fullmatch(r"1\d{10}", phone))


def safe_filename(name: str) -> str:
    allow = set(string.ascii_letters + string.digits + "._-")
    return "".join(ch if ch in allow else "_" for ch in name)


def safe_display_filename(name: str) -> str:
    value = str(name or "").strip().replace("\\", "/").split("/")[-1]
    value = re.sub(r"[\x00-\x1f\x7f]", "", value)
    value = value.replace(":", "_").replace("*", "_").replace("?", "_")
    value = value.replace("\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")
    value = value.strip().strip(".")
    return value or "unnamed"


_DOCX_REQUIRED_MEMBERS = {"[Content_Types].xml", "word/document.xml"}
_DOC_REQUIRED_HEADER = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def detect_file_magic(path: Path) -> str:
    with path.open("rb") as f:
        head = f.read(16)
    if head.startswith(b"%PDF-"):
        with path.open("rb") as f:
            try:
                f.seek(-2048, 2)
            except OSError:
                f.seek(0)
            tail = f.read()
        return ".pdf" if b"%%EOF" in tail else ""
    if head.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(path, "r") as zf:
                members = {name for name in zf.namelist() if name and not name.endswith("/")}
        except zipfile.BadZipFile:
            return ""
        return ".docx" if _DOCX_REQUIRED_MEMBERS.issubset(members) else ""
    if head.startswith(_DOC_REQUIRED_HEADER):
        return ".doc"
    try:
        raw = path.read_bytes()
        if not raw.strip():
            return ""
        if b"\x00" in raw[:4096]:
            return ""
        raw.decode("utf-8")
        return ".txt"
    except Exception:
        return ""
    return ""


def _resolve_antiword_map_file() -> str | None:
    antiword_home = str(environ.get("ANTIWORDHOME") or "").strip()
    candidates = []
    if antiword_home:
        candidates.append(Path(antiword_home) / "UTF-8.txt")
    candidates.extend(
        [
            Path("/usr/share/antiword/UTF-8.txt"),
            Path("/usr/local/share/antiword/UTF-8.txt"),
        ]
    )
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def _extract_text_from_doc(path: Path) -> str:
    antiword_bin = shutil.which("antiword")
    if not antiword_bin:
        raise ValueError("legacy .doc extractor unavailable")

    command = [antiword_bin]
    map_file = _resolve_antiword_map_file()
    if map_file:
        command.extend(["-m", map_file])
    command.append(str(path))
    result = subprocess.run(command, capture_output=True, check=False)
    if result.returncode != 0:
        message = result.stderr.decode("utf-8", errors="ignore").strip() or "antiword failed"
        raise ValueError(message)
    return result.stdout.decode("utf-8", errors="ignore").replace("\ufeff", "").strip()


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".doc":
        return _extract_text_from_doc(path)
    if suffix == ".docx":
        doc = Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        return "\n".join(parts)
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)
    raise ValueError("unsupported file type")


def count_billable_chars(text: str) -> int:
    filtered = [ch for ch in text if ch.isalnum() or ("\u4e00" <= ch <= "\u9fff")]
    return len(filtered)
